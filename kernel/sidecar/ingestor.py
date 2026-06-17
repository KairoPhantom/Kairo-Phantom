"""
Kairo Phantom — Ingestor (SPEC §S2, §S4)

Ingests document files and produces Chunks with non-null page + bbox.
Uses text extraction for .txt/.md files and structured parsing for
.pdf/.docx/.xlsx/.pptx files.

INVARIANT: Every Chunk has non-null page and bbox after ingestion.

The kernel imports NOTHING from /domains or /legacy.
"""

from __future__ import annotations

import hashlib
import logging
import pathlib
import re
from datetime import datetime, timezone
from dataclasses import replace

from kernel.core.data_model import BBox, Chunk, Document, Page

logger = logging.getLogger(__name__)

# Approximate characters per line for bbox estimation in text files
_CHARS_PER_LINE = 80
_LINES_PER_PAGE = 60


class IngestorImpl:
    """Concrete Ingestor implementation.

    Ingests files and produces Chunks with guaranteed page + bbox.
    Supports: .txt, .md, .pdf (via Docling or fitz fallback), .docx (via python-docx).
    """

    def ingest(self, path: str) -> tuple[list[Chunk], Document, list[Page]]:
        """Ingest a document file and return Chunks, Document, and Pages.

        Every Chunk MUST have non-null page and bbox.
        Raises FileNotFoundError if the file doesn't exist.
        Raises ValueError if the file type is unsupported.
        """
        filepath = pathlib.Path(path)
        if not filepath.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        # Compute document metadata
        content_bytes = filepath.read_bytes()
        sha256 = hashlib.sha256(content_bytes).hexdigest()
        suffix = filepath.suffix.lower()

        if suffix in (".txt", ".md"):
            chunks, page_count, pages = self._ingest_text(filepath, sha256)
        elif suffix == ".pdf":
            chunks, page_count, pages = self._ingest_pdf(filepath, sha256)
        elif suffix == ".docx":
            chunks, page_count, pages = self._ingest_docx(filepath, sha256)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        # Create the Document
        doc = Document(
            source_path=str(filepath),
            sha256=sha256,
            page_count=page_count,
            ingested_at=datetime.now(timezone.utc),
        )

        # Set doc_id on all chunks and pages and verify invariants
        updated_chunks: list[Chunk] = []
        for chunk in chunks:
            updated_chunk = replace(chunk, doc_id=doc.doc_id)
            if updated_chunk.bbox is None:
                raise RuntimeError(
                    f"Ingestor invariant violated: chunk {updated_chunk.chunk_id} has no bbox"
                )
            if updated_chunk.page < 1:
                raise RuntimeError(
                    f"Ingestor invariant violated: chunk {updated_chunk.chunk_id} has page={updated_chunk.page} (must be >= 1)"
                )
            updated_chunks.append(updated_chunk)

        updated_pages: list[Page] = []
        for page in pages:
            updated_page = replace(page, doc_id=doc.doc_id)
            updated_pages.append(updated_page)

        logger.info(
            "Ingested %s: %d chunks, %d pages, sha256=%s",
            filepath.name,
            len(updated_chunks),
            page_count,
            sha256[:16],
        )

        return updated_chunks, doc, updated_pages

    def _ingest_text(
        self, filepath: pathlib.Path, sha256: str
    ) -> tuple[list[Chunk], int, list[Page]]:
        """Ingest a text/markdown file into chunks by paragraph."""
        text = filepath.read_text(encoding="utf-8", errors="replace")
        paragraphs = self._split_paragraphs(text)

        chunks: list[Chunk] = []
        current_line = 0
        page_count = 1

        for para_text in paragraphs:
            if not para_text.strip():
                current_line += 1
                continue

            para_lines = para_text.count("\n") + 1
            page = (current_line // _LINES_PER_PAGE) + 1
            page_count = max(page_count, page)

            # Estimate bbox based on line position within page
            page_offset = current_line % _LINES_PER_PAGE
            y0 = page_offset * (1.0 / _LINES_PER_PAGE)
            y1 = min(1.0, (page_offset + para_lines) * (1.0 / _LINES_PER_PAGE))

            chunk = Chunk(
                page=page,
                bbox=BBox(x0=0.0, y0=y0, x1=1.0, y1=y1),
                text=para_text.strip(),
                source_type="text",
            )
            chunks.append(chunk)
            current_line += para_lines + 1  # +1 for blank line between paragraphs

        # Generate Page objects
        pages: list[Page] = []
        for p in range(1, page_count + 1):
            pages.append(Page(
                doc_id="",
                index=p,
                width_px=800,
                height_px=1000,
                image_sha256="",
            ))

        return chunks, page_count, pages

    def _ingest_pdf(
        self, filepath: pathlib.Path, sha256: str
    ) -> tuple[list[Chunk], int, list[Page]]:
        """Ingest a PDF file. Tries Docling first, then fitz (PyMuPDF)."""
        # Try Docling first
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(str(filepath))
            doc_docling = result.document

            page_count = len(doc_docling.pages) if hasattr(doc_docling, "pages") else 0
            chunks: list[Chunk] = []
            pages: list[Page] = []

            # Create Page objects
            for index, page_obj in doc_docling.pages.items():
                width = 800
                height = 1000
                if hasattr(page_obj, "size"):
                    width = int(page_obj.size.width)
                    height = int(page_obj.size.height)
                pages.append(Page(
                    doc_id="",
                    index=index,
                    width_px=width,
                    height_px=height,
                    image_sha256="",
                ))

            # Accessing elements
            elements = []
            if hasattr(doc_docling, "texts"):
                elements.extend(doc_docling.texts)
            if hasattr(doc_docling, "tables"):
                elements.extend(doc_docling.tables)
            if not elements and hasattr(doc_docling, "elements"):
                elements = list(doc_docling.elements)

            for item in elements:
                if hasattr(item, "prov") and item.prov:
                    page_info = item.prov[0]
                    page_no = page_info.page_no
                    bbox = page_info.bbox

                    text = ""
                    if hasattr(item, "text"):
                        text = item.text
                    elif hasattr(item, "export_to_markdown"):
                        text = item.export_to_markdown()

                    text = text.strip()
                    if not text:
                        continue

                    # Get page dimensions to normalize coordinates
                    page_width = 1.0
                    page_height = 1.0
                    if page_no in doc_docling.pages:
                        p_obj = doc_docling.pages[page_no]
                        if hasattr(p_obj, "size"):
                            page_width = p_obj.size.width
                            page_height = p_obj.size.height

                    x0 = max(0.0, min(bbox.left / page_width if page_width > 0 else 0, 1.0))
                    y0 = max(0.0, min(bbox.top / page_height if page_height > 0 else 0, 1.0))
                    x1 = max(x0, min(bbox.right / page_width if page_width > 0 else 1, 1.0))
                    y1 = max(y0, min(bbox.bottom / page_height if page_height > 0 else 1, 1.0))

                    chunk = Chunk(
                        page=page_no,
                        bbox=BBox(x0=x0, y0=y0, x1=x1, y1=y1),
                        text=text,
                        source_type="pdf_docling",
                    )
                    chunks.append(chunk)

            if chunks:
                return chunks, page_count, pages
        except Exception as e:
            logger.warning("Docling failed to parse PDF, trying PyMuPDF: %s", e)

        # Fallback to PyMuPDF
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF (fitz) not available — falling back to text extraction")
            chunks, page_count, pages = self._ingest_text(filepath, sha256)
            return chunks, page_count, pages

        doc = fitz.open(str(filepath))
        chunks = []
        pages = []
        page_count = len(doc)

        # Ensure page images dir exists
        page_images_dir = pathlib.Path(".kairo/page_images")
        page_images_dir.mkdir(parents=True, exist_ok=True)

        for page_num in range(page_count):
            page = doc[page_num]
            page_rect = page.rect

            # Save page image for click-to-source UX
            image_sha = ""
            try:
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                image_sha = hashlib.sha256(img_data).hexdigest()
                image_path = page_images_dir / f"{image_sha}.png"
                if not image_path.exists():
                    image_path.write_bytes(img_data)
            except Exception as e:
                logger.warning("Failed to render page image: %s", e)

            pages.append(Page(
                doc_id="",
                index=page_num + 1,
                width_px=int(page_rect.width),
                height_px=int(page_rect.height),
                image_sha256=image_sha,
            ))

            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:  # text blocks only
                    continue

                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                    block_text += "\n"

                block_text = block_text.strip()
                if not block_text:
                    continue

                # Normalize bbox to [0,1] range
                bbox = block["bbox"]
                chunk = Chunk(
                    page=page_num + 1,
                    bbox=BBox(
                        x0=max(0.0, min(bbox[0] / page_rect.width, 1.0)),
                        y0=max(0.0, min(bbox[1] / page_rect.height, 1.0)),
                        x1=max(0.0, min(bbox[2] / page_rect.width, 1.0)),
                        y1=max(0.0, min(bbox[3] / page_rect.height, 1.0)),
                    ),
                    text=block_text,
                    source_type="pdf_text",
                )
                chunks.append(chunk)

        doc.close()
        return chunks, page_count, pages

    def _ingest_docx(
        self, filepath: pathlib.Path, sha256: str
    ) -> tuple[list[Chunk], int, list[Page]]:
        """Ingest a DOCX file. Uses python-docx if available."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning(
                "python-docx not available — falling back to text extraction"
            )
            return self._ingest_text(filepath, sha256)

        doc = DocxDocument(str(filepath))
        chunks: list[Chunk] = []
        current_line = 0
        page_count = 1

        for para in doc.paragraphs:
            if not para.text.strip():
                current_line += 1
                continue

            para_lines = max(1, len(para.text) // _CHARS_PER_LINE + 1)
            page = (current_line // _LINES_PER_PAGE) + 1
            page_count = max(page_count, page)

            page_offset = current_line % _LINES_PER_PAGE
            y0 = page_offset * (1.0 / _LINES_PER_PAGE)
            y1 = min(1.0, (page_offset + para_lines) * (1.0 / _LINES_PER_PAGE))

            chunk = Chunk(
                page=page,
                bbox=BBox(x0=0.0, y0=y0, x1=1.0, y1=y1),
                text=para.text.strip(),
                source_type="docx_paragraph",
            )
            chunks.append(chunk)
            current_line += para_lines + 1

        # Generate Page objects
        pages: list[Page] = []
        for p in range(1, page_count + 1):
            pages.append(Page(
                doc_id="",
                index=p,
                width_px=800,
                height_px=1000,
                image_sha256="",
            ))

        return chunks, page_count, pages

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        """Split text into paragraphs by double newline or single newline
        when lines start with numbered/bulleted items."""
        # Split on double newlines first
        blocks = re.split(r"\n\s*\n", text)
        return [b for b in blocks if b.strip()]
