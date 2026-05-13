"""
setup_fixtures.py — Creates all test documents required by tests-scenario.md
Run ONCE before starting the test loop.
"""
import os, random, shutil
from pathlib import Path

BASE = Path("C:/tests")
BASE.mkdir(exist_ok=True)
(BASE / "logs").mkdir(exist_ok=True)
(BASE / "screenshots").mkdir(exist_ok=True)
(BASE / "results").mkdir(exist_ok=True)

print("Creating test fixtures in C:\\tests\\...")

# ── 1. report_informal.docx ──────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Q3 Performance Review", 0)
    doc.add_paragraph(
        "we gotta improve our numbers cuz theyre not looking good lol. "
        "The team did alright but we need way more customers."
    )
    doc.add_paragraph(
        "honestly its kinda frustrating cuz everyone worked super hard but "
        "the results arent there yet. gonna need to figure this out asap."
    )
    doc.save(BASE / "report_informal.docx")
    print("  ✅ report_informal.docx")
except Exception as e:
    print(f"  ❌ report_informal.docx: {e}")

# ── 2. contract.docx (NDA template) ──────────────────────────────────────────
try:
    from docx import Document
    doc = Document()
    doc.add_heading("NON-DISCLOSURE AGREEMENT", 0)
    clauses = [
        ("1. PARTIES", "This Agreement is entered into between Company A (\"Disclosing Party\") and Company B (\"Receiving Party\")."),
        ("2. DEFINITION OF CONFIDENTIAL INFORMATION", "Confidential Information means any data or information that is proprietary to the Disclosing Party and not generally known to the public."),
        ("3. OBLIGATIONS OF RECEIVING PARTY", "The Receiving Party agrees to hold Confidential Information in strict confidence and not to disclose it to any third party."),
        ("4. EXCLUSIONS", "This Agreement does not apply to information that is publicly known, or that was known to the Receiving Party before disclosure."),
        ("5. TERM", "This Agreement shall remain in effect for a period of three (3) years from the date of signing."),
        ("6. REMEDIES", "The Receiving Party acknowledges that breach of this Agreement may cause irreparable harm for which monetary damages would be inadequate."),
        ("7. GOVERNING LAW", "This Agreement shall be governed by the laws of the State of Delaware."),
        ("8. ENTIRE AGREEMENT", "This Agreement constitutes the entire agreement between the parties with respect to its subject matter."),
    ]
    for title, text in clauses:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
    doc.save(BASE / "contract.docx")
    print("  ✅ contract.docx")
except Exception as e:
    print(f"  ❌ contract.docx: {e}")

# ── 3. blank_deck.pptx ───────────────────────────────────────────────────────
try:
    from pptx import Presentation
    prs = Presentation()
    prs.save(BASE / "blank_deck.pptx")
    print("  ✅ blank_deck.pptx")
except Exception as e:
    print(f"  ❌ blank_deck.pptx: {e}")

# ── 4. spreadsheet_broken.xlsx ───────────────────────────────────────────────
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales Data"
    ws.append(["Date", "Product", "Region", "Revenue", "Units", "Margin"])
    ws.append(["2026-01-15", "Widget A", "North", 15000, 300, "=E2/D2"])  # valid
    ws.append(["2026-01-20", "Widget B", "South", 0, 150, "=D3/E3"])      # #DIV/0!
    ws.append(["2026-02-01", "Widget C", "East", 12000, "#REF!", "=D4*E4"])  # #VALUE!
    ws.append(["2026-02-15", "Widget A", "West", 9000, 180, "=D5/0"])     # #DIV/0!
    ws.append(["Jan 15 2026", "Widget B", "North", 18000, 360, "=E6/D6"]) # bad date
    ws.append(["2026-03-01", "Widget C", "South", 14000, 280, "=D7/E7"])  # valid
    wb.save(BASE / "spreadsheet_broken.xlsx")
    print("  ✅ spreadsheet_broken.xlsx")
except Exception as e:
    print(f"  ❌ spreadsheet_broken.xlsx: {e}")

# ── 5. vscode-project/ (TypeScript) ──────────────────────────────────────────
try:
    proj = BASE / "vscode-project"
    proj.mkdir(exist_ok=True)
    (proj / "utils.ts").write_text("""
export function formatCurrency(amount: number, currency: string = 'USD'): string {
  return new Intl.NumberFormat('en-US', { style: 'currency', currency }).format(amount);
}

export function validateEmail(email: string): boolean {
  const re = /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/;
  return re.test(email);
}

export function debounce<T extends (...args: any[]) => any>(fn: T, ms: number): T {
  let timer: ReturnType<typeof setTimeout>;
  return ((...args: any[]) => {
    clearTimeout(timer);
    timer = setTimeout(() => fn(...args), ms);
  }) as T;
}
""", encoding="utf-8")
    (proj / "api.ts").write_text("""
import { validateEmail } from './utils';

interface User {
  id: number;
  name: string;
  email: string;
}

// Function that fetches user data from API, validates the response,
// and returns typed User object
export async function fetchUser(id: number): Promise<User> {
  const response = await fetch(`/api/users/${id}`);
  if (!response.ok) throw new Error(`HTTP ${response.status}`);
  const data = await response.json();
  if (!validateEmail(data.email)) throw new Error('Invalid user email');
  return data as User;
}
""", encoding="utf-8")
    (proj / "main.ts").write_text("""
import { fetchUser } from './api';
import { formatCurrency } from './utils';

async function main() {
  const user = await fetchUser(1);
  console.log(`User: ${user.name} (${user.email})`);
}

main().catch(console.error);
""", encoding="utf-8")
    (proj / "tsconfig.json").write_text('{"compilerOptions":{"target":"ES2020","module":"commonjs","strict":true}}', encoding="utf-8")
    print("  ✅ vscode-project/")
except Exception as e:
    print(f"  ❌ vscode-project/: {e}")

# ── 6. vscode-buggy/ (Python with bugs) ──────────────────────────────────────
try:
    buggy = BASE / "vscode-buggy"
    buggy.mkdir(exist_ok=True)
    (buggy / "buggy.py").write_text("""
def calculate_average(numbers):
    total = 0
    # Bug 1: off-by-one - should be range(len(numbers))
    for i in range(len(numbers) + 1):
        total = total + numbers[i]
    return total / len(numbers)

def find_duplicates(items):
    seen = []
    duplicates = []
    for item in items:
        if item in seen:
            duplicates.append(item)
        seen.append(item)
    # Bug 2: should return duplicates, not seen
    return seen

def process_data(data):
    result = []
    for x in data:
        if x > 0:
            y = x * 2
        elif x == 0:
            y = 1
        result.append(y)  # Bug 3: y not defined if x < 0 (scope)
    return result
""", encoding="utf-8")
    print("  ✅ vscode-buggy/")
except Exception as e:
    print(f"  ❌ vscode-buggy/: {e}")

# ── 7. Ensure existing fixtures are correct ───────────────────────────────────
if not (BASE / "notes.txt").exists():
    (BASE / "notes.txt").write_text(
        "Meeting Notes - Team Sync 2026-05-14\n"
        "Attendees: Sandip, Arjun, Priya\n"
        "Action items:\n"
        "- Launch Kairo Phantom by end of month\n"
        "- Fix memory persistence bug\n"
        "- Review investor deck\n",
        encoding="utf-8"
    )
    print("  ✅ notes.txt")

print("\n✅ All fixtures created. Ready to start testing.")
print(f"   C:\\tests\\ contents: {[f.name for f in BASE.iterdir()]}")
