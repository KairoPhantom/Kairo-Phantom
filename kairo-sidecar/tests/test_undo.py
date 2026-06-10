import os
import sys
import tempfile
import shutil
import pytest
import asyncio
from pathlib import Path
from docx import Document

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.humanized_injector import HumanizedInjector

@pytest.fixture
def temp_docx():
    """Create a temporary Word doc."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    doc.add_heading("Production Gate Document", level=1)
    doc.add_paragraph("Adjacent paragraph 1.")
    doc.add_paragraph("Target paragraph.")
    doc.add_paragraph("Adjacent paragraph 2.")
    doc.save(file_path)
    
    yield file_path
    
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        # remove temp directory
        os.rmdir(temp_dir)
    except Exception:
        pass

def test_ctrl_z_undo_restores_original(temp_docx):
    injector = HumanizedInjector()
    
    original_size = os.path.getsize(temp_docx)
    
    # Read original text
    doc_orig = Document(temp_docx)
    original_text = [p.text for p in doc_orig.paragraphs]
    
    # Start injection session (creates snapshot of file)
    loop = asyncio.new_event_loop()
    session = loop.run_until_complete(injector.begin_session(temp_docx))
    loop.close()
    
    # Modify doc
    doc_mod = Document(temp_docx)
    doc_mod.add_paragraph("New paragraph injected.")
    doc_mod.save(temp_docx)
    
    assert os.path.getsize(temp_docx) != original_size
    
    # Trigger rollback
    success = injector.undo_injection(session)
    assert success is True
    
    # Verify exact contents restored
    doc_restored = Document(temp_docx)
    restored_text = [p.text for p in doc_restored.paragraphs]
    assert restored_text == original_text
