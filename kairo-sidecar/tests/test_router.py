"""
tests/test_router.py — Domain Master Router Integration Tests
=============================================================

Covers the 10 scenarios specified in the implementation prompt:

Test 1:  Word request → WordMaster called (not ExcelMaster)
Test 2:  Ambiguous prompt → clarification returned (not operations)
Test 3:  Quality failure → retry with feedback
Test 4:  MemMachine query called with correct domain tag
Test 5:  MemMachine record called after successful operation
Test 6:  Browser + Notion URL → Yjs CRDT path taken
Test 7:  PDF domain → PDFMaster called, output is new file creation op
Test 8:  Code domain Python file → CodeMaster called (tree-sitter context)
Test 9:  End-to-end Word: prompt → operations → write → verify file
Test 10: End-to-end Excel: prompt → formula → validate → write → verify formula

Plus the existing backward-compatible legacy API tests (Tests 11-23).
"""

import os
import sys
import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch, call, PropertyMock

from docx import Document
import openpyxl

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.router import (
    DomainMasterRouter, KairoRequest, KairoResponse,
    ReasoningStep, OutputVerifier, QualityReport,
)
from sidecar.masters.word_master import WordMaster
from sidecar.masters.excel_master import ExcelMaster


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def temp_docx():
    """Create a temporary Word doc with real content."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("First paragraph text.")
    doc.add_paragraph("Second paragraph with some content.")
    doc.save(file_path)
    yield file_path
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def temp_xlsx():
    """Create a temporary Excel workbook with real data."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Product"
    ws["B1"] = "Revenue"
    ws["C1"] = "Cost"
    ws["A2"] = "Widget A"
    ws["B2"] = 10000.0
    ws["C2"] = 6000.0
    wb.save(file_path)
    yield file_path
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def router():
    """DomainMasterRouter instance."""
    return DomainMasterRouter()


# ---------------------------------------------------------------------------
# Helper — build a mock LLM response
# ---------------------------------------------------------------------------

def mock_llm_response(operations=None, confidence=0.95, reasoning="Test"):
    r = MagicMock()
    r.operations = operations or []
    r.confidence = confidence
    r.reasoning = reasoning
    r.model_dump.return_value = {
        "operations": [
            op.model_dump() if hasattr(op, "model_dump") else op
            for op in (operations or [])
        ],
        "confidence": confidence,
        "reasoning": reasoning,
    }
    return r


# ===========================================================================
# TEST 1: Word request → WordMaster called, NOT ExcelMaster
# ===========================================================================

def test_route_word_uses_word_master_not_excel(router, temp_docx):
    """
    route(KairoRequest(domain='word')) must dispatch to WordMaster.
    ExcelMaster.extract_context must NOT be called.
    """
    mock_resp = mock_llm_response()

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        with patch.object(
            router.masters["excel"], "extract_context",
            side_effect=AssertionError("ExcelMaster should NOT be called for a Word request!"),
        ):
            req = KairoRequest(
                user_prompt="write a memo summary",
                domain="word",
                file_path=temp_docx,
                cursor_info=0,
            )
            resp = router.route(req)

    assert resp.type == "operations", f"Expected 'operations', got '{resp.type}': {resp.error}"
    assert resp.domain == "word"


# ===========================================================================
# TEST 2: Ambiguous prompt → clarification returned (not operations)
# ===========================================================================

def test_route_ambiguous_returns_clarification(router, temp_docx):
    """
    When the orchestrator classifies a request as is_ambiguous=True with
    confidence < 0.55, route() must return a 'clarification' response.
    """
    from sidecar.schemas.domain_schemas import OrchestratorResponse

    mock_classification = OrchestratorResponse(
        domain="word",
        confidence=0.40,
        is_ambiguous=True,
        clarifying_question="Do you want to insert a new paragraph or replace an existing one?",
        ambiguity_reason="Target paragraph not specified",
    )

    with patch.object(
        router.reasoning_step, "classify", return_value=mock_classification
    ):
        req = KairoRequest(
            user_prompt="fix it",
            domain="word",
            file_path=temp_docx,
        )
        resp = router.route(req)

    assert resp.type == "clarification", f"Expected 'clarification', got '{resp.type}'"
    assert "insert" in resp.question.lower() or "paragraph" in resp.question.lower()
    assert resp.confidence < 0.55


# ===========================================================================
# TEST 3: Quality failure → retry with feedback
# ===========================================================================

def test_route_quality_failure_triggers_retry(router, temp_docx):
    """
    When quality gates fail on the first LLM call, route() must call
    call_with_schema a SECOND time (the retry) with the quality issues appended.
    """
    # First response — will fail quality (contains a leaked keyword)
    bad_resp = MagicMock()
    bad_resp.operations = []
    bad_resp.confidence = 0.9
    bad_resp.model_dump.return_value = {
        "operations": [],
        "confidence": 0.9,
        "reasoning": "waza_agent leaked",   # triggers leakage quality check
    }

    # Second (retry) response — clean
    good_resp = mock_llm_response()

    call_count = {"n": 0}

    def side_effect(prompt, schema, model="ollama/qwen2.5:7b"):
        call_count["n"] += 1
        if call_count["n"] == 1:
            return bad_resp
        return good_resp

    with patch("sidecar.router.call_with_schema", side_effect=side_effect):
        req = KairoRequest(
            user_prompt="write a summary",
            domain="word",
            file_path=temp_docx,
        )
        resp = router.route(req)

    # Must have called LLM at least twice (initial + retry)
    assert call_count["n"] >= 2, (
        f"Expected at least 2 LLM calls (initial + retry), got {call_count['n']}"
    )
    assert resp.type in ("operations", "error")


# ===========================================================================
# TEST 4: MemMachine query called with correct domain tag
# ===========================================================================

def test_route_queries_mem_machine_with_correct_domain(router, temp_docx):
    """
    route() must call mem_machine.query() with domain='word' for a Word request.
    """
    mock_resp = mock_llm_response()
    query_calls = []

    original_query = router.mem_machine.query if router.mem_machine else None

    def capture_query(**kwargs):
        query_calls.append(kwargs)
        return ""

    if router.mem_machine:
        router.mem_machine.query = capture_query

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="write a heading",
            domain="word",
            file_path=temp_docx,
            user_id="test_user_42",
        )
        router.route(req)

    if router.mem_machine:
        assert len(query_calls) >= 1, "mem_machine.query() was never called"
        assert query_calls[0].get("domain") == "word", (
            f"Expected domain='word', got {query_calls[0].get('domain')!r}"
        )
        # Restore
        router.mem_machine.query = original_query


# ===========================================================================
# TEST 5: MemMachine record called after successful operation
# ===========================================================================

def test_route_records_to_mem_machine_on_success(router, temp_docx):
    """
    After a successful route(), mem_machine.record_interaction() must be called
    with the correct domain.
    """
    mock_resp = mock_llm_response()
    record_calls = []

    original_record = router.mem_machine.record_interaction if router.mem_machine else None

    def capture_record(**kwargs):
        record_calls.append(kwargs)
        return True

    if router.mem_machine:
        router.mem_machine.record_interaction = capture_record

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="add a summary paragraph",
            domain="word",
            file_path=temp_docx,
        )
        resp = router.route(req)

    if router.mem_machine:
        assert len(record_calls) >= 1, "mem_machine.record_interaction() was never called"
        assert record_calls[0].get("domain") == "word", (
            f"Expected domain='word', got {record_calls[0].get('domain')!r}"
        )
        router.mem_machine.record_interaction = original_record

    assert resp.type == "operations"


# ===========================================================================
# TEST 6: Browser + Notion URL → Yjs CRDT path taken
# ===========================================================================

def test_route_browser_notion_yjs_crdt(router):
    """
    Browser domain with a Notion URL must route to BrowserMaster.
    The response injection_method should indicate crdt_yjs.
    """
    mock_resp = MagicMock()
    mock_resp.operations = []
    mock_resp.confidence = 0.95
    mock_resp.model_dump.return_value = {
        "injection_method": "crdt_yjs",
        "content": "Notion page content here",
        "is_collaborative_editor": True,
        "safety_check": {
            "is_password_field": False,
            "is_payment_field": False,
            "is_auto_submit": False,
        },
        "confidence": 0.95,
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="write a project summary",
            domain="browser",
            file_path="https://www.notion.so/myworkspace/project-doc-abc123",
        )
        resp = router.route(req)

    assert resp.type == "operations"
    assert resp.domain == "browser"
    # The raw_data must carry the injection_method
    assert resp.raw_data.get("injection_method") == "crdt_yjs"


# ===========================================================================
# TEST 7: PDF domain → PDFMaster called, output contains new file creation op
# ===========================================================================

def test_route_pdf_creates_file_op(router):
    """
    PDF domain routing must call PDFMaster and the response should describe
    a new file creation operation (output_filename present in raw_data).
    """
    mock_resp = MagicMock()
    mock_resp.operations = []
    mock_resp.confidence = 0.9
    mock_resp.model_dump.return_value = {
        "output_type": "docx",
        "operations": [],
        "output_filename": "converted_output.docx",
        "confidence": 0.9,
        "extraction_quality": "high",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="convert this PDF to Word",
            domain="pdf",
            file_path="/path/to/report.pdf",
        )
        resp = router.route(req)

    assert resp.type == "operations"
    assert resp.domain == "pdf"
    assert "converted_output.docx" in resp.raw_data.get("output_filename", ""), (
        "output_filename missing from PDF response"
    )


# ===========================================================================
# TEST 8: Code domain Python file → CodeMaster called with language context
# ===========================================================================

def test_route_code_domain_python_file(router):
    """
    Code domain for a .py file must route to CodeMaster.
    The response must be type='operations' and domain='code'.
    """
    mock_resp = MagicMock()
    mock_resp.operations = []
    mock_resp.confidence = 0.93
    mock_resp.reasoning = "Added docstring to main function"
    mock_resp.model_dump.return_value = {
        "operations": [],
        "confidence": 0.93,
        "reasoning": "Added docstring to main function",
    }

    # Patch CodeMaster.extract_context to confirm it receives the Python file
    file_paths_seen = []

    original_extract = router.masters["code"].extract_context

    def capture_extract(file_path, cursor_info=None):
        file_paths_seen.append(file_path)
        return original_extract(file_path, cursor_info)

    router.masters["code"].extract_context = capture_extract

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="add a docstring to main()",
            domain="code",
            file_path="main.py",
            cursor_info=10,
        )
        resp = router.route(req)

    router.masters["code"].extract_context = original_extract  # restore

    assert resp.type == "operations"
    assert resp.domain == "code"
    assert resp.raw_data.get("confidence") == 0.93
    # CodeMaster received the correct file path
    assert file_paths_seen and file_paths_seen[0] == "main.py"


# ===========================================================================
# TEST 9: End-to-end Word: prompt → operations → write → verify file
# ===========================================================================

def test_route_word_end_to_end_write_file(router, temp_docx):
    """
    Full end-to-end Word pipeline:
    - Mock LLM returns an insert_paragraph operation
    - route() validates and returns the op
    - We then apply_operations() to the file
    - Open the saved file and verify the new paragraph exists
    """
    from sidecar.schemas.docx_schema import DocxResponse, InsertParagraphOp, ParagraphRun

    new_para_text = "E2E_KAIRO_TEST_PARAGRAPH_UNIQUE_STRING"

    # Mock LLM — returns a proper DocxResponse with one insert_paragraph op
    insert_op = InsertParagraphOp(
        type="insert_paragraph",
        after_paragraph_index=0,
        style="Normal",
        runs=[ParagraphRun(text=new_para_text, bold=False, italic=False)],
    )
    mock_resp = DocxResponse(
        operations=[insert_op],
        confidence=0.95,
        reasoning="E2E test paragraph insertion",
    )

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt=f"write a test paragraph: {new_para_text}",
            domain="word",
            file_path=temp_docx,
            cursor_info=0,
        )
        resp = router.route(req)

    assert resp.type == "operations", f"Expected 'operations', got: {resp.type}, error: {resp.error}"
    assert len(resp.operations) == 1, f"Expected 1 op, got {len(resp.operations)}"

    # Apply the validated ops to the file
    apply_result = router.masters["word"].apply_operations(temp_docx, resp.operations)
    assert apply_result.get("applied_count", 0) >= 1, (
        f"apply_operations failed: {apply_result}"
    )

    # Verify the paragraph now exists in the file
    saved_doc = Document(temp_docx)
    texts = [p.text for p in saved_doc.paragraphs]
    assert any(new_para_text in t for t in texts), (
        f"E2E paragraph '{new_para_text}' not found in saved document.\n"
        f"All paragraphs: {texts}"
    )


# ===========================================================================
# TEST 10: End-to-end Excel: prompt → formula → validate → write → verify
# ===========================================================================

def test_route_excel_end_to_end_write_formula(router, temp_xlsx):
    """
    Full end-to-end Excel pipeline:
    - Mock LLM returns a write_cell operation with a real SUM formula
    - route() validates the formula
    - We then apply_operations() to the .xlsx file
    - Reload and verify the formula cell contains =SUM(B2:B5)
    """
    from sidecar.schemas.xlsx_schema import ExcelResponse, WriteCellOp

    target_formula = "=SUM(B2:B5)"
    target_cell = "D2"

    write_op = WriteCellOp(
        type="write_cell",
        sheet="Sheet",
        cell=target_cell,
        formula=target_formula,
        value=None,
    )
    mock_resp = ExcelResponse(
        operations=[write_op],
        confidence=0.97,
        reasoning="E2E Excel formula write test",
    )

    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt=f"write SUM formula for revenue in {target_cell}",
            domain="excel",
            file_path=temp_xlsx,
            cursor_info=target_cell,
        )
        resp = router.route(req)

    assert resp.type == "operations", (
        f"Expected 'operations', got '{resp.type}': {resp.error}"
    )

    # Apply to file
    apply_result = router.masters["excel"].apply_operations(temp_xlsx, resp.operations)
    assert apply_result.get("applied_count", 0) >= 1 or apply_result.get("errors") == [], (
        f"apply_operations returned unexpected result: {apply_result}"
    )

    # Verify formula written to file
    wb = openpyxl.load_workbook(temp_xlsx, data_only=False)
    ws = wb.active
    cell_value = ws[target_cell].value
    assert cell_value == target_formula, (
        f"Expected formula '{target_formula}' in {target_cell}, got '{cell_value}'"
    )


# ===========================================================================
# BACKWARD-COMPATIBLE LEGACY TESTS (Tests 11-23)
# These exercise route_llm_request() to ensure nothing broke.
# ===========================================================================

@pytest.mark.anyio
async def test_route_word_request(temp_docx):
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test reasoning"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.95,
        "reasoning": "Test reasoning",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response) as mock_call:
        res = await router.route_llm_request(
            domain="word",
            file_path=temp_docx,
            user_instruction="write a memo summary",
            mem_context="Use bullet points",
            cursor_info=0,
        )
        assert res["ok"] is True
        assert res["data"]["confidence"] == 0.95
        assert mock_call.called


@pytest.mark.anyio
async def test_route_excel_request(temp_xlsx):
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.9
    mock_response.reasoning = "Excel total formula"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.9,
        "reasoning": "Excel total formula",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response) as mock_call:
        res = await router.route_llm_request(
            domain="excel",
            file_path=temp_xlsx,
            user_instruction="calculate sum",
            mem_context="No preference",
            cursor_info="A2",
        )
        assert res["ok"] is True
        assert res["data"]["confidence"] == 0.9
        assert mock_call.called


@pytest.mark.anyio
async def test_invalid_domain_fails(temp_docx):
    router = DomainMasterRouter()
    res = await router.route_llm_request(
        domain="unsupported_domain_xyz",
        file_path=temp_docx,
        user_instruction="test",
        mem_context="",
        cursor_info=0,
    )
    assert res["ok"] is False
    assert "Unsupported routing domain" in res["error"]


@pytest.mark.anyio
async def test_mem_machine_context_integration(temp_docx):
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test"
    mock_response.model_dump.return_value = {"operations": [], "confidence": 0.95, "reasoning": "Test"}

    with patch("sidecar.router.call_with_schema", return_value=mock_response) as mock_call:
        await router.route_llm_request(
            domain="word",
            file_path=temp_docx,
            user_instruction="write a memo",
            mem_context="USER_MEM_MACHINE_STYLE_GUIDE",
            cursor_info=0,
        )
        passed_prompt = mock_call.call_args[0][0]
        assert "USER_MEM_MACHINE_STYLE_GUIDE" in passed_prompt


@pytest.mark.anyio
async def test_excel_grid_context_generation(temp_xlsx):
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test"
    mock_response.model_dump.return_value = {"operations": [], "confidence": 0.95, "reasoning": "Test"}

    with patch("sidecar.router.call_with_schema", return_value=mock_response) as mock_call:
        await router.route_llm_request(
            domain="excel",
            file_path=temp_xlsx,
            user_instruction="write a total formula",
            mem_context="",
            cursor_info="A2",
        )
        passed_prompt = mock_call.call_args[0][0]
        assert "Excel Context" in passed_prompt or "DOCUMENT CONTEXT" in passed_prompt
        assert "A2" in passed_prompt


@pytest.mark.anyio
async def test_word_style_fuzzy_matching_via_router(temp_docx):
    router = DomainMasterRouter()
    mock_op = MagicMock()
    mock_op.model_dump.return_value = {
        "type": "insert_paragraph",
        "after_paragraph_index": 0,
        "style": "Heading1",
        "runs": [{"text": "Conformed heading text", "bold": False, "italic": False}],
    }

    mock_response = MagicMock()
    mock_response.operations = [mock_op]
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test style matching"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.95,
        "reasoning": "Test style matching",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        res = await router.route_llm_request(
            domain="word",
            file_path=temp_docx,
            user_instruction="write a heading",
            mem_context="",
            cursor_info=0,
        )
        assert res["ok"] is True
        assert res["data"]["operations"][0]["style"] == "Heading 1"


@pytest.mark.anyio
async def test_excel_circular_reference_rejected_via_router(temp_xlsx):
    router = DomainMasterRouter()
    mock_op = MagicMock()
    mock_op.model_dump.return_value = {
        "type": "write_cell",
        "cell": "A2",
        "formula": "=SUM(A1:A2)",
    }

    mock_response = MagicMock()
    mock_response.operations = [mock_op]
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test circular reference"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.95,
        "reasoning": "Test circular reference",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        res = await router.route_llm_request(
            domain="excel",
            file_path=temp_xlsx,
            user_instruction="write formula",
            mem_context="",
            cursor_info="A2",
        )
        assert res["ok"] is True
        assert len(res["data"]["operations"]) == 0


@pytest.mark.anyio
async def test_excel_locale_adjusted_via_router(temp_xlsx):
    router = DomainMasterRouter()
    mock_op = MagicMock()
    mock_op.model_dump.return_value = {
        "type": "write_cell",
        "cell": "B2",
        "formula": "=SUM(A1,A2)",
    }

    mock_response = MagicMock()
    mock_response.operations = [mock_op]
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test locale conversion"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.95,
        "reasoning": "Test locale conversion",
    }

    mock_context = MagicMock()
    mock_context.locale = "eu"
    mock_context.active_cell = "B2"
    mock_context.cells = []

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        with patch.object(router.excel_extractor, "extract", return_value=mock_context):
            res = await router.route_llm_request(
                domain="excel",
                file_path=temp_xlsx,
                user_instruction="sum values",
                mem_context="",
                cursor_info="B2",
            )
            assert res["ok"] is True
            assert res["data"]["operations"][0]["formula"] == "=SUM(A1;A2)"


@pytest.mark.anyio
async def test_router_error_handling(temp_docx):
    router = DomainMasterRouter()
    with patch.object(router.word_extractor, "extract", side_effect=RuntimeError("Hard disk error")):
        res = await router.route_llm_request(
            domain="word",
            file_path=temp_docx,
            user_instruction="write a memo",
            mem_context="",
            cursor_info=0,
        )
        assert res["ok"] is False
        assert "Routing failed" in res["error"]


@pytest.mark.anyio
async def test_router_maps_file_extensions_correctly(temp_docx):
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test"
    mock_response.model_dump.return_value = {"operations": [], "confidence": 0.95, "reasoning": "Test"}

    with patch("sidecar.router.call_with_schema", return_value=mock_response) as mock_call:
        res = await router.route_llm_request(
            domain="docx",
            file_path=temp_docx,
            user_instruction="write a summary",
            mem_context="",
            cursor_info=0,
        )
        assert res["ok"] is True
        assert mock_call.called


@pytest.mark.anyio
async def test_browser_notion_yjs_path():
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test"
    mock_response.model_dump.return_value = {
        "injection_method": "crdt_yjs",
        "content": "Notion text",
        "is_collaborative_editor": True,
        "safety_check": {"is_password_field": False, "is_payment_field": False, "is_auto_submit": False},
        "confidence": 0.95,
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        res = await router.route_llm_request(
            domain="browser",
            file_path="https://www.notion.so/workspace/doc1",
            user_instruction="write a summary",
            mem_context="",
            cursor_info=0,
        )
        assert res["ok"] is True
        assert res["data"]["injection_method"] == "crdt_yjs"


@pytest.mark.anyio
async def test_pdf_domain_creates_new_file_op():
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.9
    mock_response.model_dump.return_value = {
        "output_type": "docx",
        "operations": [],
        "output_filename": "output_doc.docx",
        "confidence": 0.9,
        "extraction_quality": "high",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        res = await router.route_llm_request(
            domain="pdf",
            file_path="test_file.pdf",
            user_instruction="convert to word",
            mem_context="",
            cursor_info=0,
        )
        assert res["ok"] is True
        assert res["data"]["output_filename"] == "output_doc.docx"


@pytest.mark.anyio
async def test_code_domain_called():
    router = DomainMasterRouter()
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.92
    mock_response.reasoning = "Code edit"
    mock_response.model_dump.return_value = {
        "operations": [],
        "confidence": 0.92,
        "reasoning": "Code edit",
    }

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        res = await router.route_llm_request(
            domain="code",
            file_path="main.py",
            user_instruction="add docstring",
            mem_context="",
            cursor_info=10,
        )
        assert res["ok"] is True
        assert res["data"]["confidence"] == 0.92


# ===========================================================================
# Additional edge-case tests for the new route() API
# ===========================================================================

def test_route_normalises_docx_domain(router, temp_docx):
    """'docx' domain alias must be normalised to 'word'."""
    mock_resp = mock_llm_response()
    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="add a summary",
            domain="docx",
            file_path=temp_docx,
        )
        resp = router.route(req)
    assert resp.domain == "word"


def test_route_normalises_xlsx_domain(router, temp_xlsx):
    """'xlsx' domain alias must be normalised to 'excel'."""
    mock_resp = mock_llm_response()
    with patch("sidecar.router.call_with_schema", return_value=mock_resp):
        req = KairoRequest(
            user_prompt="calculate sum",
            domain="xlsx",
            file_path=temp_xlsx,
        )
        resp = router.route(req)
    assert resp.domain == "excel"


def test_route_error_returns_error_response(router, temp_docx):
    """Extraction failure must return KairoResponse(type='error') not raise."""
    with patch.object(
        router.masters["word"], "extract_context",
        side_effect=OSError("File not found"),
    ):
        req = KairoRequest(
            user_prompt="write something",
            domain="word",
            file_path=temp_docx,
        )
        resp = router.route(req)
    assert resp.type == "error"
    assert "File not found" in resp.error or len(resp.error) > 0


def test_word_master_facade_imports():
    """WordMaster must be importable and expose the required interface."""
    from sidecar.masters.word_master import WordMaster
    m = WordMaster()
    assert hasattr(m, "extract_context")
    assert hasattr(m, "build_prompt")
    assert hasattr(m, "validate_operations")
    assert hasattr(m, "apply_operations")
    assert hasattr(m, "get_schema_class")


def test_excel_master_facade_imports():
    """ExcelMaster must be importable and expose the required interface."""
    from sidecar.masters.excel_master import ExcelMaster
    m = ExcelMaster()
    assert hasattr(m, "extract_context")
    assert hasattr(m, "build_prompt")
    assert hasattr(m, "validate_operations")
    assert hasattr(m, "apply_operations")
    assert hasattr(m, "get_schema_class")


def test_kairo_request_defaults():
    """KairoRequest default values must be sensible."""
    req = KairoRequest(user_prompt="hello", domain="word")
    assert req.user_id == "local"
    assert req.file_path == ""
    assert req.cursor_info is None


def test_kairo_response_defaults():
    """KairoResponse default values must be sensible."""
    resp = KairoResponse()
    assert resp.type == "operations"
    assert resp.operations == []
    assert resp.error == ""


def test_output_verifier_detects_leakage():
    """OutputVerifier must flag system prompt keyword leakage."""
    ov = OutputVerifier()
    report = ov.run_all_checks(
        output='{"text": "waza_agent leaked here"}',
        domain="word",
    )
    assert not report.all_passed
    assert any("leakage" in i.lower() for i in report.issues)


def test_output_verifier_passes_clean_response():
    """OutputVerifier must pass a clean, non-empty JSON response."""
    ov = OutputVerifier()
    report = ov.run_all_checks(
        output='{"operations": [{"type": "insert_paragraph"}], "confidence": 0.95}',
        domain="word",
    )
    assert report.all_passed


def test_output_verifier_detects_empty_response():
    """OutputVerifier must flag an empty response."""
    ov = OutputVerifier()
    report = ov.run_all_checks(output="", domain="word")
    assert not report.all_passed
    assert any("empty" in i.lower() for i in report.issues)


def test_output_verifier_legacy_task_kwarg():
    """OutputVerifier.run_all_checks() must accept old-style task= kwarg."""
    ov = OutputVerifier()
    # Should not raise even when `task` is provided instead of `domain`
    report = ov.run_all_checks(
        output='{"operations": [], "confidence": 0.9}',
        task=None,
    )
    assert isinstance(report, QualityReport)


def test_config_kairo_toml_exists():
    """config/kairo.toml must exist and contain db_path."""
    config_path = Path(__file__).parent.parent / "config" / "kairo.toml"
    assert config_path.exists(), f"config/kairo.toml not found at {config_path}"
    content = config_path.read_text(encoding="utf-8")
    assert "db_path" in content, "kairo.toml must contain db_path setting"


def test_memmachine_db_path_from_config():
    """MemMachineClient must resolve a valid db_path (from env, config, or default)."""
    from sidecar.mem_machine import MemMachineClient
    client = MemMachineClient()
    assert client.db_path, "db_path must not be empty"
    # The path must point to an existing directory (DB file may not exist yet)
    parent = os.path.dirname(client.db_path)
    os.makedirs(parent, exist_ok=True)
    assert os.path.isdir(parent), f"DB parent directory doesn't exist: {parent}"
