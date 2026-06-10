import os
import sys
import tempfile
import shutil
import pytest
from pathlib import Path
from docx import Document
import openpyxl

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.masters.word.writer import WordWriter
from sidecar.masters.word.context_extractor import WordContextExtractor
from sidecar.masters.excel_master import ExcelWriter

@pytest.fixture
def temp_docx():
    """Create a temporary Word doc."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    doc.add_paragraph("Paragraph 0 (unchanged)")
    doc.add_paragraph("Paragraph 1 (unchanged)")
    doc.add_paragraph("Paragraph 2 (target)")
    doc.add_paragraph("Paragraph 3 (unchanged)")
    doc.save(file_path)
    yield file_path
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        os.rmdir(temp_dir)
    except Exception:
        pass

@pytest.fixture
def temp_xlsx():
    """Create a temporary Excel workbook."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Revenue"
    ws["A2"] = 500.0
    ws["B1"] = "Expenses"
    ws["B2"] = 300.0
    ws["C1"] = "Net"
    wb.save(file_path)
    yield file_path
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        os.rmdir(temp_dir)
    except Exception:
        pass

def test_adjacent_elements_unchanged(temp_docx, temp_xlsx):
    # 1. Test Word adjacent paragraph preservation
    doc_before = Document(temp_docx)
    para1_text_before = doc_before.paragraphs[1].text
    para3_text_before = doc_before.paragraphs[3].text
    
    writer = WordWriter()
    # Edit paragraph 2 (target)
    writer._edit_paragraph(doc_before, 2, "Paragraph 2 (EDITED)")
    writer.atomic_save_docx(doc_before, temp_docx)
    
    doc_after = Document(temp_docx)
    assert doc_after.paragraphs[2].text == "Paragraph 2 (EDITED)"
    assert doc_after.paragraphs[1].text == para1_text_before
    assert doc_after.paragraphs[3].text == para3_text_before

    # 2. Test Excel adjacent cells preservation
    excel_writer = ExcelWriter()
    op = {
        "type": "write_cell",
        "cell": "B2",
        "value": 450.0
    }
    excel_writer.apply_operations(temp_xlsx, [op])
    
    wb = openpyxl.load_workbook(temp_xlsx)
    ws = wb.active
    assert ws["B2"].value == 450.0
    assert ws["A2"].value == 500.0
    assert ws["A1"].value == "Revenue"
