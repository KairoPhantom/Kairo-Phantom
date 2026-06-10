"""
tests/test_e2e_docx.py — E2E Word Document Tests
==================================================
Tests WordWriter.apply_operations() and WordContextExtractor.extract() directly.
No LLM or named pipe required.
"""

import os
import sys
import tempfile
import pytest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.masters.word_master import WordWriter, WordContextExtractor, WordContext


# ---------------------------------------------------------------------------
# Helper: Create a temp docx with optional paragraphs and headings
# ---------------------------------------------------------------------------

def make_temp_docx(paragraphs=None, headings=None):
    doc = Document()
    if headings:
        for h_text, h_level in headings:
            doc.add_heading(h_text, level=h_level)
    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def get_context(path):
    """Use WordContextExtractor to get a real context from the file."""
    return WordContextExtractor().extract(path, 0)


# ---------------------------------------------------------------------------
# W-01: insert_paragraph with style='Heading 2'
# ---------------------------------------------------------------------------

def test_w01_insert_paragraph_heading2():
    """insert_paragraph with style='Heading 2' → inserted para has style 'Heading 2'."""
    path = make_temp_docx(paragraphs=["First paragraph", "Second paragraph"])
    try:
        context = get_context(path)
        ops = [
            {
                "type": "insert_paragraph",
                "after_paragraph_index": 0,
                "style": "Heading 2",
                "runs": [{"text": "My New Heading", "bold": False, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("applied_count", 0) >= 1 or result.get("errors") == [], (
            f"apply_operations failed: {result}"
        )

        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 2" in styles, (
            f"Heading 2 style not found. Styles found: {styles}"
        )
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        # Clean up atomic backup if exists
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-02: insert_paragraph with style='List Bullet'
# ---------------------------------------------------------------------------

def test_w02_insert_paragraph_list_bullet():
    """insert_paragraph with style='List Bullet' → style name contains 'Bullet'."""
    path = make_temp_docx(paragraphs=["Paragraph A", "Paragraph B"])
    try:
        context = get_context(path)

        # Check if 'List Bullet' style exists in document, otherwise use a fallback
        bullet_styles = [s for s in context.styles["paragraph"] if "bullet" in s.lower() or "Bullet" in s]
        style_to_use = bullet_styles[0] if bullet_styles else "List Bullet"

        ops = [
            {
                "type": "insert_paragraph",
                "after_paragraph_index": 0,
                "style": style_to_use,
                "runs": [{"text": "Bullet item text", "bold": False, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("errors", []) == [] or result.get("applied_count", 0) >= 1, (
            f"apply_operations errors: {result.get('errors')}"
        )

        doc = Document(path)
        style_names = [p.style.name for p in doc.paragraphs]
        # The inserted paragraph should have a bullet-like style
        found_bullet = any("bullet" in s.lower() or "Bullet" in s for s in style_names)
        assert found_bullet, (
            f"No bullet style found in paragraphs. Styles: {style_names}"
        )
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-03: replace_paragraph → verify surrounding paragraphs unchanged
# ---------------------------------------------------------------------------

def test_w03_replace_paragraph_surrounding_unchanged():
    """replace_paragraph → surrounding paragraphs are unchanged."""
    path = make_temp_docx(paragraphs=["Keep me before", "Replace me", "Keep me after"])
    try:
        context = get_context(path)

        # Determine index of "Replace me" (may be offset by doc metadata paragraphs)
        doc_before = Document(path)
        replace_idx = None
        for i, p in enumerate(doc_before.paragraphs):
            if p.text == "Replace me":
                replace_idx = i
                break
        assert replace_idx is not None, "Could not find 'Replace me' paragraph"

        ops = [
            {
                "type": "replace_paragraph",
                "paragraph_index": replace_idx,
                "style": "Normal",
                "runs": [{"text": "Replaced content", "bold": False, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("errors") == [] or result.get("applied_count", 0) >= 1

        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert "Keep me before" in texts, f"'Keep me before' missing. Texts: {texts}"
        assert "Keep me after" in texts, f"'Keep me after' missing. Texts: {texts}"
        assert "Replace me" not in texts, f"'Replace me' still present. Texts: {texts}"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-04: insert_table with 3 headers, 2 rows
# ---------------------------------------------------------------------------

def test_w04_insert_table_3_headers_2_rows():
    """insert_table with 3 headers and 2 rows → table exists with correct dimensions."""
    path = make_temp_docx(paragraphs=["Before table", "After table"])
    try:
        context = get_context(path)

        ops = [
            {
                "type": "insert_table",
                "after_paragraph_index": 0,
                "headers": ["Name", "Age", "City"],
                "rows": [
                    ["Alice", "30", "London"],
                    ["Bob", "25", "Paris"],
                ],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("errors") == [] or result.get("applied_count", 0) >= 1, (
            f"insert_table failed: {result}"
        )

        doc = Document(path)
        assert len(doc.tables) >= 1, "No table found in document"
        table = doc.tables[0]
        # Header row + 2 data rows = 3 total rows
        assert len(table.rows) == 3, f"Expected 3 rows (header + 2), got {len(table.rows)}"
        assert len(table.columns) == 3, f"Expected 3 cols, got {len(table.columns)}"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-05: delete_paragraph → verify paragraph count decreases by 1
# ---------------------------------------------------------------------------

def test_w05_delete_paragraph_count_decreases():
    """delete_paragraph → paragraph count decreases by exactly 1."""
    path = make_temp_docx(paragraphs=["Para One", "Para Two", "Para Three"])
    try:
        doc_before = Document(path)
        count_before = len(doc_before.paragraphs)

        # Find Para Two
        del_idx = None
        for i, p in enumerate(doc_before.paragraphs):
            if p.text == "Para Two":
                del_idx = i
                break
        assert del_idx is not None, "Could not find 'Para Two'"

        context = get_context(path)
        ops = [
            {
                "type": "delete_paragraph",
                "paragraph_index": del_idx,
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("applied_count", 0) >= 1 or result.get("errors") == [], (
            f"delete_paragraph failed: {result}"
        )

        doc_after = Document(path)
        count_after = len(doc_after.paragraphs)
        assert count_after == count_before - 1, (
            f"Expected {count_before - 1} paragraphs, got {count_after}"
        )
        texts_after = [p.text for p in doc_after.paragraphs]
        assert "Para Two" not in texts_after, f"'Para Two' still present: {texts_after}"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-06: insert_paragraph after_index=-1 → paragraph appended to end
# ---------------------------------------------------------------------------

def test_w06_insert_paragraph_append_to_end():
    """insert_paragraph with after_paragraph_index=-1 → paragraph appended to end."""
    path = make_temp_docx(paragraphs=["First", "Second", "Third"])
    try:
        context = get_context(path)
        sentinel = "APPENDED_SENTINEL_TEXT"
        ops = [
            {
                "type": "insert_paragraph",
                "after_paragraph_index": -1,
                "style": "Normal",
                "runs": [{"text": sentinel, "bold": False, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("applied_count", 0) >= 1 or result.get("errors") == []

        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert sentinel in texts, f"Sentinel not found in paragraphs. Texts: {texts}"
        # Verify it's at the end (last non-empty paragraph)
        non_empty = [t for t in texts if t.strip()]
        assert non_empty[-1] == sentinel, (
            f"Sentinel not at end. Non-empty texts: {non_empty}"
        )
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-07: Fuzzy style match → 'Heading1' resolves to a valid heading style
# ---------------------------------------------------------------------------

def test_w07_fuzzy_style_match():
    """Fuzzy style match → 'Heading1' resolves to a heading style that exists."""
    from sidecar.masters.word_master import WordOperationValidator

    path = make_temp_docx(paragraphs=["Sample text"])
    try:
        context = get_context(path)
        validator = WordOperationValidator()

        op = {
            "type": "insert_paragraph",
            "after_paragraph_index": 0,
            "style": "Heading1",  # Non-canonical form — should fuzzy-match to 'Heading 1'
            "runs": [{"text": "Fuzzy matched heading", "bold": False, "italic": False}],
        }
        result = validator.validate(op, context)
        assert result.valid, f"Validation failed: {result.error}"
        # Style should have been corrected to a real heading style
        corrected_style = op.get("style", "")
        assert corrected_style in context.styles["paragraph"], (
            f"Corrected style '{corrected_style}' not in available styles"
        )
        assert "heading" in corrected_style.lower(), (
            f"Corrected style '{corrected_style}' does not look like a heading"
        )
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# W-08: Multiple operations in one call → all applied correctly
# ---------------------------------------------------------------------------

def test_w08_multiple_operations_all_applied():
    """Multiple operations in one call → all applied correctly."""
    path = make_temp_docx(paragraphs=["Para A", "Para B", "Para C"])
    try:
        context = get_context(path)
        sentinel1 = "MULTI_OP_SENTINEL_1"
        sentinel2 = "MULTI_OP_SENTINEL_2"

        ops = [
            {
                "type": "insert_paragraph",
                "after_paragraph_index": 0,
                "style": "Normal",
                "runs": [{"text": sentinel1, "bold": False, "italic": False}],
            },
            {
                "type": "insert_paragraph",
                "after_paragraph_index": -1,
                "style": "Normal",
                "runs": [{"text": sentinel2, "bold": False, "italic": False}],
            },
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)

        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert sentinel1 in texts, f"Sentinel 1 not found. Texts: {texts}"
        assert sentinel2 in texts, f"Sentinel 2 not found. Texts: {texts}"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-09: Empty document → insert_paragraph works
# ---------------------------------------------------------------------------

def test_w09_empty_document_insert_paragraph():
    """Empty document → insert_paragraph works without crash."""
    doc = Document()
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    path = tmp.name
    try:
        context = get_context(path)
        sentinel = "EMPTY_DOC_SENTINEL"
        ops = [
            {
                "type": "insert_paragraph",
                "after_paragraph_index": -1,
                "style": "Normal",
                "runs": [{"text": sentinel, "bold": False, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        # Should not raise; applied_count >= 1 or errors == []
        assert isinstance(result, dict), "Result should be a dict"

        doc_after = Document(path)
        texts = [p.text for p in doc_after.paragraphs]
        assert sentinel in texts, f"Sentinel not found in empty doc. Texts: {texts}"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass


# ---------------------------------------------------------------------------
# W-10: extract() on real docx → returns correct paragraph count and styles
# ---------------------------------------------------------------------------

def test_w10_extract_returns_correct_context():
    """extract() on a real docx → returns correct paragraph count and available styles."""
    path = make_temp_docx(
        headings=[("Main Title", 1), ("Section Header", 2)],
        paragraphs=["Body text one", "Body text two", "Body text three"],
    )
    try:
        extractor = WordContextExtractor()
        context = extractor.extract(path, 0)

        # Paragraph count should cover headings + body
        assert context.total_paragraphs >= 5, (
            f"Expected >= 5 paragraphs (2 headings + 3 body), got {context.total_paragraphs}"
        )

        # Styles dict should have paragraph key with at least 'Normal'
        assert "paragraph" in context.styles, "Missing 'paragraph' key in styles"
        assert len(context.styles["paragraph"]) > 0, "Paragraph styles list is empty"
        assert "Normal" in context.styles["paragraph"], (
            f"'Normal' not found in styles: {context.styles['paragraph'][:10]}"
        )

        # Paragraphs list should be populated
        assert len(context.paragraphs) >= 5, (
            f"Expected >= 5 paragraph dicts, got {len(context.paragraphs)}"
        )

        # Check a sample paragraph has required fields
        sample = context.paragraphs[0]
        assert "index" in sample, "Paragraph missing 'index'"
        assert "style" in sample, "Paragraph missing 'style'"
        assert "text" in sample, "Paragraph missing 'text'"
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# W-11: append_to_run → text appended to existing paragraph
# ---------------------------------------------------------------------------

def test_w11_append_to_run():
    """append_to_run → appends text to an existing paragraph."""
    path = make_temp_docx(paragraphs=["Original text"])
    try:
        context = get_context(path)
        ops = [
            {
                "type": "append_to_run",
                "paragraph_index": 0,
                "runs": [{"text": " - appended run", "bold": True, "italic": False}],
            }
        ]
        writer = WordWriter()
        result = writer.apply_operations(path, ops, context)
        assert result.get("errors") == [] or result.get("applied_count", 0) >= 1

        doc = Document(path)
        # Note: python-docx might add metadata paragraphs, find our target
        target_para = None
        for p in doc.paragraphs:
            if "Original text" in p.text:
                target_para = p
                break
        assert target_para is not None
        assert target_para.text == "Original text - appended run"
        # Since it was created as a single run "Original text", adding one more makes it at least 2 runs
        assert len(target_para.runs) >= 2
        assert target_para.runs[-1].bold is True
    finally:
        try:
            os.unlink(path)
        except Exception:
            pass
        for suffix in [".kairo_tmp", ".kairo_bak"]:
            try:
                os.unlink(path + suffix)
            except Exception:
                pass

