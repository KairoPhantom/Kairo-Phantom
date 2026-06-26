import os
import sys
import tempfile
import pytest
from pathlib import Path
from docx import Document
from unittest.mock import patch

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.masters.word.writer import WordWriter


@pytest.fixture
def temp_docx():
    """Create a temporary Word doc."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    doc.add_paragraph("Original paragraph")
    doc.save(file_path)
    yield file_path
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        os.rmdir(temp_dir)
    except Exception:
        pass


def test_atomic_save_crash_safety(temp_docx):
    writer = WordWriter()
    doc = Document(temp_docx)

    # Save original size and modification time
    original_size = os.path.getsize(temp_docx)

    # Simulate a crash during replace
    def faulty_replace(src, dst):
        raise OSError("Simulated crash mid-save")

    with patch("os.replace", side_effect=faulty_replace):
        with pytest.raises(OSError):
            writer.atomic_save_docx(doc, temp_docx)

    # The original file must exist and be completely untouched
    assert os.path.exists(temp_docx)
    assert os.path.getsize(temp_docx) == original_size

    # Verify the contents are still original
    doc_check = Document(temp_docx)
    assert doc_check.paragraphs[0].text == "Original paragraph"
