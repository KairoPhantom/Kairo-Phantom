import os
import sys
import tempfile
import pytest
from pathlib import Path
from docx import Document

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.masters.word.writer import WordWriter
from sidecar.masters.word.validator import WordOperationValidator

@pytest.fixture
def docx_with_custom_style():
    """Create a temporary Word doc with custom style."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "custom_style.docx")
    doc = Document()
    
    # Add a custom paragraph style
    styles = doc.styles
    custom_style = styles.add_style('MyCustomStyle', 1) # 1 = PARAGRAPH
    
    doc.add_paragraph("Paragraph 1", style='MyCustomStyle')
    doc.save(file_path)
    
    yield file_path
    
    # Clean up file
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        os.rmdir(temp_dir)
    except Exception:
        pass

def test_word_style_conformance(docx_with_custom_style):
    doc = Document(docx_with_custom_style)
    available_styles = [s.name for s in doc.styles]
    
    assert "MyCustomStyle" in available_styles
    
    validator = WordOperationValidator()
    matched_style, was_corrected = validator.validate_style("MyCustomStyle", available_styles)
    assert matched_style == "MyCustomStyle"
    assert was_corrected is False
    
    # Test fuzzy style matching
    fuzzy_matched, was_corrected_fuzzy = validator.validate_style("MyCustomstyle", available_styles)
    assert fuzzy_matched == "MyCustomStyle"
    assert was_corrected_fuzzy is True

    # Test writer uses custom style
    writer = WordWriter()
    new_para = writer._insert_paragraph(doc, after_idx=0, text="New custom style paragraph", style_name="MyCustomStyle")
    assert new_para.style.name == "MyCustomStyle"
