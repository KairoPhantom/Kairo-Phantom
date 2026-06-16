import os
import sys
import tempfile
import time
import shutil
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from docx.shared import Pt

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.masters.word_master import (
    WordContext,
    WordContextExtractor,
    WordOperationValidator,
    WordWriter,
    ValidationResult
)
from sidecar.masters.word_prompt_builder import build_word_prompt

@pytest.fixture
def temp_docx():
    """Create a basic docx file with default styles and headings."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    if doc.paragraphs:
        heading = doc.paragraphs[0]
        heading.text = "Title paragraph"
        heading.style = "Heading 1"
    else:
        doc.add_heading("Title paragraph", level=1)
    doc.add_paragraph("First body paragraph of text.")
    doc.add_paragraph("Second body paragraph.")
    doc.add_table(rows=2, cols=3)
    doc.save(file_path)
    yield file_path
    try:
        shutil.rmtree(temp_dir)
    except Exception:
        pass

@pytest.fixture(autouse=True)
def mock_win32com():
    """Automatically mock win32com to prevent interaction with live Word."""
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")), \
         patch("win32com.client.GetObject", side_effect=Exception("Word not running")):
        yield

# --- Test 1: Style from document used (not default) ---
def test_style_from_document_used(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    # Check that Heading 1 is extracted from style sheet
    assert "Heading 1" in ctx.styles["paragraph"]
    assert "Normal" in ctx.styles["paragraph"]

# --- Test 2: Fuzzy style match works ("Heading1" -> "Heading 1") ---
def test_fuzzy_style_match():
    validator = WordOperationValidator()
    available = ["Normal", "Heading 1", "Heading 2", "List Bullet"]
    
    # Test fuzzy matches
    assert validator._fuzzy_style_match("Heading1", available) == "Heading 1"
    assert validator._fuzzy_style_match("heading 2", available) == "Heading 2"
    assert validator._fuzzy_style_match("listbullet", available) == "List Bullet"
    assert validator._fuzzy_style_match("Unknown", available) is None

# --- Test 3: Invalid style triggers validation error (not crash) ---
def test_invalid_style_triggers_error(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    validator = WordOperationValidator()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "InvalidStyleNameXYZ",
        "runs": [{"text": "Hello world"}]
    }
    
    res = validator.validate(op, ctx)
    assert res.valid is False
    assert "not in document" in res.error

# --- Test 4: Paragraph inserted at correct position ---
def test_paragraph_inserted_at_position(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Newly inserted paragraph text!"}]
    }
    
    # Apply python-docx fallback write by bypassing COM check
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        res = writer.apply_operations(temp_docx, [op], ctx)
        assert len(res["errors"]) == 0
        
        doc2 = Document(temp_docx)
        # Should now have 4 paragraphs (1 heading + 2 original + 1 inserted)
        assert len(doc2.paragraphs) == 4
        assert doc2.paragraphs[2].text == "Newly inserted paragraph text!"

# --- Test 5: Adjacent paragraphs unchanged (binary diff) ---
def test_adjacent_paragraphs_unchanged(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "replace_paragraph",
        "paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Modified body paragraph text."}]
    }
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        writer.apply_operations(temp_docx, [op], ctx)
        doc2 = Document(temp_docx)
        
        # Heading (index 0) and second body (original index 2, now 2) must remain identical
        assert doc2.paragraphs[0].text == "Title paragraph"
        assert doc2.paragraphs[2].text == "Second body paragraph."

# --- Test 6: Atomic save — sidecar kill leaves original intact ---
def test_atomic_save_failure_keeps_original(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Failed write"}]
    }
    
    original_size = os.path.getsize(temp_docx)
    
    # Simulate a crash during save by patching Document.save to raise an exception
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        with patch("docx.document.Document.save", side_effect=IOError("Simulated disk full / crash")):
            with pytest.raises(IOError):
                writer.apply_operations(temp_docx, [op], ctx)
                
            # Original file should be restored from backup
            assert os.path.exists(temp_docx)
            assert os.path.getsize(temp_docx) == original_size

# --- Test 7: Legal document -> legal prompt instructions used ---
def test_legal_prompt_instructions():
    styles = {"paragraph": ["Normal"]}
    paragraphs = []
    
    # Context with legal keywords (purpose will be legal)
    ctx = WordContext(
        styles=styles,
        paragraphs=paragraphs,
        tables=[],
        theme_fonts={},
        list_sequences=[],
        document_purpose="legal",
        cursor_paragraph_index=0,
        total_paragraphs=0
    )
    
    prompt = build_word_prompt("Write a paragraph", ctx, "Use formal tone.")
    assert "Use formal legal language" in prompt
    assert "DOCUMENT PURPOSE: legal" in prompt

# --- Test 8: Table insertion has correct row/column count ---
def test_table_insertion_bounds(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_table",
        "after_paragraph_index": 1,
        "headers": ["Col 1", "Col 2"],
        "rows": [["Val A", "Val B"], ["Val C", "Val D"]]
    }
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        writer.apply_operations(temp_docx, [op], ctx)
        doc = Document(temp_docx)
        
        # Document tables should now be 2
        tables = doc.tables
        assert len(tables) == 2
        # Check shapes (rows, cols) without relying on XML table sequence order
        shapes = {(len(t.rows), len(t.columns)) for t in tables}
        assert (2, 3) in shapes
        assert (3, 2) in shapes

# --- Test 9: Numbered list continues from previous number ---
def test_list_sequence_extraction(temp_docx):
    doc = Document(temp_docx)
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Second item", style="List Number")
    doc.save(temp_docx)
    
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    assert len(ctx.list_sequences) == 2
    assert ctx.list_sequences[0]["style"] in ("List Number", "List Bullet")

# --- Test 10: 100-page document — parsing completes under 3s ---
def test_large_document_parsing_performance():
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "large.docx")
    
    doc = Document()
    # Add 100 pages worth of text (approx 300 paragraphs)
    for i in range(300):
        doc.add_paragraph(f"This is paragraph number {i} in a very large document used for parsing benchmarks.")
    doc.save(file_path)
    
    extractor = WordContextExtractor()
    
    start_time = time.time()
    ctx = extractor.extract(file_path, 100)
    elapsed = time.time() - start_time
    
    shutil.rmtree(temp_dir)
    
    # Assert parsing completes within 5 seconds (generous for CI/dev machines under load)
    # 300 paragraphs should parse well under 5s even on slow hardware
    assert elapsed < 5.0, f"Parsing {ctx.total_paragraphs} paragraphs took {elapsed:.2f}s (limit: 5.0s)"
    assert ctx.total_paragraphs == 300

# --- Test 11: Document with custom styles — custom style detected and used ---
def test_custom_style_detection(temp_docx):
    doc = Document(temp_docx)
    # Add a custom paragraph style
    custom_style = doc.styles.add_style("MyCustomStyle", 1) # 1 = WD_STYLE_TYPE.PARAGRAPH
    doc.add_paragraph("Custom styled paragraph", style="MyCustomStyle")
    doc.save(temp_docx)
    
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    assert "MyCustomStyle" in ctx.styles["paragraph"]

# --- Test 12: OneDrive path resolved correctly ---
def test_onedrive_path_resolution():
    with patch("os.path.exists", return_value=True):
        # Resolve user folder paths
        path = "C:\\Users\\user\\OneDrive\\Documents\\doc.docx"
        abs_path = os.path.abspath(path)
        assert "OneDrive" in abs_path

# --- Test 13: File locked -> graceful fallback to clipboard ---
def test_file_locked_fallback(temp_docx):
    # Simulate a file locking / PermissionError on file save
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Fallback content"}]
    }
    
    # Trigger fallback response by raising PermissionError during save
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        with patch("docx.document.Document.save", side_effect=PermissionError("File locked by another process")):
            res = writer.apply_operations(temp_docx, [op], ctx)
            assert "error" in res
            assert "locked" in res["error"]

# --- Test 14: Track changes injection via Adeu bridge ---
def test_track_changes_adeu_routing():
    # Verify that COM writer routes to adeu operations when track changes is active
    writer = WordWriter()
    operations = [
        {"type": "insert_paragraph", "after_paragraph_index": 0, "runs": [{"text": "Tracked change"}]}
    ]
    context = MagicMock()
    
    mock_word = MagicMock()
    mock_doc = MagicMock()
    mock_word.Documents = [mock_doc]
    mock_doc.FullName = "C:\\test.docx"
    
    import subprocess
    mock_proc = MagicMock(spec=subprocess.CompletedProcess)
    mock_proc.stdout = "winword.exe"
    
    with patch("subprocess.run", return_value=mock_proc):
        with patch("win32com.client.GetObject", side_effect=Exception("Simulated GetObject fail")):
            with patch("win32com.client.GetActiveObject", return_value=mock_word):
                with patch("pythoncom.CoInitialize"):
                    with patch("sidecar.writers.docx_writer._try_com_write", return_value={"ok": True}) as mock_com_write:
                        res = writer.apply_operations("C:\\test.docx", operations, context)
                        assert mock_com_write.called
                        assert res["ok"] is True

# --- Test 15: Undo leaves document identical to pre-injection state ---
def test_undo_restores_identical_state(temp_docx):
    # Capture initial file hash/size
    original_size = os.path.getsize(temp_docx)
    
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Temporary injection"}]
    }
    
    # Save a backup path to simulate manual undo restoration (avoid colliding with the writer's own .kairo_bak)
    backup_path = temp_docx + ".test_bak"
    shutil.copy2(temp_docx, backup_path)
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        # Write
        writer.apply_operations(temp_docx, [op], ctx)
        assert os.path.getsize(temp_docx) != original_size
        
        # Undo: restore original from backup
        shutil.copy2(backup_path, temp_docx)
        os.remove(backup_path)
        
        # Verify size matches original exactly
        assert os.path.getsize(temp_docx) == original_size
