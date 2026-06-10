import os
import sys
import time
import socket
import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
import openpyxl

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.router import DomainMasterRouter, OutputVerifier
from sidecar.masters.word_master import WordContextExtractor, WordOperationValidator, WordWriter
from sidecar.masters.excel_master import ExcelContextExtractor, ExcelOperationValidator, ExcelWriter
from sidecar.masters.other_masters import TerminalMaster, EmailMaster
from sidecar.mem_machine import MemMachineClient
from sidecar.schemas.domain_schemas import TerminalResponse, EmailResponse

@pytest.fixture
def temp_docx():
    """Create a temporary Word doc."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.docx")
    doc = Document()
    if doc.paragraphs:
        heading = doc.paragraphs[0]
        heading.text = "Production Gate Document"
        heading.style = "Heading 1"
    else:
        doc.add_heading("Production Gate Document", level=1)
    doc.add_paragraph("Adjacent paragraph 1.")
    doc.add_paragraph("Target paragraph.")
    doc.add_paragraph("Adjacent paragraph 2.")
    doc.save(file_path)
    yield file_path
    try:
        shutil.rmtree(temp_dir)
    except Exception:
        pass

@pytest.fixture(autouse=True)
def mock_win32com():
    """Automatically mock win32com to prevent interaction with live Word."""
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")), \
         patch("win32com.client.GetObject", side_effect=Exception("Word not running")):
        yield

@pytest.fixture
def temp_xlsx():
    """Create a temporary Excel workbook."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "test.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Revenue"
    ws["A2"] = 500.0
    ws["B1"] = "Expenses"
    ws["B2"] = 300.0
    ws["C1"] = "Net"
    wb.save(file_path)
    yield file_path
    shutil.rmtree(temp_dir)

# ===========================================================================
# PR-01: Word injection uses correct paragraph style
# ===========================================================================
def test_pr01_word_style_conformance(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    # Valid style matches Heading 1
    validator = WordOperationValidator()
    assert validator._fuzzy_style_match("Heading1", ctx.styles["paragraph"]) == "Heading 1"

    # Invalid style is rejected
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "UnmatchedStyleNameXYZ",
        "runs": [{"text": "Hello world"}]
    }
    res = validator.validate(op, ctx)
    assert res.valid is False

# ===========================================================================
# PR-02: GRP never injects without Tab approval (Esc test)
# ===========================================================================
def test_pr02_terminal_show_only_enforcement():
    # Terminal master must always enforce show_only injection
    master = TerminalMaster()
    resp = TerminalResponse(
        injection_method="show_only",
        command="git status",
        explanation="shows working tree status",
        danger_level="safe",
        confidence=1.0,
        reasoning="user requested git status"
    )
    ops = master.validate_operations(resp, {"shell_type": "powershell"})
    assert len(ops) == 1
    assert ops[0]["injection_method"] == "show_only"

# ===========================================================================
# PR-03: System prompt never leaks
# ===========================================================================
def test_pr03_system_prompt_leakage_detector():
    verifier = OutputVerifier()
    
    # A normal clean output should pass
    clean_report = verifier.run_all_checks('{"operations": []}', "word", "context")
    assert clean_report.all_passed is True

    # System instruction terms should trigger a leak report
    leaked_report = verifier.run_all_checks('{"operations": [], "reasoning": "using waza_agent system prompt details"}', "word", "context")
    assert leaked_report.all_passed is False
    assert any("leakage" in issue for issue in leaked_report.issues)

# ===========================================================================
# PR-04: Zero external connections in offline mode
# ===========================================================================
def test_pr04_offline_mode_isolation():
    # Attempting to open a socket connection should raise an error in this test
    # to guarantee our masters perform extraction/routing/validation 100% offline.
    original_connect = socket.socket.connect
    
    def blocked_connect(self, address):
        raise socket.error("Network connection blocked in offline test mode.")
        
    with patch("socket.socket.connect", side_effect=blocked_connect):
        # Instantiate objects and check offline behaviors
        router = DomainMasterRouter()
        # Ensure it starts without making any network connections
        assert router is not None

# ===========================================================================
# PR-05: Ctrl+Z undoes entire injection
# ===========================================================================
def test_pr05_ctrl_z_undo_simulation(temp_docx):
    # Undo simulation: sidecar saves a backup, then restores it on undo command
    original_size = os.path.getsize(temp_docx)
    
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    
    # Simulate the frontend / extension holding a backup for undo (Ctrl+Z)
    backup_path = temp_docx + ".undo_bak"
    shutil.copy2(temp_docx, backup_path)
    
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 2,
        "style": "Normal",
        "runs": [{"text": "Injected paragraph"}]
    }
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        res = writer.apply_operations(temp_docx, [op], ctx)
        assert len(res.get("errors", [])) == 0
        
        # Verify document was modified (size changed)
        assert os.path.getsize(temp_docx) != original_size
        
        # Undo simulation: restore backup
        shutil.copy2(backup_path, temp_docx)
        os.remove(backup_path)
        
        # Verify file is restored and has the original content/size
        assert os.path.getsize(temp_docx) == original_size
        doc = Document(temp_docx)
        assert len(doc.paragraphs) == 4 # Heading 1 + 3 original paras (no injected one)
        assert doc.paragraphs[2].text == "Target paragraph."

# ===========================================================================
# PR-06: Excel formulas validated before injection
# ===========================================================================
def test_pr06_excel_formula_validation(temp_xlsx):
    extractor = ExcelContextExtractor()
    ctx = extractor.extract(temp_xlsx, "C2")
    
    validator = ExcelOperationValidator()
    
    # 1. Circular reference test: targeting C2, formula references C2
    circular_op = {
        "type": "write_cell",
        "cell": "C2",
        "formula": "=SUM(A2:C2)"
    }
    res = validator.validate(circular_op, ctx)
    assert res.valid is False
    assert "circular reference" in res.error.lower()
    
    # 2. Balanced parentheses auto-fix
    unbalanced_op = {
        "type": "write_cell",
        "cell": "C2",
        "formula": "=SUM(A2:B2"
    }
    res_unbalanced = validator.validate(unbalanced_op, ctx)
    assert res_unbalanced.valid is True
    assert res_unbalanced.op["formula"] == "=SUM(A2:B2)"

# ===========================================================================
# PR-07: Adjacent cells/paragraphs never modified
# ===========================================================================
def test_pr07_adjacent_elements_unchanged(temp_docx, temp_xlsx):
    # Test Word adjacent paragraphs preservation
    word_extractor = WordContextExtractor()
    word_ctx = word_extractor.extract(temp_docx, 2)
    
    word_writer = WordWriter()
    word_op = {
        "type": "replace_paragraph",
        "paragraph_index": 2, # target "Target paragraph."
        "style": "Normal",
        "runs": [{"text": "Replaced target paragraph."}]
    }
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        word_writer.apply_operations(temp_docx, [word_op], word_ctx)
        doc = Document(temp_docx)
        # Adjacent paragraphs must be untouched
        assert doc.paragraphs[1].text == "Adjacent paragraph 1."
        assert doc.paragraphs[3].text == "Adjacent paragraph 2."

    # Test Excel adjacent cells preservation
    excel_extractor = ExcelContextExtractor()
    excel_ctx = excel_extractor.extract(temp_xlsx, "C2")
    
    excel_writer = ExcelWriter()
    excel_op = {
        "type": "write_cell",
        "cell": "C2",
        "value": 200.0
    }
    excel_writer.apply_operations(temp_xlsx, [excel_op])
    
    wb = openpyxl.load_workbook(temp_xlsx)
    ws = wb.active
    # Target changed
    assert ws["C2"].value == 200.0
    # Adjacent cells untouched
    assert ws["A2"].value == 500.0
    assert ws["B2"].value == 300.0

# ===========================================================================
# PR-08: Sidecar crash leaves original file intact
# ===========================================================================
def test_pr08_atomic_save_crash_safety(temp_docx):
    extractor = WordContextExtractor()
    ctx = extractor.extract(temp_docx, 1)
    
    writer = WordWriter()
    op = {
        "type": "insert_paragraph",
        "after_paragraph_index": 1,
        "style": "Normal",
        "runs": [{"text": "Failed insert"}]
    }
    
    original_size = os.path.getsize(temp_docx)
    
    with patch("win32com.client.GetActiveObject", side_effect=Exception("Word not running")):
        # Mock document save to throw exception, simulating crash
        with patch("docx.document.Document.save", side_effect=IOError("Simulated disk crash")):
            with pytest.raises(IOError):
                writer.apply_operations(temp_docx, [op], ctx)
            
            # File must still exist and be intact with original size
            assert os.path.exists(temp_docx)
            assert os.path.getsize(temp_docx) == original_size

# ===========================================================================
# PR-09: First token under 2s (7B) / 600ms (3B)
# ===========================================================================
def test_pr09_first_token_latency():
    # Context assemblies from preloaded cache must be sub-100ms to guarantee fast first-token start
    from sidecar.kairo_eye.context_assembler import ContextAssembler
    assembler = ContextAssembler()
    preloaded = {"paragraphs": list(range(50))}
    
    start = time.time()
    ctx = assembler.assemble(
        preloaded_ctx=preloaded,
        cursor_pos=0,
        mem_ctx="No style preferences.",
        domain="word",
        file_path="C:/dummy.docx"
    )
    elapsed = time.time() - start
    assert elapsed < 0.1 # <100ms assembly time leaves 500ms+ margin for model generation start

@pytest.mark.anyio
async def test_pr13_domain_routing_accuracy(temp_docx):
    router = DomainMasterRouter()
    
    mock_response = MagicMock()
    mock_response.operations = []
    mock_response.confidence = 0.95
    mock_response.reasoning = "Test"
    mock_response.model_dump.return_value = {"operations": [], "confidence": 0.95, "reasoning": "Test"}

    with patch("sidecar.router.call_with_schema", return_value=mock_response):
        # Route docx extension/domain
        res_docx = await router.route_llm_request(
            domain="docx",
            file_path=temp_docx,
            user_instruction="test docx",
            mem_context="",
            cursor_info=0
        )
        assert res_docx["ok"] is True
        
        # Route xlsx extension/domain
        res_xlsx = await router.route_llm_request(
            domain="xlsx",
            file_path=temp_docx,
            user_instruction="test xlsx",
            mem_context="",
            cursor_info="A1"
        )
        assert res_xlsx["ok"] is True

    # Process classification
    from sidecar.kairo_eye.app_watcher import AppWatcher, Domain
    watcher = AppWatcher()
    assert watcher.get_domain_for_process("winword.exe") == Domain.WORD
    assert watcher.get_domain_for_process("excel.exe") == Domain.EXCEL
    assert watcher.get_domain_for_process("powerpnt.exe") == Domain.POWERPOINT
    assert watcher.get_domain_for_process("chrome.exe") == Domain.BROWSER

# ===========================================================================
# PR-14: MemMachine style recall across sessions
# ===========================================================================
def test_pr14_mem_machine_session_recall():
    with tempfile.NamedTemporaryFile(suffix=".db", delete=False) as f:
        temp_db_path = f.name
        
    try:
        client = MemMachineClient(db_path=temp_db_path)
        
        # Record preference in session 1
        client.record_interaction(
            domain="word",
            task_type="insert",
            user_prompt="draft a formal email",
            style_notes="User prefers signature with 'Best regards'"
        )
        
        # Close client, re-open client to simulate session 2
        client2 = MemMachineClient(db_path=temp_db_path)
        notes = client2.query(domain="word")
        
        assert "best regards" in notes.lower()
    finally:
        if os.path.exists(temp_db_path):
            os.remove(temp_db_path)
