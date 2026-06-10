"""
docling_parser.py
=================
Structured document parser integration using Docling as the primary engine,
with fallback to mammoth / docx2python / python-docx.

Public API
----------
DoclingParser          – class with .parse(file_path, format) method
parse_docx_structured  – convenience function for WordContext-compatible output

Cache Strategy
--------------
An in-process LRU cache keyed on (absolute_path, mtime) avoids re-parsing
unchanged files between calls within the same process lifetime.
"""

import os
import datetime
import logging
import threading
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger("kairo-sidecar.docling_parser")

# ---------------------------------------------------------------------------
# Optional dependency probes
# ---------------------------------------------------------------------------

try:
    from docling.document_converter import DocumentConverter as _DoclingConverter
    _DOCLING_AVAILABLE = True
    log.debug("Docling is available.")
except ImportError:
    _DoclingConverter = None  # type: ignore
    _DOCLING_AVAILABLE = False
    log.debug("Docling not installed – will use fallback parsers.")

try:
    import mammoth as _mammoth
    _MAMMOTH_AVAILABLE = True
except ImportError:
    _mammoth = None  # type: ignore
    _MAMMOTH_AVAILABLE = False

try:
    from docx2python import docx2python as _docx2python
    _DOCX2PYTHON_AVAILABLE = True
except ImportError:
    _docx2python = None  # type: ignore
    _DOCX2PYTHON_AVAILABLE = False

try:
    from docx import Document as _DocxDocument
    from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
    _PYTHON_DOCX_AVAILABLE = True
except ImportError:
    _DocxDocument = None  # type: ignore
    _WD_STYLE_TYPE = None  # type: ignore
    _PYTHON_DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# LRU-style in-process cache (keyed by (abs_path, mtime))
# ---------------------------------------------------------------------------

_CACHE_MAX_SIZE = 32
_cache: Dict[Tuple[str, float], Dict[str, Any]] = {}
_cache_order: List[Tuple[str, float]] = []
_cache_lock = threading.Lock()


def _cache_get(abs_path: str) -> Optional[Dict[str, Any]]:
    """Return cached parse result if the file hasn't changed since last parse."""
    try:
        mtime = os.path.getmtime(abs_path)
        key = (abs_path, mtime)
        with _cache_lock:
            return _cache.get(key)
    except OSError:
        return None


def _cache_set(abs_path: str, result: Dict[str, Any]) -> None:
    """Store parse result keyed by (abs_path, current mtime)."""
    try:
        mtime = os.path.getmtime(abs_path)
        key = (abs_path, mtime)
        with _cache_lock:
            if key not in _cache:
                # Evict oldest entry if at capacity
                while len(_cache_order) >= _CACHE_MAX_SIZE:
                    oldest = _cache_order.pop(0)
                    _cache.pop(oldest, None)
                _cache_order.append(key)
            _cache[key] = result
    except OSError:
        pass


# ---------------------------------------------------------------------------
# Heading style helpers
# ---------------------------------------------------------------------------

_HEADING_STYLE_MAP = {
    "title": ("Heading 1", 1),
    "chapter": ("Heading 1", 1),
}

def _heading_level_from_style(style_name: str) -> int:
    """
    Extract numeric heading level from style name strings like
    'Heading 1', 'heading1', 'h2', 'Title', etc.
    Returns 0 for non-heading styles.
    """
    lower = style_name.lower().strip()
    # Explicit level patterns: "Heading 3", "heading3", "h3"
    for prefix in ("heading ", "heading", "h"):
        if lower.startswith(prefix):
            remainder = lower[len(prefix):].strip()
            if remainder.isdigit():
                return min(int(remainder), 6)
    # Fuzzy: title → level 1
    if lower in ("title", "chapter"):
        return 1
    return 0


def _map_heading_to_word_style(level: int) -> str:
    """Return Word-compatible style string for a given heading level."""
    if level == 0:
        return "Normal"
    return f"Heading {level}"


# ---------------------------------------------------------------------------
# DoclingParser class
# ---------------------------------------------------------------------------

class DoclingParser:
    """
    Structured document parser backed by Docling with graceful fallbacks.

    parse(file_path, format) → {
        "paragraphs": [
            {
                "index": int,
                "text": str,
                "style": str,        # e.g. "Heading 1", "Normal", "List Bullet"
                "level": int,        # heading level (0 = body text)
                "page": int | None,  # page number if available
                "runs": [
                    {"text": str, "bold": bool, "italic": bool}
                ]
            },
            ...
        ],
        "tables": [
            {
                "after_paragraph_index": int,
                "rows": [[str, ...], ...]
            },
            ...
        ],
        "metadata": {
            "tier": str,                # "docling" | "mammoth" | "docx2python" | "python-docx"
            "total_paragraphs": int,
            "table_count": int,
            "file_path": str,
            "parse_timestamp": str
        }
    }
    """

    def parse(self, file_path: str, format: str = "docx") -> Dict[str, Any]:  # noqa: A002
        """
        Parse *file_path* and return the structured document dict.

        Parameters
        ----------
        file_path : str
            Absolute or relative path to the document.
        format : str
            Hint for the document format ('docx', 'pdf', 'pptx', …).
            Currently only 'docx' has full fallback support.
        """
        abs_path = str(Path(file_path).resolve())

        # 1. Check LRU cache first
        cached = _cache_get(abs_path)
        if cached is not None:
            log.debug("DoclingParser: cache hit for %s", abs_path)
            return cached

        result = None
        tier = "unknown"

        # 2. Tier A: Docling
        if _DOCLING_AVAILABLE:
            try:
                result = self._parse_with_docling(abs_path)
                tier = "docling"
                log.info("DoclingParser: Docling succeeded for %s", abs_path)
            except Exception as exc:
                log.debug("DoclingParser: Docling failed (%s), trying fallbacks.", exc)

        # 3. Tier B: Mammoth (docx only)
        if result is None and format.lower() in ("docx", "doc") and _MAMMOTH_AVAILABLE:
            try:
                result = self._parse_with_mammoth(abs_path)
                tier = "mammoth"
                log.info("DoclingParser: Mammoth succeeded for %s", abs_path)
            except Exception as exc:
                log.debug("DoclingParser: Mammoth failed (%s).", exc)

        # 4. Tier C: docx2python (docx only)
        if result is None and format.lower() in ("docx", "doc") and _DOCX2PYTHON_AVAILABLE:
            try:
                result = self._parse_with_docx2python(abs_path)
                tier = "docx2python"
                log.info("DoclingParser: docx2python succeeded for %s", abs_path)
            except Exception as exc:
                log.debug("DoclingParser: docx2python failed (%s).", exc)

        # 5. Tier D: python-docx (docx only, always available in this project)
        if result is None and format.lower() in ("docx", "doc") and _PYTHON_DOCX_AVAILABLE:
            try:
                result = self._parse_with_python_docx(abs_path)
                tier = "python-docx"
                log.info("DoclingParser: python-docx succeeded for %s", abs_path)
            except Exception as exc:
                log.warning("DoclingParser: python-docx also failed (%s).", exc)

        if result is None:
            raise RuntimeError(
                f"DoclingParser: all parsing tiers failed for {abs_path!r}. "
                "Ensure at least python-docx is installed."
            )

        # Stamp metadata tier
        result.setdefault("metadata", {})["tier"] = tier
        result["metadata"]["file_path"] = abs_path
        result["metadata"]["parse_timestamp"] = datetime.datetime.now().isoformat()

        # Store in cache
        _cache_set(abs_path, result)
        return result

    # ------------------------------------------------------------------
    # Tier A: Docling
    # ------------------------------------------------------------------

    def _parse_with_docling(self, abs_path: str) -> Dict[str, Any]:
        """Use Docling DocumentConverter to parse and map to Kairo schema."""
        converter = _DoclingConverter()
        doc_result = converter.convert(abs_path)
        dl_doc = doc_result.document  # DoclingDocument

        paragraphs: List[Dict[str, Any]] = []
        tables: List[Dict[str, Any]] = []
        p_idx = 0

        # Docling exposes items via .body.children or iterate_items()
        # We use the export API: dl_doc.export_to_dict() as a reliable
        # cross-version approach, then fall back to iterating texts/tables.
        try:
            raw = dl_doc.export_to_dict()
        except AttributeError:
            raw = {}

        body_items = raw.get("body", {}).get("children", [])

        def _process_items(items: list) -> None:
            nonlocal p_idx
            for item in items:
                item_type = item.get("$type", item.get("type", ""))
                label = item.get("label", "")

                # ---- Paragraph / Heading / List ----
                if item_type in (
                    "TextItem", "SectionHeaderItem", "ListItem",
                    "text", "section_header", "list_item",
                ):
                    text = (item.get("text") or "").strip()
                    level = 0
                    style = "Normal"
                    list_level = item.get("enumerated_level", item.get("level", None))

                    # Determine heading level
                    if item_type in ("SectionHeaderItem", "section_header"):
                        level = item.get("heading_level", 1)
                        style = _map_heading_to_word_style(level)
                    elif label.lower().startswith("section_header"):
                        level = item.get("heading_level", 1)
                        style = _map_heading_to_word_style(level)
                    elif label.lower() in ("list_item", "list-item"):
                        style = "List Bullet"
                    elif label.lower() in ("list_number", "list-number"):
                        style = "List Number"
                    else:
                        # Check prov for heading style names
                        prov_style = item.get("style", "")
                        if prov_style:
                            level = _heading_level_from_style(prov_style)
                            style = _map_heading_to_word_style(level) if level else prov_style

                    # Build runs from inline formatting if available
                    runs = _extract_docling_runs(item, text)

                    paragraphs.append({
                        "index": p_idx,
                        "text": text,
                        "style": style,
                        "level": level,
                        "page": _extract_page(item),
                        "runs": runs,
                    })
                    p_idx += 1

                # ---- Table ----
                elif item_type in ("TableItem", "table"):
                    rows = _extract_docling_table_rows(item)
                    tables.append({
                        "after_paragraph_index": max(p_idx - 1, 0),
                        "rows": rows,
                    })

                # ---- Recurse into groups ----
                children = item.get("children", [])
                if children:
                    _process_items(children)

        _process_items(body_items)

        # Fallback: if body_items empty, try iterating via dl_doc API directly
        if not paragraphs:
            paragraphs, tables, p_idx = _iterate_docling_doc(dl_doc)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "total_paragraphs": len(paragraphs),
                "table_count": len(tables),
            },
        }

    # ------------------------------------------------------------------
    # Tier B: Mammoth
    # ------------------------------------------------------------------

    def _parse_with_mammoth(self, abs_path: str) -> Dict[str, Any]:
        from bs4 import BeautifulSoup  # type: ignore

        with open(abs_path, "rb") as fh:
            result = _mammoth.convert_to_html(fh)
        html = result.value
        soup = BeautifulSoup(html, "html.parser")

        paragraphs: List[Dict[str, Any]] = []
        tables: List[Dict[str, Any]] = []
        p_idx = 0

        for child in soup.children:
            if not child or isinstance(child, str):
                continue
            tag = getattr(child, "name", None)
            if not tag:
                continue

            if tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6",
                       "ul", "ol", "li"):
                # Handle list items
                if tag in ("ul", "ol"):
                    for li in child.find_all("li"):
                        li_text = li.get_text().strip()
                        list_style = "List Bullet" if tag == "ul" else "List Number"
                        # Estimate nesting level by count of parent ul/ol
                        nest = len(li.find_parents(["ul", "ol"]))
                        paragraphs.append({
                            "index": p_idx,
                            "text": li_text,
                            "style": list_style,
                            "level": nest,
                            "page": None,
                            "runs": [{"text": li_text, "bold": False, "italic": False}],
                        })
                        p_idx += 1
                    continue

                text = child.get_text().strip()
                level = 0
                style = "Normal"
                if tag.startswith("h") and len(tag) == 2:
                    try:
                        level = int(tag[1])
                    except ValueError:
                        level = 1
                    style = _map_heading_to_word_style(level)

                runs = []
                for item in child.children:
                    if isinstance(item, str):
                        if item:
                            runs.append({"text": item, "bold": False, "italic": False})
                    else:
                        item_tag = getattr(item, "name", "")
                        bold = item_tag in ("strong", "b") or bool(item.find("strong"))
                        italic = item_tag in ("em", "i") or bool(item.find("em"))
                        runs.append({
                            "text": item.get_text(),
                            "bold": bold,
                            "italic": italic,
                        })
                if not runs:
                    runs = [{"text": text, "bold": False, "italic": False}]

                paragraphs.append({
                    "index": p_idx,
                    "text": text,
                    "style": style,
                    "level": level,
                    "page": None,
                    "runs": runs,
                })
                p_idx += 1

            elif tag == "table":
                rows = []
                for tr in child.find_all("tr"):
                    rows.append([td.get_text().strip() for td in tr.find_all(["td", "th"])])
                tables.append({
                    "after_paragraph_index": max(p_idx - 1, 0),
                    "rows": rows,
                })

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "total_paragraphs": len(paragraphs),
                "table_count": len(tables),
            },
        }

    # ------------------------------------------------------------------
    # Tier C: docx2python
    # ------------------------------------------------------------------

    def _parse_with_docx2python(self, abs_path: str) -> Dict[str, Any]:
        content = _docx2python(abs_path)
        paragraphs: List[Dict[str, Any]] = []
        tables: List[Dict[str, Any]] = []
        p_idx = 0

        def _flatten_body(body):
            nonlocal p_idx
            for section in body:
                for column in section:
                    for para_parts in column:
                        if isinstance(para_parts, list):
                            # Could be table row or paragraph parts
                            flat = " ".join(str(p) for p in para_parts).strip()
                            if flat:
                                level = _heading_level_from_style("")
                                paragraphs.append({
                                    "index": p_idx,
                                    "text": flat,
                                    "style": "Normal",
                                    "level": 0,
                                    "page": None,
                                    "runs": [{"text": flat, "bold": False, "italic": False}],
                                })
                                p_idx += 1
                        elif isinstance(para_parts, str) and para_parts.strip():
                            paragraphs.append({
                                "index": p_idx,
                                "text": para_parts.strip(),
                                "style": "Normal",
                                "level": 0,
                                "page": None,
                                "runs": [{"text": para_parts.strip(), "bold": False, "italic": False}],
                            })
                            p_idx += 1

        _flatten_body(content.body)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "total_paragraphs": len(paragraphs),
                "table_count": len(tables),
            },
        }

    # ------------------------------------------------------------------
    # Tier D: python-docx
    # ------------------------------------------------------------------

    def _parse_with_python_docx(self, abs_path: str) -> Dict[str, Any]:
        from docx.oxml.text.paragraph import CT_P  # type: ignore
        from docx.oxml.table import CT_Tbl  # type: ignore

        doc = _DocxDocument(abs_path)
        paragraphs: List[Dict[str, Any]] = []
        tables: List[Dict[str, Any]] = []

        for i, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else "Normal"
            level = _heading_level_from_style(style_name)
            if not level and style_name.lower() in ("title",):
                level = 1
                style_name = "Heading 1"

            # Detect list style
            list_level = None
            try:
                if (p._p.pPr is not None
                        and p._p.pPr.numPr is not None
                        and p._p.pPr.numPr.ilvl is not None):
                    list_level = p._p.pPr.numPr.ilvl.val
            except Exception:
                pass

            runs = []
            for r in p.runs:
                runs.append({
                    "text": r.text,
                    "bold": bool(r.bold),
                    "italic": bool(r.italic),
                })

            paragraphs.append({
                "index": i,
                "text": p.text.strip(),
                "style": style_name,
                "level": level if level else (list_level if list_level is not None else 0),
                "page": None,
                "runs": runs,
            })

        # Walk body XML to get accurate table positions
        paragraph_idx = -1
        table_idx = 0
        for child in doc.element.body:
            if isinstance(child, CT_P):
                paragraph_idx += 1
            elif isinstance(child, CT_Tbl):
                if table_idx < len(doc.tables):
                    table_el = doc.tables[table_idx]
                    table_idx += 1
                    rows = [[cell.text.strip() for cell in row.cells]
                            for row in table_el.rows]
                    tables.append({
                        "after_paragraph_index": max(paragraph_idx, 0),
                        "rows": rows,
                    })

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "total_paragraphs": len(paragraphs),
                "table_count": len(tables),
            },
        }


# ---------------------------------------------------------------------------
# Docling helper utilities
# ---------------------------------------------------------------------------

def _extract_page(item: dict) -> Optional[int]:
    """Try to pull page number from Docling item provenance."""
    for prov in item.get("prov", []):
        page = prov.get("page_no", prov.get("page", None))
        if page is not None:
            return int(page)
    return None


def _extract_docling_runs(item: dict, fallback_text: str) -> List[Dict[str, Any]]:
    """
    Build a runs list from Docling inline formatting annotations.
    Falls back to a single run with the full text if no annotation data.
    """
    # Docling ≥ 2.x stores inline annotations under item["annotations"] or
    # item["inline_annotations"]; earlier versions may not have them.
    annotations = item.get("annotations", item.get("inline_annotations", []))
    if not annotations:
        return [{"text": fallback_text, "bold": False, "italic": False}]

    runs = []
    for ann in annotations:
        seg = ann.get("text", ann.get("segment", fallback_text))
        bold = ann.get("bold", False)
        italic = ann.get("italic", False)
        runs.append({"text": seg, "bold": bool(bold), "italic": bool(italic)})
    return runs if runs else [{"text": fallback_text, "bold": False, "italic": False}]


def _extract_docling_table_rows(item: dict) -> List[List[str]]:
    """Convert a Docling TableItem dict into a list of row lists."""
    # Docling TableItem may expose data via item["data"]["grid"] or
    # item["table_cells"] depending on version.
    grid = item.get("data", {}).get("grid", [])
    if grid:
        rows = []
        for row in grid:
            rows.append([cell.get("text", "") for cell in row])
        return rows

    # Fallback: table_cells flat list with row_span / col_span
    cells = item.get("table_cells", [])
    if cells:
        max_row = max(c.get("row", 0) for c in cells) + 1
        max_col = max(c.get("col", 0) for c in cells) + 1
        grid_out = [[""] * max_col for _ in range(max_row)]
        for c in cells:
            grid_out[c.get("row", 0)][c.get("col", 0)] = c.get("text", "")
        return grid_out

    return []


def _iterate_docling_doc(dl_doc) -> Tuple[List, List, int]:
    """
    Fallback iteration when export_to_dict() body is empty.
    Uses the Docling document's iterate_items() or texts / tables iterators.
    """
    paragraphs: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []
    p_idx = 0

    # Try iterate_items (Docling ≥ 2.x)
    try:
        for item, _ in dl_doc.iterate_items():
            item_cls = type(item).__name__
            text = getattr(item, "text", "") or ""
            text = text.strip()
            level = 0
            style = "Normal"

            if "SectionHeader" in item_cls or "Heading" in item_cls:
                level = getattr(item, "heading_level", 1) or 1
                style = _map_heading_to_word_style(level)
            elif "Table" in item_cls:
                # Extract table using export methods
                rows: List[List[str]] = []
                try:
                    export = item.export_to_dataframe()
                    rows = export.values.tolist()
                    rows = [[str(v) for v in row] for row in rows]
                except Exception:
                    pass
                tables.append({
                    "after_paragraph_index": max(p_idx - 1, 0),
                    "rows": rows,
                })
                continue
            elif "List" in item_cls:
                style = "List Bullet"

            if text:
                paragraphs.append({
                    "index": p_idx,
                    "text": text,
                    "style": style,
                    "level": level,
                    "page": None,
                    "runs": [{"text": text, "bold": False, "italic": False}],
                })
                p_idx += 1
    except AttributeError:
        pass

    return paragraphs, tables, p_idx


# ---------------------------------------------------------------------------
# Singleton parser instance
# ---------------------------------------------------------------------------

_parser_instance = DoclingParser()


# ---------------------------------------------------------------------------
# Public convenience function: parse_docx_structured
# ---------------------------------------------------------------------------

def parse_docx_structured(file_path: str) -> Dict[str, Any]:
    """
    Parse a .docx file using Docling (with fallbacks) and return a dict
    that is compatible with WordContext.paragraphs + WordContext.tables.

    Returns
    -------
    dict with keys:
        "paragraphs"  – list of paragraph dicts (index, text, style, level, page, runs)
        "tables"      – list of table dicts (after_paragraph_index, rows)
        "metadata"    – parse metadata (tier, total_paragraphs, table_count, ...)
    """
    return _parser_instance.parse(file_path, format="docx")
