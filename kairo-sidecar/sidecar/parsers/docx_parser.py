import datetime
import os
import traceback
import logging
from pathlib import Path
from typing import Dict, Any

from sidecar.parsers.cache import DocumentCache

log = logging.getLogger("kairo-sidecar.docx_parser")

def parse_docx(file_path: str) -> dict:
    """
    Parses a .docx file and returns structured JSON conforming to:
    {
      "paragraphs": [{"index": int, "text": str, "style": str, "level": int, "runs": [{"text": str, "bold": bool, "italic": bool}]}],
      "tables": [{"after_paragraph_index": int, "rows": [[str]]}],
      "metadata": {"total_paragraphs": int, "heading_count": int, "table_count": int, "file_path": str, "parse_timestamp": str}
    }
    
    Uses 4-tier parsing chain:
    Tier 0: win32com (reads live Word document — bypasses file lock)
    Tier 1: MinerU (optional, highly accurate layout parser)
    Tier 2: Mammoth (HTML extraction + BeautifulSoup)
    Tier 3: python-docx (direct XML extraction)
    """
    file_path_abs = str(Path(file_path).resolve())

    # Tier 0: Try COM first (Word is almost always open during testing)
    try:
        result = _parse_docx_com(file_path_abs)
        if result:
            log.info(f"Tier 0 parser (COM/Word) succeeded for {file_path_abs}")
            # Don't cache COM result — it changes per keystroke
            return result
    except Exception as e:
        log.debug(f"Tier 0 COM parser failed (normal if Word not open): {e}")

    # 1. Check Cache (only for file-based parsers)
    cached_data = DocumentCache.get(file_path_abs)
    if cached_data:
        log.info(f"Using cached parse output for: {file_path_abs}")
        return cached_data

    # 2. Try parsing
    result = None
    errors = []

    # Tier 1: MinerU
    try:
        log.info(f"Attempting Tier 1 parser (MinerU) for {file_path_abs}")
        result = _parse_docx_mineru(file_path_abs)
        log.info("Tier 1 parser (MinerU) succeeded")
    except Exception as e:
        errors.append(f"MinerU failed: {e}")
        log.debug(traceback.format_exc())

    # Tier 2: Mammoth
    if not result:
        try:
            log.info(f"Attempting Tier 2 parser (Mammoth) for {file_path_abs}")
            result = _parse_docx_mammoth(file_path_abs)
            log.info("Tier 2 parser (Mammoth) succeeded")
        except Exception as e:
            errors.append(f"Mammoth failed: {e}")
            log.debug(traceback.format_exc())

    # Tier 3: python-docx
    if not result:
        try:
            log.info(f"Attempting Tier 3 parser (python-docx) for {file_path_abs}")
            result = _parse_docx_python_docx(file_path_abs)
            log.info("Tier 3 parser (python-docx) succeeded")
        except Exception as e:
            errors.append(f"python-docx failed: {e}")
            log.error(f"All parsers failed for {file_path_abs}. Errors: {errors}")
            raise RuntimeError(f"All document parsers failed: {errors}")

    # Save to Cache on success
    DocumentCache.set(file_path_abs, result)
    return result


def _parse_docx_com(file_path: str) -> dict:
    """
    Reads the live open Word document via COM automation.
    This bypasses the file lock that Word holds when a doc is open.
    Returns None if Word is not running or the file is not open.
    """
    import win32com.client
    import pythoncom
    import subprocess
    pythoncom.CoInitialize()
    
    word = None
    target_doc = None
    path = Path(file_path).resolve()
    
    word_running = False
    try:
        out = subprocess.run(["tasklist", "/FI", "IMAGENAME eq winword.exe"], capture_output=True, text=True)
        word_running = "winword.exe" in out.stdout.lower()
    except Exception:
        pass
        
    if word_running:
        try:
            target_doc = win32com.client.GetObject(str(path))
            word = target_doc.Application
        except Exception:
            try:
                word = win32com.client.GetActiveObject("Word.Application")
                for doc in word.Documents:
                    try:
                        if Path(doc.FullName).resolve() == path:
                            target_doc = doc
                            break
                    except Exception:
                        pass
            except Exception:
                pass
    
    if not target_doc:
        return None  # File not open in Word
    
    paragraphs = []
    tables = []
    p_idx = 0
    
    for i in range(1, target_doc.Paragraphs.Count + 1):
        p = target_doc.Paragraphs(i)
        text = p.Range.Text.rstrip("\r\n\x0d\x07")  # Strip Word paragraph marks
        style_name = "Normal"
        level = 0
        try:
            style_name = p.Style.NameLocal
        except Exception:
            pass
        
        if "Heading" in style_name:
            try:
                level = int(style_name.split()[-1])
            except Exception:
                level = 1
        
        runs = [{"text": text, "bold": bool(p.Range.Bold), "italic": bool(p.Range.Italic)}]
        paragraphs.append({
            "index": p_idx,
            "text": text,
            "style": style_name,
            "level": level,
            "runs": runs
        })
        p_idx += 1
    
    # Extract tables
    for t_idx in range(1, target_doc.Tables.Count + 1):
        table = target_doc.Tables(t_idx)
        rows = []
        for r in range(1, table.Rows.Count + 1):
            row = []
            for c in range(1, table.Columns.Count + 1):
                try:
                    row.append(table.Cell(r, c).Range.Text.rstrip("\r\n\x0d\x07"))
                except Exception:
                    row.append("")
            rows.append(row)
        tables.append({"after_paragraph_index": p_idx - 1, "rows": rows})
    
    headings_count = sum(1 for p in paragraphs if p["level"] > 0)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {
            "total_paragraphs": len(paragraphs),
            "heading_count": headings_count,
            "table_count": len(tables),
            "file_path": file_path,
            "parse_timestamp": datetime.datetime.now().isoformat(),
            "source": "com"
        }
    }



def _parse_docx_mineru(file_path: str) -> dict:
    """
    MinerU 3.1 DOCX Layout Parser API interface.
    Since MinerU can be heavy to load, we wrap this in try-except in parse_docx.
    """
    # Try importing MinerU libraries
    from magic_pdf.data.data_reader_writer import FileBasedReaderWriter
    from magic_pdf.pipe.UNIPipe import UNIPipe
    
    # Create output directory for MinerU internal processing
    parent_dir = str(Path(file_path).parent)
    reader = FileBasedReaderWriter(parent_dir)
    
    # Initialize UNIPipe on docx
    with open(file_path, "rb") as f:
        docx_bytes = f.read()
        
    pipe = UNIPipe(docx_bytes, "docx", {})
    pipe.pipe_classify()
    pipe.pipe_analyze()
    model_json = pipe.pipe_to_json()
    
    # Process model_json and map to Kairo schema
    # Let's map headings, texts, tables, and runs
    paragraphs = []
    tables = []
    
    p_idx = 0
    # Process blocks from MinerU json output
    for block in model_json:
        block_type = block.get("type", "")
        if block_type in ("text", "title", "heading"):
            text = block.get("text", "").strip()
            style = "Normal"
            level = 0
            if block_type == "title":
                style = "Heading1"
                level = 1
            elif block_type == "heading":
                level = block.get("level", 1)
                style = f"Heading{level}"
                
            runs = [{"text": text, "bold": False, "italic": False}]
            paragraphs.append({
                "index": p_idx,
                "text": text,
                "style": style,
                "level": level,
                "runs": runs
            })
            p_idx += 1
        elif block_type == "table":
            # Extract rows from block
            rows = block.get("table_cells", [])
            # MinerU tables have structured format, let's build rows grid
            tables.append({
                "after_paragraph_index": p_idx - 1,
                "rows": rows
            })
            
    headings_count = sum(1 for p in paragraphs if p["level"] > 0)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {
            "total_paragraphs": len(paragraphs),
            "heading_count": headings_count,
            "table_count": len(tables),
            "file_path": file_path,
            "parse_timestamp": datetime.datetime.now().isoformat()
        }
    }


def _parse_docx_mammoth(file_path: str) -> dict:
    """
    Converts .docx to HTML via Mammoth, then extracts structured document layout using BeautifulSoup.
    """
    import mammoth
    from bs4 import BeautifulSoup
    
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
        
    soup = BeautifulSoup(html, "html.parser")
    paragraphs = []
    tables = []
    
    p_idx = 0
    # Mammoth outputs a flat body with elements. We traverse linearly.
    for child in soup.children:
        if not child or isinstance(child, str):
            continue
        if child.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6"):
            text = child.get_text().strip()
            level = 0
            style = "Normal"
            if child.name.startswith("h"):
                try:
                    level = int(child.name[1])
                except ValueError:
                    level = 1
                style = f"Heading{level}"
            
            # Map child text runs to runs
            runs = []
            for item in child.children:
                if isinstance(item, str):
                    if item:
                        runs.append({"text": item, "bold": False, "italic": False})
                else:
                    bold = item.name == "strong" or bool(item.find("strong"))
                    italic = item.name == "em" or bool(item.find("em"))
                    runs.append({
                        "text": item.get_text(),
                        "bold": bold,
                        "italic": italic
                    })
            
            if not runs:
                runs = [{"text": text, "bold": False, "italic": False}]
                
            paragraphs.append({
                "index": p_idx,
                "text": text,
                "style": style,
                "level": level,
                "runs": runs
            })
            p_idx += 1
            
        elif child.name == "table":
            rows = []
            for tr in child.find_all("tr"):
                rows.append([td.get_text().strip() for td in tr.find_all(["td", "th"])])
            tables.append({
                "after_paragraph_index": p_idx - 1,
                "rows": rows
            })
            
    headings_count = sum(1 for p in paragraphs if p["level"] > 0)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {
            "total_paragraphs": len(paragraphs),
            "heading_count": headings_count,
            "table_count": len(tables),
            "file_path": file_path,
            "parse_timestamp": datetime.datetime.now().isoformat()
        }
    }


def _parse_docx_python_docx(file_path: str) -> dict:
    """
    Extracts text, styles, and runs directly using python-docx.
    Linear body traversal identifies exact relative positioning for tables.
    """
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    doc = Document(file_path)
    paragraphs = []
    tables = []
    
    # Enumerate paragraphs
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else "Normal"
        level = 0
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip())
            except ValueError:
                level = 1
        runs = []
        for r in p.runs:
            runs.append({
                "text": r.text,
                "bold": bool(r.bold),
                "italic": bool(r.italic)
            })
        paragraphs.append({
            "index": i,
            "text": text,
            "style": style,
            "level": level,
            "runs": runs
        })
        
    # Walk doc XML body to place tables at correct paragraph indices
    body_elements = doc.element.body
    paragraph_idx = -1
    table_idx = 0
    
    for child in body_elements:
        if isinstance(child, CT_P):
            paragraph_idx += 1
        elif isinstance(child, CT_Tbl):
            if table_idx < len(doc.tables):
                table_el = doc.tables[table_idx]
                table_idx += 1
                rows = []
                for r in table_el.rows:
                    rows.append([cell.text.strip() for cell in r.cells])
                tables.append({
                    "after_paragraph_index": paragraph_idx,
                    "rows": rows
                })
                
    headings_count = sum(1 for p in paragraphs if p["level"] > 0)
    
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {
            "total_paragraphs": len(paragraphs),
            "heading_count": headings_count,
            "table_count": len(tables),
            "file_path": file_path,
            "parse_timestamp": datetime.datetime.now().isoformat()
        }
    }
