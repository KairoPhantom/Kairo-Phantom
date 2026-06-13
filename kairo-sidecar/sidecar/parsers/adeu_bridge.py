"""
Adeu Bridge — Native Track Changes for Kairo Phantom Word Specialist.
=====================================================================
Provides zero-clipboard, format-preserving edit injection into DOCX files
via the Adeu Virtual DOM (https://github.com/dealfluence/adeu).

Injection ladder (tried in order):
  1. process_active_word_batch  — live Word COM injection (Track Changes, no clipboard)
  2. process_document_batch     — file-based Track Changes via Adeu Python SDK
  3. Fallback to caller         — safe-docx or clipboard (managed by injector.rs)

All functions return dicts matching Kairo's sidecar JSON envelope:
  {"ok": bool, "data": {...}, "error": str | None}
"""

from __future__ import annotations

import json
import logging
import os
import shutil
import subprocess
import tempfile
import traceback
from io import BytesIO
from pathlib import Path
from typing import Any

log = logging.getLogger("kairo-sidecar.adeu_bridge")


# ---------------------------------------------------------------------------
# Capability detection (safe — never raises)
# ---------------------------------------------------------------------------

def _get_adeu_cmd() -> str:
    """Resolve the absolute path to adeu.exe, falling back to python scripts folder on Windows."""
    which_res = shutil.which("adeu")
    if which_res is not None:
        return which_res
    import sys
    if sys.platform == "win32":
        import os
        user_profile = os.environ.get("USERPROFILE", "")
        # Standard Python User Scripts path
        fallback_path = os.path.join(user_profile, "AppData", "Roaming", "Python", f"Python{sys.version_info.major}{sys.version_info.minor}", "Scripts", "adeu.exe")
        if os.path.exists(fallback_path):
            return fallback_path
        # Alternative path
        app_data = os.environ.get("APPDATA", "")
        fallback_path_2 = os.path.join(app_data, "Python", "Scripts", "adeu.exe")
        if os.path.exists(fallback_path_2):
            return fallback_path_2
    return "adeu"


def _adeu_installed() -> bool:
    """True if the `adeu` CLI is available in PATH or resolved locally."""
    cmd = _get_adeu_cmd()
    if cmd == "adeu":
        return shutil.which("adeu") is not None
    return Path(cmd).exists()


def _adeu_sdk_available() -> bool:
    """True if the adeu Python package is importable."""
    try:
        import adeu  # noqa: F401
        return True
    except ImportError:
        return False


def _word_is_open_with_file(file_path: str) -> bool:
    """
    Windows-only: check if WINWORD.EXE has the file open.
    Returns False on non-Windows or any error — fail-safe.
    """
    try:
        import sys
        if sys.platform != "win32":
            return False
        
        import subprocess
        word_running = False
        try:
            out = subprocess.run(["tasklist", "/FI", "IMAGENAME eq winword.exe"], capture_output=True, text=True)
            word_running = "winword.exe" in out.stdout.lower()
        except Exception:
            pass

        if not word_running:
            return False

        import win32com.client
        target = Path(file_path).resolve()
        try:
            doc = win32com.client.GetObject(str(target))
            return True
        except Exception:
            try:
                word = win32com.client.GetActiveObject("Word.Application")
                for doc in word.Documents:
                    if Path(doc.FullName).resolve() == target:
                        return True
            except Exception:
                pass
        return False
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def adeu_read_document(file_path: str) -> dict:
    """
    Extract document text + paragraph structure via `adeu extract`.

    Returns:
        {
            "ok": True,
            "data": {
                "full_text": str,          # CriticMarkup markdown
                "paragraphs": [...],       # indexed paragraph list with style
                "headings": [...],         # indexed headings list
                "paragraph_count": int,
                "format": "criticmarkup",
            }
        }
    """
    if not _adeu_installed():
        return _unavailable("adeu CLI not installed — run: pip install adeu")

    file_path = str(Path(file_path).resolve())
    if not Path(file_path).exists():
        return _error(f"File not found: {file_path}")

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".md")
    os.close(tmp_fd)
    try:
        result = subprocess.run(
            [_get_adeu_cmd(), "extract", file_path, "-o", tmp_path],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )
        with open(tmp_path, "r", encoding="utf-8") as f:
            markdown = f.read()

        paragraphs = _parse_markdown_to_paragraphs(markdown)
        for p in paragraphs:
            if p["is_heading"]:
                p["style"] = f"Heading {p['level']}"
            else:
                p["style"] = "Normal"

        headings = [
            {
                "index": p["index"],
                "level": p["level"],
                "text": p["text"],
            }
            for p in paragraphs if p["is_heading"]
        ]

        log.info(
            "adeu_read_document OK: %d chars, %d paragraphs, %d headings from %s",
            len(markdown), len(paragraphs), len(headings), file_path,
        )
        return {
            "ok": True,
            "data": {
                "full_text": markdown,
                "paragraphs": paragraphs,
                "headings": headings,
                "paragraph_count": len(paragraphs),
                "format": "criticmarkup",
                "file_path": file_path,
            },
        }
    except subprocess.TimeoutExpired:
        return _error("adeu extract timed out after 30 s")
    except subprocess.CalledProcessError as e:
        return _error(f"adeu extract failed (exit {e.returncode}): {e.stderr[:500]}")
    except Exception:
        return _error(traceback.format_exc())
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def adeu_apply_edits(
    file_path: str,
    edits: list[dict],
    output_path: str | None = None,
    author: str = "Kairo AI",
) -> dict:
    """
    Apply a list of text edits as native DOCX Track Changes (w:ins/w:del).

    `edits` is a list of dicts with keys:
        target_text  (str, required) — exact text to find and replace
        new_text     (str, required) — replacement text
        comment      (str, optional) — comment/rationale for the change

    Returns the path of the saved redlined file.
    """
    file_path = str(Path(file_path).resolve())
    if not Path(file_path).exists():
        return _error(f"File not found: {file_path}")
    if not edits:
        return _error("edits list is empty — nothing to apply")

    res = None
    # --- Strategy 1: Live Word COM injection (best quality)
    if _word_is_open_with_file(file_path):
        live_result = _adeu_live_apply(file_path, edits, author)
        if live_result["ok"]:
            res = live_result
        else:
            log.warning("Live COM injection failed, falling back to SDK: %s", live_result.get("error"))

    if res is None:
        # --- Strategy 2: Python SDK file-based injection
        if _adeu_sdk_available():
            res = _adeu_sdk_apply(file_path, edits, output_path, author)
        # --- Strategy 3: CLI-based fallback (adeu apply)
        elif _adeu_installed():
            res = _adeu_cli_apply(file_path, edits, output_path, author)
        # --- Strategy 4: Fallback Track Changes using python-docx XML w:ins/w:del revision marks
        else:
            log.info("Adeu not available, using python-docx XML tracked changes fallback")
            res = _python_docx_tracked_fallback(file_path, edits, output_path, author)

    if res.get("ok") and "data" in res and res["data"] is not None:
        # Guarantee keys in both top-level and data
        data = res["data"]
        applied = res.get("applied_count", data.get("applied_count", len(edits)))
        skipped = res.get("skipped_count", data.get("skipped_count", 0))
        backend = res.get("backend", data.get("backend", "unknown"))
        
        data["applied_count"] = applied
        data["skipped_count"] = skipped
        data["backend"] = backend
        
        res["applied_count"] = applied
        res["skipped_count"] = skipped
        res["backend"] = backend

    return res


def adeu_read_live_document() -> dict:
    """
    Read the currently active Word document via Adeu's COM bridge.
    Windows only; requires `adeu extract --live` CLI command.
    """
    if not _adeu_installed():
        return _unavailable("adeu CLI not installed")
    try:
        result = subprocess.run(
            [_get_adeu_cmd(), "extract", "--live"],
            capture_output=True,
            text=True,
            timeout=20,
        )
        if result.returncode != 0:
            return _error(f"adeu extract --live failed: {result.stderr[:400]}")
        markdown = result.stdout
        paragraphs = _parse_markdown_to_paragraphs(markdown)
        for p in paragraphs:
            if p["is_heading"]:
                p["style"] = f"Heading {p['level']}"
            else:
                p["style"] = "Normal"

        headings = [
            {
                "index": p["index"],
                "level": p["level"],
                "text": p["text"],
            }
            for p in paragraphs if p["is_heading"]
        ]

        return {
            "ok": True,
            "data": {
                "full_text": markdown,
                "paragraphs": paragraphs,
                "headings": headings,
                "paragraph_count": len(paragraphs),
                "format": "criticmarkup",
            },
        }
    except subprocess.TimeoutExpired:
        return _error("adeu extract --live timed out")
    except Exception:
        return _error(traceback.format_exc())


def adeu_sanitize(file_path: str, output_path: str | None = None) -> dict:
    """
    Strip dangerous metadata from DOCX before sharing (adeu sanitize).
    Removes author names, comments history, revision tracking metadata.
    """
    file_path = str(Path(file_path).resolve())
    if not Path(file_path).exists():
        return _error(f"File not found: {file_path}")
    if not _adeu_installed():
        return _unavailable("adeu CLI not installed")

    if output_path is None:
        base, ext = os.path.splitext(file_path)
        output_path = f"{base}_sanitized{ext}"

    try:
        subprocess.run(
            [_get_adeu_cmd(), "sanitize", file_path, "-o", output_path],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
        log.info("adeu_sanitize OK: %s → %s", file_path, output_path)
        return {"ok": True, "data": {"output_path": output_path}}
    except subprocess.CalledProcessError as e:
        return _error(f"adeu sanitize failed: {e.stderr[:400]}")
    except subprocess.TimeoutExpired:
        return _error("adeu sanitize timed out")
    except Exception:
        return _error(traceback.format_exc())


# ---------------------------------------------------------------------------
# Internal strategies
# ---------------------------------------------------------------------------

def _adeu_sdk_apply(
    file_path: str,
    edits: list[dict],
    output_path: str | None,
    author: str,
) -> dict:
    """File-based Track Changes via adeu Python SDK (no COM, works headless)."""
    try:
        from adeu import RedlineEngine, ModifyText  # type: ignore

        if output_path is None:
            base, ext = os.path.splitext(file_path)
            output_path = f"{base}_redlined{ext}"

        with open(file_path, "rb") as f:
            stream = BytesIO(f.read())

        engine = RedlineEngine(stream, author=author)

        modify_ops = []
        for edit in edits:
            target = edit.get("target_text", "").strip()
            new = edit.get("new_text", "").strip()
            comment = edit.get("comment", "") or ""
            if not target or not new:
                log.warning("Skipping edit with empty target_text or new_text: %r", edit)
                continue
            modify_ops.append(
                ModifyText(target_text=target, new_text=new, comment=comment or None)
            )

        if not modify_ops:
            return _error("All edits were invalid (empty target_text or new_text)")

        # Pre-validate to surface missing-text errors before applying
        validation_errors = engine.validate_edits(modify_ops)
        if validation_errors:
            log.warning(
                "adeu validate_edits found %d issues: %s",
                len(validation_errors), validation_errors[:3]
            )
            # Non-fatal: apply anyway (engine skips unresolvable edits gracefully)

        # apply_edits returns (applied_count, skipped_count)
        applied_count, skipped_count = engine.apply_edits(modify_ops)

        with open(output_path, "wb") as f:
            f.write(engine.save_to_stream().getvalue())

        log.info(
            "adeu_sdk_apply OK: %d applied, %d skipped → %s",
            applied_count, skipped_count, output_path
        )
        return {
            "ok": True,
            "applied_count": applied_count,
            "skipped_count": skipped_count,
            "backend": "adeu_sdk",
            "data": {
                "output_path": output_path,
                "applied_count": applied_count,
                "skipped_count": skipped_count,
                "validation_warnings": validation_errors[:5] if validation_errors else [],
                "backend": "adeu_sdk",
            },
        }
    except ImportError:
        return _unavailable("adeu Python SDK import failed (pip install adeu)")
    except Exception:
        return _error(traceback.format_exc())

def _python_docx_tracked_fallback(
    file_path: str,
    edits: list[dict],
    output_path: str | None,
    author: str,
) -> dict:
    """Fallback Track Changes using python-docx XML w:ins/w:del revision marks."""
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import datetime

        if output_path is None:
            base, ext = os.path.splitext(file_path)
            output_path = f"{base}_redlined{ext}"

        doc = Document(file_path)
        
        # Enable track revisions via document settings XML
        settings_element = doc.settings.element
        track_revisions = settings_element.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings_element.append(track_revisions)
        timestamp = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

        applied_count = 0
        for edit in edits:
            target = edit.get("target_text", "").strip()
            new = edit.get("new_text", "")
            if not target:
                continue
            if new is None:
                new = ""

            for p in doc.paragraphs:
                if target in p.text:
                    text = p.text
                    p.text = ""
                    parts = text.split(target, 1)
                    if parts[0]:
                        p.add_run(parts[0])

                    # Add w:del
                    w_del = OxmlElement('w:del')
                    w_del.set(qn('w:author'), author)
                    w_del.set(qn('w:date'), timestamp)
                    w_del.set(qn('w:id'), str(applied_count * 2))
                    r_del = OxmlElement('w:r')
                    del_text = OxmlElement('w:delText')
                    del_text.text = target
                    r_del.append(del_text)
                    w_del.append(r_del)
                    p._p.append(w_del)

                    # Add w:ins (only if new is not empty)
                    if new:
                        if "\n" in new:
                            lines = new.split("\n")
                            # First line goes to the current paragraph
                            if lines[0]:
                                w_ins = OxmlElement('w:ins')
                                w_ins.set(qn('w:author'), author)
                                w_ins.set(qn('w:date'), timestamp)
                                w_ins.set(qn('w:id'), str(applied_count * 2 + 1))
                                r_ins = OxmlElement('w:r')
                                ins_text = OxmlElement('w:t')
                                ins_text.text = lines[0]
                                r_ins.append(ins_text)
                                w_ins.append(r_ins)
                                p._p.append(w_ins)
                            
                            # The remaining lines are added as new paragraphs with track changes (w:ins)
                            current_p = p
                            for idx, line in enumerate(lines[1:]):
                                next_p_elem = OxmlElement('w:p')
                                current_p._p.addnext(next_p_elem)
                                from docx.text.paragraph import Paragraph
                                next_p = Paragraph(next_p_elem, doc)
                                
                                w_ins_next = OxmlElement('w:ins')
                                w_ins_next.set(qn('w:author'), author)
                                w_ins_next.set(qn('w:date'), timestamp)
                                w_ins_next.set(qn('w:id'), str(applied_count * 1000 + idx))
                                r_ins_next = OxmlElement('w:r')
                                ins_text_next = OxmlElement('w:t')
                                ins_text_next.text = line
                                r_ins_next.append(ins_text_next)
                                w_ins_next.append(r_ins_next)
                                next_p_elem.append(w_ins_next)
                                current_p = next_p
                        else:
                            w_ins = OxmlElement('w:ins')
                            w_ins.set(qn('w:author'), author)
                            w_ins.set(qn('w:date'), timestamp)
                            w_ins.set(qn('w:id'), str(applied_count * 2 + 1))
                            r_ins = OxmlElement('w:r')
                            ins_text = OxmlElement('w:t')
                            ins_text.text = new
                            r_ins.append(ins_text)
                            w_ins.append(r_ins)
                            p._p.append(w_ins)

                    if len(parts) > 1 and parts[1]:
                        p.add_run(parts[1])

                    applied_count += 1
                    break

        doc.save(output_path)
        skipped_count = len(edits) - applied_count
        return {
            "ok": True,
            "applied_count": applied_count,
            "skipped_count": skipped_count,
            "backend": "python_docx_fallback",
            "data": {
                "output_path": output_path,
                "applied_count": applied_count,
                "skipped_count": skipped_count,
                "backend": "python_docx_fallback",
            }
        }
    except Exception as e:
        return _error(f"Tracked changes fallback failed: {e}")


def _adeu_cli_apply(
    file_path: str,
    edits: list[dict],
    output_path: str | None,
    author: str,
) -> dict:
    """CLI-based fallback using `adeu apply` subcommand."""
    if output_path is None:
        base, ext = os.path.splitext(file_path)
        output_path = f"{base}_redlined{ext}"

    # Write edits to a temp JSON file
    tmp_fd, tmp_json = tempfile.mkstemp(suffix=".json")
    os.close(tmp_fd)
    try:
        with open(tmp_json, "w", encoding="utf-8") as f:
            json.dump(edits, f)

        subprocess.run(
            [
                _get_adeu_cmd(), "apply", file_path,
                "--edits", tmp_json,
                "--author", author,
                "-o", output_path,
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )
        log.info("adeu_cli_apply OK: %d edits → %s", len(edits), output_path)
        return {
            "ok": True,
            "applied_count": len(edits),
            "skipped_count": 0,
            "backend": "adeu_cli",
            "data": {
                "output_path": output_path,
                "applied_count": len(edits),
                "skipped_count": 0,
                "backend": "adeu_cli",
            },
        }
    except subprocess.CalledProcessError as e:
        return _error(f"adeu apply failed: {e.stderr[:500]}")
    except subprocess.TimeoutExpired:
        return _error("adeu apply CLI timed out after 30 s")
    except Exception:
        return _error(traceback.format_exc())
    finally:
        try:
            os.unlink(tmp_json)
        except OSError:
            pass


def _adeu_live_apply(
    file_path: str,
    edits: list[dict],
    author: str,
) -> dict:
    """
    Live COM-macro injection via `adeu apply --live`.
    Only called when Word has the file open.
    """
    tmp_fd, tmp_json = tempfile.mkstemp(suffix=".json")
    os.close(tmp_fd)
    try:
        with open(tmp_json, "w", encoding="utf-8") as f:
            json.dump(edits, f)

        result = subprocess.run(
            [
                _get_adeu_cmd(), "apply",
                tmp_json,
                "--live",
                "--author", author,
            ],
            capture_output=True,
            text=True,
            timeout=20,
        )
        if result.returncode != 0:
            return _error(f"adeu apply --live failed: {result.stderr[:400]}")

        log.info("adeu_live_apply OK: %d edits injected via COM", len(edits))
        return {
            "ok": True,
            "applied_count": len(edits),
            "skipped_count": 0,
            "backend": "adeu_live_com",
            "data": {
                "applied_count": len(edits),
                "skipped_count": 0,
                "backend": "adeu_live_com",
            },
        }
    except subprocess.TimeoutExpired:
        return _error("adeu process-active timed out")
    except Exception:
        return _error(traceback.format_exc())
    finally:
        try:
            os.unlink(tmp_json)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_markdown_to_paragraphs(markdown: str) -> list[dict]:
    """
    Split CriticMarkup-flavoured markdown into indexed paragraph blocks.
    Each block gets an `is_heading` flag, `level`, and `is_table` flag.
    """
    paragraphs: list[dict] = []
    for i, block in enumerate(markdown.split("\n\n")):
        stripped = block.strip()
        if not stripped:
            continue
        level = 0
        if stripped.startswith("# "):
            level = 1
        elif stripped.startswith("## "):
            level = 2
        elif stripped.startswith("### "):
            level = 3
        elif stripped.startswith("#### "):
            level = 4
        paragraphs.append(
            {
                "index": i,
                "text": stripped,
                "level": level,
                "is_heading": level > 0,
                "is_table": ("|" in stripped and stripped.startswith("|")),
            }
        )
    return paragraphs


def _error(msg: str) -> dict:
    log.error("AdeuBridge error: %s", msg)
    return {"ok": False, "data": None, "error": msg}


def _unavailable(msg: str) -> dict:
    log.warning("AdeuBridge unavailable: %s", msg)
    return {"ok": False, "data": None, "error": msg, "unavailable": True}
