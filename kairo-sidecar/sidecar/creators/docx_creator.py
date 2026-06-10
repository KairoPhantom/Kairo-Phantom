"""
sidecar/creators/docx_creator.py — Kairo Phantom Create-From-Scratch DOCX
==========================================================================
Creates a new .docx file from a structured content dict and opens it in Word.

Usage
-----
    from sidecar.creators.docx_creator import DocxCreator

    creator = DocxCreator()
    path = creator.create_and_open({
        "title": "Quarterly Business Review",
        "author": "Kairo Phantom",
        "sections": [
            {
                "heading": "Executive Summary",
                "level": 1,
                "paragraphs": ["Q3 revenue exceeded targets by 12%..."]
            },
            {
                "heading": "Key Metrics",
                "level": 2,
                "table": {
                    "headers": ["Metric", "Target", "Actual"],
                    "rows": [["Revenue", "$1M", "$1.12M"]]
                }
            }
        ]
    })
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List

log = logging.getLogger("kairo-sidecar.docx_creator")

# Default Kairo output directory
KAIRO_DOCS_DIR = Path.home() / "Documents" / "Kairo"


class DocxCreator:
    """
    Creates new .docx documents from a structured content dict.
    Saves to ~/Documents/Kairo/ and opens in Word via os.startfile().
    No CUA required — pure python-docx.
    """

    def create(
        self,
        content: Dict[str, Any],
        output_path: Optional[str] = None,
    ) -> str:
        """
        Creates a .docx from structured content and saves to disk.

        Parameters
        ----------
        content     : Dict with keys:
                      - title (str): Document title
                      - author (str): Author name
                      - sections (list): Each section dict may have:
                          - heading (str): Section heading text
                          - level (int): Heading level 1-4 (default 1)
                          - paragraphs (list[str]): Body paragraphs
                          - bullets (list[str]): Bullet list items
                          - table (dict): {"headers": [...], "rows": [[...]]}
        output_path : Full path to save the file. Auto-generated if None.

        Returns
        -------
        str — absolute path of the created file.
        """
        from docx import Document

        doc = Document()

        # Set core properties
        title = content.get("title", "Untitled Document")
        author = content.get("author", "Kairo Phantom")
        try:
            doc.core_properties.author = author
            doc.core_properties.title = title
        except Exception:
            pass

        # Title heading (level 0 = title style)
        if title:
            doc.add_heading(title, level=0)

        # Process sections
        for section in content.get("sections", []):
            heading_text = section.get("heading", "")
            level = max(1, min(4, int(section.get("level", 1))))

            if heading_text:
                doc.add_heading(heading_text, level=level)

            # Body paragraphs
            for para_text in section.get("paragraphs", []):
                if para_text:
                    doc.add_paragraph(str(para_text))

            # Bullet list items
            for bullet in section.get("bullets", []):
                if bullet:
                    para = doc.add_paragraph(str(bullet), style="List Bullet")

            # Table
            table_data = section.get("table")
            if table_data and isinstance(table_data, dict):
                headers = table_data.get("headers", [])
                rows_data = table_data.get("rows", [])
                cols = len(headers)
                if cols > 0:
                    table = doc.add_table(rows=1, cols=cols)
                    try:
                        table.style = "Table Grid"
                    except Exception:
                        pass
                    # Header row
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr_cells[i].text = str(h)
                    # Data rows
                    for row_data in rows_data:
                        row = table.add_row()
                        for i, cell_val in enumerate(row_data):
                            if i < cols:
                                row.cells[i].text = str(cell_val)

        # Resolve output path
        if not output_path:
            KAIRO_DOCS_DIR.mkdir(parents=True, exist_ok=True)
            safe_title = "".join(
                c if c.isalnum() or c in " _-" else "_" for c in title
            ).strip()[:50] or "Document"
            output_path = str(KAIRO_DOCS_DIR / f"{safe_title}.docx")

        doc.save(output_path)
        log.info(f"DocxCreator: saved to {output_path}")
        return output_path

    def create_and_open(self, content: Dict[str, Any]) -> str:
        """
        Creates a .docx and opens it in the default Word application.
        Returns the path of the created file.
        """
        path = self.create(content)
        try:
            os.startfile(path)
            log.info(f"DocxCreator: opened {path}")
        except Exception as e:
            log.warning(f"DocxCreator: could not auto-open {path}: {e}")
        return path
