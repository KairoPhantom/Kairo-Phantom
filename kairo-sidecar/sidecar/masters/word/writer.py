"""
sidecar/masters/word/writer.py
================================
Atomic, XML-level Word document writer using python-docx.

All insertions use ref_para._element.addnext() rather than doc.add_paragraph()
(which always appends to the end). Saves use os.replace() atomicity.
"""

from __future__ import annotations

import os
import tempfile
import logging

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

log = logging.getLogger("kairo-sidecar.word.writer")


class WordWriter:
    """
    XML-level paragraph writer for python-docx Documents.

    All mutating methods operate directly on the OOXML element tree so that
    insertion position is exact and formatting runs are preserved.
    """

    # ── Paragraph insertion ───────────────────────────────────────────────────

    def _insert_paragraph(
        self,
        doc: Document,
        after_idx: int,
        text: str,
        style_name: str = "Normal",
    ) -> Paragraph:
        """
        Insert a new paragraph AFTER the paragraph at *after_idx*.

        Implementation uses XML-level insertion via addnext() — NOT
        doc.add_paragraph() which always appends to the end of the body.

        Parameters
        ----------
        doc : Document
            Open python-docx Document.
        after_idx : int
            Zero-based paragraph index. The new paragraph is placed immediately
            after this index. Pass -1 to append at the end of the body.
        text : str
            Plain-text content for the new paragraph.
        style_name : str
            Style name that must already exist in the document.

        Returns
        -------
        Paragraph
            The newly-created Paragraph object.
        """
        # Create a bare <w:p> element
        p_elem = OxmlElement("w:p")

        if 0 <= after_idx < len(doc.paragraphs):
            ref_para = doc.paragraphs[after_idx]
            # CRITICAL: addnext places p_elem immediately AFTER ref_para in the body
            ref_para._element.addnext(p_elem)
        else:
            # Append to body
            doc.element.body.append(p_elem)

        new_para = Paragraph(p_elem, doc)

        # Apply style (best-effort; fall back to Normal if not found)
        try:
            new_para.style = doc.styles[style_name]
        except (KeyError, Exception) as exc:
            log.debug(
                f"WordWriter._insert_paragraph: style '{style_name}' not found ({exc}); using Normal"
            )
            try:
                new_para.style = doc.styles["Normal"]
            except Exception:
                pass

        # Add a single run with the text
        if text:
            new_para.add_run(text)

        return new_para

    # ── Paragraph editing ─────────────────────────────────────────────────────

    def _edit_paragraph(self, doc: Document, idx: int, new_text: str) -> None:
        """
        Replace the text of paragraph *idx* while preserving the paragraph element.

        Clears all existing runs and adds a single new run with *new_text*.
        The paragraph's style is preserved unchanged.

        Parameters
        ----------
        doc : Document
        idx : int
            Zero-based paragraph index.
        new_text : str
            Replacement text.
        """
        if idx < 0 or idx >= len(doc.paragraphs):
            log.warning(
                f"WordWriter._edit_paragraph: index {idx} out of range (0-{len(doc.paragraphs)-1})"
            )
            return

        para = doc.paragraphs[idx]

        # Remove all <w:r> (run) children from the paragraph element
        for run_elem in list(para._p.findall(qn("w:r"))):
            para._p.remove(run_elem)

        # Add a fresh run with the replacement text
        para.add_run(new_text)

    # ── Style change ──────────────────────────────────────────────────────────

    def _change_style(self, doc: Document, idx: int, style_name: str) -> None:
        """
        Change the paragraph style at *idx* without altering its text content.

        Parameters
        ----------
        doc : Document
        idx : int
        style_name : str
            Must be an existing style name in this document.
        """
        if idx < 0 or idx >= len(doc.paragraphs):
            log.warning(f"WordWriter._change_style: index {idx} out of range")
            return

        para = doc.paragraphs[idx]
        try:
            para.style = doc.styles[style_name]
        except (KeyError, Exception) as exc:
            log.warning(f"WordWriter._change_style: style '{style_name}' not found ({exc})")

    # ── Paragraph deletion ────────────────────────────────────────────────────

    def _delete_paragraph(self, doc: Document, idx: int) -> None:
        """
        Remove paragraph *idx* from the document body using XML removal.

        Implementation: para._element.getparent().remove(para._element)

        Parameters
        ----------
        doc : Document
        idx : int
            Zero-based paragraph index.
        """
        if idx < 0 or idx >= len(doc.paragraphs):
            log.warning(f"WordWriter._delete_paragraph: index {idx} out of range")
            return

        para = doc.paragraphs[idx]
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)
        else:
            log.warning(f"WordWriter._delete_paragraph: paragraph {idx} has no parent element")

    # ── Atomic save ───────────────────────────────────────────────────────────

    def atomic_save_docx(self, doc: Document, path: str) -> None:
        """
        Save *doc* to *path* atomically by:
          1. Writing to a sibling .tmp file in the same directory.
          2. Calling os.replace() to swap it in — single-syscall, kernel-guaranteed atomic.

        This ensures the original file is NEVER partially overwritten.
        If the save fails, the .tmp file is cleaned up and the original is intact.

        Parameters
        ----------
        doc : Document
            The python-docx Document to save.
        path : str
            Absolute path where the document should be written.

        Raises
        ------
        OSError / PermissionError
            If the file is locked by another process (e.g. Word has it open).
        """
        # Write to a temp file in the same directory (same filesystem → atomic rename)
        dir_name = os.path.dirname(os.path.abspath(path))
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx.tmp", dir=dir_name)
        os.close(tmp_fd)

        try:
            doc.save(tmp_path)
            # os.replace is atomic on POSIX and uses MoveFileExW on Windows
            os.replace(tmp_path, path)
            log.debug(f"WordWriter.atomic_save_docx: saved to {path}")
        except Exception as exc:
            # Clean up the temp file; original is untouched
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            log.error(f"WordWriter.atomic_save_docx: failed to save to {path}: {exc}")
            raise
