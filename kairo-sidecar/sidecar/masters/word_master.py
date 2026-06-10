import os
import shutil
import logging
import traceback
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

log = logging.getLogger("kairo-sidecar.word_master")

# ---------------------------------------------------------------------------
# Optional Docling integration
# ---------------------------------------------------------------------------
try:
    from sidecar.parsers.docling_parser import parse_docx_structured as _docling_parse_docx
    from sidecar.parsers.docling_parser import _DOCLING_AVAILABLE as _REAL_DOCLING_AVAILABLE
    _DOCLING_AVAILABLE = _REAL_DOCLING_AVAILABLE
    log.debug(f"word_master: Docling integration available: {_DOCLING_AVAILABLE}")
except ImportError as e:
    _docling_parse_docx = None  # type: ignore
    _DOCLING_AVAILABLE = False
    log.debug(f"word_master: Docling not available ({e}), using python-docx only.")

@dataclass
class WordContext:
    styles: Dict[str, List[str]]
    paragraphs: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    theme_fonts: Dict[str, str]
    list_sequences: List[Dict[str, Any]]
    document_purpose: str
    cursor_paragraph_index: int
    total_paragraphs: int

    def to_dict(self):
        return asdict(self)

class ValidationResult:
    def __init__(self, valid: bool, error: str = "", op: dict = None):
        self.valid = valid
        self.error = error
        self.op = op

class WordContextExtractor:
    """Reads COMPLETE document state before any LLM call."""

    def extract(self, file_path: str, cursor_paragraph_index: int) -> WordContext:
        doc = Document(file_path)

        # 1. Style inventory (sourced from THIS document, not defaults) in a single pass
        paragraph_styles = []
        character_styles = []
        table_styles = []
        style_id_to_name = {}
        for s in doc.styles:
            name = s.name
            if name:
                style_id_to_name[s.style_id] = name
            stype = s.type
            if stype == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append(name)
            elif stype == WD_STYLE_TYPE.CHARACTER:
                character_styles.append(name)
            elif stype == WD_STYLE_TYPE.TABLE:
                table_styles.append(name)
        
        styles = {
            "paragraph": paragraph_styles,
            "character": character_styles,
            "table": table_styles,
        }

        # 2. Paragraph inventory with positions
        paragraphs = []
        _docling_used = False

        if _DOCLING_AVAILABLE and _docling_parse_docx is not None:
            try:
                docling_result = _docling_parse_docx(file_path)
                docling_paras = docling_result.get("paragraphs", [])
                if docling_paras:
                    # Merge Docling output with WordContext-compatible shape
                    for dp in docling_paras:
                        text = dp.get("text", "")
                        runs = dp.get("runs", [])
                        paragraphs.append({
                            "index": dp.get("index", len(paragraphs)),
                            "style": dp.get("style", "Normal"),
                            "text": text[:200],
                            "level": dp.get("level", None),
                            "is_empty": len(text.strip()) == 0,
                            "has_runs": len(runs) > 0,
                            "page": dp.get("page"),
                            "runs": runs,
                        })
                    _docling_used = True
                    log.info(
                        "WordContextExtractor: Docling paragraph inventory used (tier=%s, count=%d)",
                        docling_result.get("metadata", {}).get("tier", "docling"),
                        len(paragraphs),
                    )
            except Exception as exc:
                log.debug(f"WordContextExtractor: Docling parse failed ({exc}), falling back to python-docx.")

        # Fallback: python-docx paragraph extraction
        list_sequences = []
        if not _docling_used:
            for i, para in enumerate(doc.paragraphs):
                pPr = para._p.pPr
                level = None
                style_name = None

                if pPr is not None:
                    try:
                        numPr = pPr.numPr
                        if numPr is not None and numPr.ilvl is not None:
                            level = numPr.ilvl.val
                    except Exception as exc:
                        log.debug(f"Failed to extract list level: {exc}")

                    pStyle = pPr.pStyle
                    if pStyle is not None and pStyle.val is not None:
                        style_name = style_id_to_name.get(pStyle.val, pStyle.val)

                if not style_name:
                    sid = para._p.style
                    if sid:
                        style_name = style_id_to_name.get(sid)
                        if not style_name:
                            try:
                                style_name = para.style.name if para.style else sid
                            except Exception:
                                style_name = sid
                    else:
                        style_name = "Normal"

                # Speed up text extraction
                t_elements = para._p.xpath('.//w:t')
                para_text = "".join(t.text for t in t_elements if t.text is not None)

                # Speed up run traversal
                has_runs = bool(para._p.xpath('.//w:r'))

                paragraphs.append({
                    "index": i,
                    "style": style_name,
                    "text": para_text[:200],  # truncated for context
                    "level": level,
                    "is_empty": len(para_text.strip()) == 0,
                    "has_runs": has_runs,
                })

                # Build list_sequences in the first and only pass!
                if style_name and ("List" in style_name or "bullet" in style_name.lower() or "number" in style_name.lower()):
                    list_sequences.append({
                        "index": i,
                        "style": style_name,
                        "text": para_text[:100]
                    })
        else:
            for p in paragraphs:
                style_name = p.get("style", "")
                if style_name and ("List" in style_name or "bullet" in style_name.lower() or "number" in style_name.lower()):
                    list_sequences.append({
                        "index": p["index"],
                        "style": style_name,
                        "text": p["text"][:100]
                    })

        # 3. Table inventory (pre-map table elements to paragraph index positions to avoid O(N*M))
        tbl_to_para_index = {}
        p_idx = -1
        for child in doc.element.body:
            tag = child.tag
            if tag.endswith('p'):
                p_idx += 1
            elif tag.endswith('tbl'):
                tbl_to_para_index[child] = p_idx

        tables = []
        for i, table in enumerate(doc.tables):
            tables.append({
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "after_paragraph": tbl_to_para_index.get(table._element, -1),
                "header_text": [cell.text[:50] for cell in table.rows[0].cells] if table.rows else [],
            })

        # 4. Theme fonts
        theme_fonts = self._extract_theme_fonts(doc)

        # 6. Document purpose detection
        purpose = self._detect_document_purpose(paragraphs, doc)

        # Safe guard cursor_paragraph_index
        total_paras = len(paragraphs)
        if cursor_paragraph_index < 0 or cursor_paragraph_index >= total_paras:
            cursor_paragraph_index = max(0, total_paras - 1)

        return WordContext(
            styles=styles,
            paragraphs=paragraphs,
            tables=tables,
            theme_fonts=theme_fonts,
            list_sequences=list_sequences,
            document_purpose=purpose,
            cursor_paragraph_index=cursor_paragraph_index,
            total_paragraphs=total_paras,
        )

    def _find_table_position(self, doc, table) -> int:
        p_idx = -1
        for child in doc.element.body:
            if child.tag.endswith('p'):
                p_idx += 1
            elif child.tag.endswith('tbl'):
                if child is table._element:
                     return p_idx
        return -1

    def _extract_theme_fonts(self, doc) -> Dict[str, str]:
        fonts = {"major": "Calibri", "minor": "Calibri"}
        try:
            for s in doc.styles:
                if s.font and s.font.name:
                    fonts["major"] = s.font.name
                    break
        except Exception as exc:
            log.debug(f"Failed to extract theme fonts: {exc}")
        return fonts

    def _extract_list_sequences(self, doc) -> List[Dict[str, Any]]:
        sequences = []
        for i, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else ""
            if style_name and ("List" in style_name or "bullet" in style_name.lower() or "number" in style_name.lower()):
                sequences.append({"index": i, "style": style_name, "text": p.text[:100]})
        return sequences

    def _detect_document_purpose(self, paragraphs: List[Dict[str, Any]], doc) -> str:
        # Heuristic detection using pre-extracted paragraphs list
        first_elements = [p["text"] for p in paragraphs[:15] if p["text"].strip()]
        first_text = "\n".join(first_elements)
        first_500 = first_text[:500].lower()

        first_para_style = ""
        first_para_text = ""
        if paragraphs:
            first_para_style = (paragraphs[0]["style"] or "").lower()
            first_para_text = paragraphs[0]["text"].lower()

        is_legal_first = "whereas" in first_para_text or "agreement" in first_para_text or "legal" in first_para_style

        # Presence of footnotes (direct XML check is much faster)
        has_footnotes = False
        try:
            has_footnotes = bool(doc.element.body.xpath('./w:p[position() <= 50]//w:footnoteReference'))
        except Exception as e:
            log.debug(f"Footnote detection error: {e}")

        # Presence of numbered sections like "1.1.1" or "1.2"
        has_numbered_sections = False
        numbered_pattern = r"^\s*\d+(\.\d+){1,3}\s+"
        for p in paragraphs[:50]:
            if re.match(numbered_pattern, p["text"]):
                has_numbered_sections = True
                break

        # Average sentence length
        all_text = "\n".join(p["text"] for p in paragraphs[:50] if p["text"].strip())
        sentences = [s.strip() for s in re.split(r"[.!?]", all_text) if s.strip()]
        avg_sentence_len = 0
        if sentences:
            avg_sentence_len = sum(len(s.split()) for s in sentences) / len(sentences)

        if is_legal_first or any(w in first_500 for w in ["whereas", "parties agree", "non-disclosure", "nda", "confidentiality"]):
            return "legal"
        if has_footnotes or any(w in first_500 for w in ["abstract", "introduction", "methodology", "references", "hypothesis"]):
            return "academic"
        if has_numbered_sections or any(w in first_500 for w in ["installation", "api documentation", "deployment", "configuration"]):
            return "technical"
        if avg_sentence_len > 25:
            return "legal"

        return "business_memo"


class WordOperationValidator:
    """Validates DocxOperation against THIS document's actual state."""

    def validate(self, op: dict, context: WordContext) -> ValidationResult:
        op_type = op.get("type", op.get("action", ""))
        
        if op_type in ("insert_paragraph", "replace_paragraph", "append", "insert_after_heading"):
            style = op.get("style", "Normal")
            # CRITICAL: verify style actually exists in this document
            if style not in context.styles["paragraph"]:
                corrected = self._fuzzy_style_match(style, context.styles["paragraph"])
                if corrected:
                    op["style"] = corrected
                else:
                    return ValidationResult(
                        valid=False,
                        error=f"Style '{style}' not in document. Available: {context.styles['paragraph'][:5]}"
                    )

        if op_type == "insert_paragraph":
            after_idx = op.get("after_paragraph_index", -1)
            if after_idx >= context.total_paragraphs:
                op["after_paragraph_index"] = context.total_paragraphs - 1
            elif after_idx < -1:
                op["after_paragraph_index"] = -1

        elif op_type in ("replace_paragraph", "delete_paragraph", "append_to_run"):
            idx = op.get("paragraph_index", 0)
            if context.total_paragraphs > 0:
                op["paragraph_index"] = max(0, min(idx, context.total_paragraphs - 1))
            else:
                op["paragraph_index"] = 0

        return ValidationResult(valid=True, op=op)

    def _fuzzy_style_match(self, requested: str, available: List[str]) -> str | None:
        normalized = requested.replace(" ", "").replace("_", "").replace("-", "").lower()
        
        # 1. Exact normalized match
        for style in available:
            if style.replace(" ", "").replace("_", "").replace("-", "").lower() == normalized:
                return style
                
        # 2. Substring fallback match
        for style in available:
            style_norm = style.replace(" ", "").replace("_", "").replace("-", "").lower()
            if normalized in style_norm or style_norm in normalized:
                return style
                
        # 3. Known style family aliases mapping
        aliases = {
            "heading1": ["heading 1", "heading1", "h1"],
            "heading2": ["heading 2", "heading2", "h2"],
            "heading3": ["heading 3", "heading3", "h3"],
            "heading4": ["heading 4", "heading4", "h4"],
            "listbullet": ["list bullet", "listbullet", "bullet", "bullets", "list paragraph"],
            "listnumber": ["list number", "listnumber", "number", "numbers"],
            "quote": ["quote", "blockquote"],
            "normal": ["normal", "body", "paragraph", "body text"]
        }
        for canonical, alias_list in aliases.items():
            if normalized in alias_list:
                for style in available:
                    style_norm = style.replace(" ", "").replace("_", "").replace("-", "").lower()
                    if style_norm == canonical or style_norm in alias_list:
                        return style
        return None


class WordAgent:
    """
    COM/app APIs wrapper for Word matching UseIt-AI/use-it-agent interface.
    """
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.word_app = None

    def connect(self) -> bool:
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            path = os.path.abspath(self.file_path)
            try:
                self.word_app = win32com.client.GetActiveObject("Word.Application")
                for doc in self.word_app.Documents:
                    try:
                        if os.path.abspath(doc.FullName) == path:
                            self.doc = doc
                            break
                    except Exception:
                        pass
            except Exception:
                pass

            return self.doc is not None
        except Exception:
            return False

    def apply_operations(self, operations: List[dict]) -> bool:
        if not self.doc:
            return False
        try:
            from sidecar.writers.docx_writer import _try_com_write
            com_ops = []
            for op in operations:
                cop = op.copy()
                if "type" in cop and "action" not in cop:
                    cop["action"] = cop["type"]
                com_ops.append(cop)
            res = _try_com_write(self.file_path, com_ops)
            return res.get("ok", False)
        except Exception as e:
            log.warning(f"WordAgent.apply_operations failed: {e}")
            return False


class WordWriter:
    """python-docx writer that preserves ALL formatting details."""

    def apply_operations(self, file_path: str, operations: List[dict], context: WordContext) -> dict:
        # Try WordAgent() COM path first (UseIt-AI integration)
        try:
            agent = WordAgent(file_path)
            if agent.connect():
                log.info("UseIt-AI: Applying live injection using WordAgent COM API...")
                success = agent.apply_operations(operations)
                if success:
                    return {"ok": True, "applied_count": len(operations), "errors": [], "path": file_path}
        except Exception as e:
            log.debug(f"WordAgent COM live injection failed: {e}. Falling back to standard COM/python-docx write.")
        # COM Write path for active document
        try:
            import win32com.client
            import pythoncom
            import subprocess
            pythoncom.CoInitialize()
            
            word = None
            target_doc = None
            path = os.path.abspath(file_path)
            
            word_running = False
            try:
                out = subprocess.run(["tasklist", "/FI", "IMAGENAME eq winword.exe"], capture_output=True, text=True)
                word_running = "winword.exe" in out.stdout.lower()
            except Exception as exc:
                log.debug(f"Tasklist task checking failed: {exc}")
                
            if word_running:
                try:
                    word = win32com.client.GetActiveObject("Word.Application")
                    for doc in word.Documents:
                        try:
                            if os.path.abspath(doc.FullName) == path:
                                target_doc = doc
                                break
                        except Exception as doc_exc:
                            log.debug(f"COM Doc FullName check failed: {doc_exc}")
                except Exception as act_exc:
                    log.debug(f"COM ActiveObject check failed: {act_exc}")
                        
            if target_doc:
                log.info("Applying Word operations via live COM...")
                from sidecar.writers.docx_writer import _try_com_write
                # Adapt operations key format ('type' -> 'action')
                com_ops = []
                for op in operations:
                    cop = op.copy()
                    if "type" in cop and "action" not in cop:
                        cop["action"] = cop["type"]
                    com_ops.append(cop)
                return _try_com_write(file_path, com_ops)
        except Exception as e:
            log.debug(f"Live COM check failed: {e}. Falling back to python-docx write.")

        # Fallback to python-docx file write
        backup_path = file_path + ".kairo_bak"
        tmp_path = file_path + ".kairo_tmp"

        # Ensure no stale temp/backup files exist before we start
        for p in (tmp_path, backup_path):
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception as exc:
                    log.debug(f"File removal failed: {exc}")

        try:
            # Copy backup before loading/modifying
            shutil.copy2(file_path, backup_path)

            doc = Document(file_path)

            # Pre-cache original paragraphs
            original_paragraphs = list(doc.paragraphs)

            # Sort operations in reverse index order to preserve indices
            sorted_ops = sorted(
                [op for op in operations if op.get("type", op.get("action")) in ("insert_paragraph", "replace_paragraph", "delete_paragraph", "append_to_run")],
                key=lambda x: x.get("paragraph_index", x.get("after_paragraph_index", 0)),
                reverse=True
            )

            applied = []
            errors = []

            for op in sorted_ops:
                op_type = op.get("type", op.get("action", ""))
                try:
                    if op_type == "insert_paragraph":
                        self._insert_paragraph(doc, op, context)
                        applied.append(op)
                    elif op_type == "replace_paragraph":
                        self._replace_paragraph(doc, op, original_paragraphs)
                        applied.append(op)
                    elif op_type == "delete_paragraph":
                        self._delete_paragraph(doc, op, original_paragraphs)
                        applied.append(op)
                    elif op_type == "append_to_run":
                        self._append_to_run(doc, op, original_paragraphs)
                        applied.append(op)
                except Exception as e:
                    errors.append(f"Operation {op_type} failed: {e}")

            # Non-index-shifting operations (like insert_table)
            other_ops = [op for op in operations if op.get("type", op.get("action")) not in ("insert_paragraph", "replace_paragraph", "delete_paragraph", "append_to_run")]
            for op in other_ops:
                op_type = op.get("type", op.get("action", ""))
                try:
                    if op_type == "insert_table":
                        self._insert_table(doc, op, original_paragraphs)
                        applied.append(op)
                    else:
                        errors.append(f"Unsupported operation: {op_type}")
                except Exception as e:
                    errors.append(f"Operation {op_type} failed: {e}")

            # Save atomically
            doc.save(tmp_path)
            os.replace(tmp_path, file_path)
            if os.path.exists(backup_path):
                os.remove(backup_path)
        except PermissionError:
            # Word has this file locked. COM write failed too. Rollback original file from backup
            if os.path.exists(backup_path):
                try:
                    shutil.copy2(backup_path, file_path)
                    os.remove(backup_path)
                except Exception as exc:
                    log.debug(f"Backup copy rollback failed: {exc}")
            # Clean up tmp_path if it was created
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception as exc:
                    log.debug(f"Backup temp file cleanup failed: {exc}")
            return {
                "error": "Word has this file locked. COM write failed too — close Word and retry.",
                "path": file_path,
            }
        except Exception as e:
            # Automated rollback on any other exception
            if os.path.exists(backup_path):
                try:
                    shutil.copy2(backup_path, file_path)
                    os.remove(backup_path)
                except Exception as exc:
                    log.debug(f"Rollback file copy failed: {exc}")
            # Clean up tmp_path if it was created
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception as exc:
                    log.debug(f"Temp file removal failed: {exc}")
            raise e

        return {
            "applied_count": len(applied),
            "errors": errors,
            "path": file_path
        }

    def _insert_paragraph(self, doc, op, context):
        """
        CORRECT paragraph insertion using XML manipulation.
        Never use doc.add_paragraph() for insertion — it always appends.
        Use paragraph._element.addnext() / addprevious() for correct positioning.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        after_idx = op.get("after_paragraph_index", -1)
        style = op.get("style", "Normal")

        p_elem = OxmlElement('w:p')
        if 0 <= after_idx < len(doc.paragraphs):
            ref_para = doc.paragraphs[after_idx]
            ref_para._element.addnext(p_elem)
        elif after_idx == -1 and len(doc.paragraphs) > 0:
            doc.paragraphs[-1]._element.addnext(p_elem)
        else:
            doc.element.body.append(p_elem)

        new_para = Paragraph(p_elem, doc)
        try:
            new_para.style = doc.styles[style]
        except Exception as exc:
            log.debug(f"Failed to set style on paragraph: {exc}")

        # Add runs with formatting
        for run_data in op.get("runs", []):
            run = new_para.add_run(run_data.get("text", ""))
            run.bold = run_data.get("bold", False)
            run.italic = run_data.get("italic", False)

    def _replace_paragraph(self, doc, op, original_paragraphs):
        idx = op.get("paragraph_index", -1)
        style = op.get("style", "Normal")
        
        if 0 <= idx < len(original_paragraphs):
            para = original_paragraphs[idx]
            for run in para.runs:
                run.text = ""
            try:
                para.style = doc.styles[style]
            except Exception as exc:
                log.debug(f"Failed to set style on replaced paragraph: {exc}")
            for run_data in op.get("runs", []):
                run = para.add_run(run_data.get("text", ""))
                run.bold = run_data.get("bold", False)
                run.italic = run_data.get("italic", False)

    def _delete_paragraph(self, doc, op, original_paragraphs):
        idx = op.get("paragraph_index", -1)
        if 0 <= idx < len(original_paragraphs):
            para = original_paragraphs[idx]
            para._element.getparent().remove(para._element)

    def _append_to_run(self, doc, op, original_paragraphs):
        idx = op.get("paragraph_index", -1)
        if 0 <= idx < len(original_paragraphs):
            para = original_paragraphs[idx]
            for run_data in op.get("runs", []):
                run = para.add_run(run_data.get("text", ""))
                run.bold = run_data.get("bold", False)
                run.italic = run_data.get("italic", False)

    def _insert_table(self, doc, op, original_paragraphs):
        after_idx = op.get("after_paragraph_index", -1)
        headers = op.get("headers", [])
        rows_data = op.get("rows", [])

        cols = len(headers) if headers else (max(len(r) for r in rows_data) if rows_data else 0)
        if cols == 0:
            return

        table = doc.add_table(rows=0, cols=cols)
        table.style = "Table Grid"

        if headers:
            row = table.add_row()
            for i, h in enumerate(headers):
                row.cells[i].text = str(h)

        for r in rows_data:
            row = table.add_row()
            for i, cell_text in enumerate(r):
                if i < cols:
                    row.cells[i].text = str(cell_text)

        if 0 <= after_idx < len(original_paragraphs):
            ref_para = original_paragraphs[after_idx]
            ref_para._element.addnext(table._element)
        else:
            doc.element.body.append(table._element)


# ---------------------------------------------------------------------------
# WordMaster — Unified Façade (required by DomainMasterRouter)
# Wraps WordContextExtractor + WordOperationValidator + WordWriter + prompt
# into the standard master interface: extract_context / build_prompt /
# validate_operations / apply_operations / get_schema_class.
# ---------------------------------------------------------------------------

class WordMaster:
    """
    Unified Word Domain Master.
    Implements the standard interface used by DomainMasterRouter so every
    domain can be addressed uniformly.
    """

    def __init__(self):
        self._extractor = WordContextExtractor()
        self._validator = WordOperationValidator()
        self._writer = WordWriter()

    # --- Standard interface ---------------------------------------------------

    def extract_context(self, file_path: str, cursor_info=None) -> WordContext:
        """Extract full document context from a .docx file."""
        cursor_idx = 0
        if cursor_info is not None:
            try:
                cursor_idx = int(cursor_info)
            except (TypeError, ValueError):
                cursor_idx = 0
        return self._extractor.extract(file_path, cursor_idx)

    def _generate_reasoning(self, doc_context, mem_context: str) -> str:
        reasons = []
        if not mem_context:
            mem_context = ""
        # Check for style hints in mem_context
        if 'bullet' in mem_context.lower():
            reasons.append('Using List Bullet style based on your preferred format (MemMachine recall).')
        if 'heading' in mem_context.lower() or 'h2' in mem_context.lower():
            reasons.append('Using Heading 2 because surrounding headings are H2.')
        if not reasons:
            reasons.append('Analyzing document structure and applying matching paragraph style.')
        return ' '.join(reasons[:2])

    def build_prompt(
        self,
        user_prompt: str,
        doc_context: WordContext,
        mem_context: str,
        classification=None,
    ) -> str:
        """Build a fully-assembled Word domain prompt with zero unreplaced variables."""
        self._last_reasoning = self._generate_reasoning(doc_context, mem_context)
        from sidecar.masters.word_prompt_builder import build_word_prompt
        prompt = build_word_prompt(
            user_instruction=user_prompt,
            context=doc_context,
            mem_context=mem_context,
            file_path=getattr(doc_context, "file_path", None),
            app_name="Microsoft Word",
            app_type="Word Processor",
            intent_classification=getattr(classification, "task_type", "insert") if classification else "insert",
        )
        prompt += '\nAlso output a "reasoning" field explaining: which paragraph style you chose and why, and your word count rationale based on the document context.\n'
        return prompt

    def validate_operations(self, raw_response, doc_context: WordContext) -> list:
        """Validate and fuzzy-match operations from the LLM response."""
        if hasattr(raw_response, 'reasoning'):
            raw_response.reasoning = getattr(self, '_last_reasoning', '') or getattr(raw_response, 'reasoning', '')
        validated = []
        ops = getattr(raw_response, "operations", [])
        if isinstance(ops, list):
            for op in ops:
                op_dict = op.model_dump() if hasattr(op, "model_dump") else dict(op)
                result = self._validator.validate(op_dict, doc_context)
                if result.valid:
                    validated.append(result.op)
                else:
                    log.warning(f"WordMaster rejected op: {result.error}")
        return validated

    def apply_operations(self, file_path: str, operations: list, context: WordContext = None) -> dict:
        """Write validated operations to the .docx file atomically."""
        if context is None:
            # Create a minimal context so the writer can function
            try:
                context = self._extractor.extract(file_path, 0)
            except Exception:
                context = WordContext(
                    styles={"paragraph": [], "character": [], "table": []},
                    paragraphs=[],
                    tables=[],
                    theme_fonts={"major": "Calibri", "minor": "Calibri"},
                    list_sequences=[],
                    document_purpose="general",
                    cursor_paragraph_index=0,
                    total_paragraphs=0,
                )
        return self._writer.apply_operations(file_path, operations, context)

    def get_schema_class(self):
        """Return the Pydantic schema class for LLM structured output."""
        from sidecar.schemas.docx_schema import DocxResponse
        return DocxResponse
