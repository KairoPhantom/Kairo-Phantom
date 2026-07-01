import os
import shutil
import logging
import traceback
from pathlib import Path
from typing import List, Dict, Any

from sidecar.constants import KAIRO_BACKUP_SUFFIX

try:
    import win32com.client  # noqa: F401
    import pythoncom  # noqa: F401

    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

from docx import Document

log = logging.getLogger("kairo-sidecar.docx_writer")


def _find_prompt_paragraphs(doc, com_mode=False) -> list:
    """Find indices of all paragraphs that start with '//' (user prompt lines)."""
    prompt_indices = []
    if com_mode:
        for i in range(1, doc.Paragraphs.Count + 1):
            text = doc.Paragraphs(i).Range.Text.strip()
            if text.startswith("//"):
                prompt_indices.append(i - 1)  # 0-based
    else:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("//"):
                prompt_indices.append(i)
    return prompt_indices


def write_docx(file_path: str, operations: List[Dict[str, Any]]) -> dict:
    """
    Applies a list of DocxOperations to a .docx file.
    """
    log.info(f"Incoming docx operations: {operations}")
    """
    Strategy:
    1. Try COM write first (Word has the file open/locked in almost all test scenarios)
    2. If COM unavailable, fall back to python-docx file write
    3. Automatically remove any // prompt paragraphs before injecting new content
    """
    path = Path(file_path)
    if not path.exists():
        return {"error": f"Document not found: {path}"}

    # STRATEGY: Try COM first (Word is almost always open during testing)
    if HAS_WIN32COM:
        com_result = _try_com_write(str(path), operations)
        if "error" not in com_result:
            return com_result
        log.warning(
            f"COM write attempt failed: {com_result.get('error')} — falling back to file write"
        )

    # Fallback: python-docx file write
    backup_path = path.with_suffix(path.suffix + KAIRO_BACKUP_SUFFIX)
    try:
        shutil.copy2(path, backup_path)
    except Exception as e:
        return {"error": f"Failed to create backup of document: {e}"}

    applied = []
    errors = []
    tmp_path = path.with_suffix(path.suffix + ".tmp")

    try:
        doc = Document(str(path))

        # Auto-remove // prompt paragraphs before injecting new content
        prompt_indices = _find_prompt_paragraphs(doc, com_mode=False)
        for idx in sorted(prompt_indices, reverse=True):
            try:
                para = doc.paragraphs[idx]
                para._element.getparent().remove(para._element)
                log.info(f"Auto-removed prompt paragraph at index {idx}")
            except Exception as e:
                log.warning(f"Could not remove prompt paragraph {idx}: {e}")

        # Pre-cache remaining paragraphs to prevent index shifting bugs
        original_paragraphs = list(doc.paragraphs)

        for op in operations:
            op_type = op.get("action", op.get("type", ""))

            # Normalize 'content' → 'runs'
            if "content" in op and "runs" not in op:
                content = op["content"]
                if isinstance(content, list):
                    op["runs"] = [{"text": str(c)} for c in content]
                else:
                    op["runs"] = [{"text": str(content)}]

            try:
                if op_type in ("insert_paragraph", "append", "insert_after_heading"):
                    _op_insert_paragraph(doc, op, original_paragraphs)
                    applied.append(op)
                elif op_type == "replace_paragraph":
                    _op_replace_paragraph(doc, op, original_paragraphs)
                    applied.append(op)
                elif op_type == "append_to_run":
                    _op_append_to_run(doc, op, original_paragraphs)
                    applied.append(op)
                elif op_type == "insert_table":
                    _op_insert_table(doc, op, original_paragraphs)
                    applied.append(op)
                elif op_type in ("delete_paragraph", "remove_prompt"):
                    _op_delete_paragraph(doc, op, original_paragraphs)
                    applied.append(op)
                else:
                    errors.append(f"Unknown operation type: {op_type!r}")
            except Exception as e:
                errors.append(f"Operation {op_type!r} failed: {e}")
                log.error(f"Operation error: {e}\n{traceback.format_exc()}")

        # Atomic save
        doc.save(str(tmp_path))
        os.replace(str(tmp_path), str(path))

        # Cleanup backup only when the save/replace was successful and there are no errors.
        if not errors:
            try:
                backup_path.unlink(missing_ok=True)
            except Exception:
                pass

        return {
            "applied_count": len(applied),
            "errors": errors,
            "path": str(path),
        }

    except PermissionError:
        try:
            if backup_path.exists():
                shutil.copy2(str(backup_path), str(path))
        except Exception:
            pass
        try:
            backup_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_path.unlink()
        except Exception:
            pass
        return {
            "error": "Word has this file locked. COM write failed too — close Word and retry.",
            "path": str(path),
        }
    except Exception as e:
        try:
            if backup_path.exists():
                shutil.copy2(str(backup_path), str(path))
        except Exception:
            pass
        try:
            backup_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_path.unlink()
        except Exception:
            pass
        log.error(f"Failed to write docx: {e}\n{traceback.format_exc()}")
        return {"error": f"Failed to write document: {e}", "traceback": traceback.format_exc()}


def _add_runs_to_paragraph(para, runs_data: List[Dict[str, Any]]):
    """Adds styled runs to an existing paragraph element."""
    for run_item in runs_data:
        text = run_item.get("text", "")
        if text:
            run = para.add_run(text)
            run.bold = run_item.get("bold", False)
            run.italic = run_item.get("italic", False)


def _make_paragraph(doc, style: str, runs_data: List[Dict[str, Any]]):
    """Creates a new paragraph with styled runs, not yet added to document body."""
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style]
    except Exception:
        log.warning(f"Style '{style}' not found in document, falling back to Normal")
        try:
            para.style = doc.styles["Normal"]
        except Exception:
            pass
    _add_runs_to_paragraph(para, runs_data)
    return para


def _op_insert_paragraph(doc, op: Dict[str, Any], original_paragraphs: List[Any]):
    """Inserts a styled paragraph after a specific index (relative to original document)."""
    after_idx = op.get("after_paragraph_index", op.get("index", -1))
    style = op.get("style", "Normal")
    runs_data = op.get("runs", [])

    if 0 <= after_idx < len(original_paragraphs):
        ref_para = original_paragraphs[after_idx]
        if ref_para._element is not None and ref_para._element.getparent() is not None:
            new_para = _make_paragraph(doc, style, runs_data)
            ref_para._element.addnext(new_para._element)
            return

    # Append to end of document
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style]
    except Exception:
        pass
    _add_runs_to_paragraph(para, runs_data)


def _op_replace_paragraph(doc, op: Dict[str, Any], original_paragraphs: List[Any]):
    """Replaces text and style of a paragraph at specific index (relative to original document)."""
    idx = op.get("paragraph_index", op.get("index", -1))
    style = op.get("style", "Normal")
    runs_data = op.get("runs", [])

    if 0 <= idx < len(original_paragraphs):
        para = original_paragraphs[idx]
        if para._element is not None and para._element.getparent() is not None:
            # Clear existing text safely
            for run in para.runs:
                run.text = ""
            # Apply style if found
            try:
                para.style = doc.styles[style]
            except Exception:
                pass
            _add_runs_to_paragraph(para, runs_data)


def _op_append_to_run(doc, op: Dict[str, Any], original_paragraphs: List[Any]):
    """Appends styled runs to an existing paragraph at specific index."""
    idx = op.get("paragraph_index", op.get("index", -1))
    runs_data = op.get("runs", [])
    if 0 <= idx < len(original_paragraphs):
        para = original_paragraphs[idx]
        _add_runs_to_paragraph(para, runs_data)


def _op_insert_table(doc, op: Dict[str, Any], original_paragraphs: List[Any]):
    """Inserts a table after a specific paragraph index (relative to original document)."""
    after_idx = op.get("after_paragraph_index", op.get("index", -1))
    headers = op.get("headers", [])
    rows_data = op.get("rows", [])

    # Let's count columns
    cols = len(headers) if headers else (max(len(r) for r in rows_data) if rows_data else 0)
    if cols == 0:
        return

    # Create table in document first (python-docx appends to end of body element)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"

    # Add header
    if headers:
        row = table.add_row()
        for i, h in enumerate(headers):
            row.cells[i].text = str(h)

    # Add rows
    for r in rows_data:
        row = table.add_row()
        for i, cell_text in enumerate(r):
            if i < cols:
                row.cells[i].text = str(cell_text)

    # Move table immediately after after_idx paragraph
    if 0 <= after_idx < len(original_paragraphs):
        ref_para = original_paragraphs[after_idx]
        if ref_para._element is not None and ref_para._element.getparent() is not None:
            ref_para._element.addnext(table._element)


def _op_delete_paragraph(doc, op: Dict[str, Any], original_paragraphs: List[Any]):
    """Deletes a paragraph at specific index (relative to original document)."""
    idx = op.get("paragraph_index", op.get("index", -1))
    if 0 <= idx < len(original_paragraphs):
        para = original_paragraphs[idx]
        if para._element is not None and para._element.getparent() is not None:
            para._element.getparent().remove(para._element)


def _get_com_style(doc, style_name: str):
    mapped = style_name
    if style_name.startswith("Heading") and len(style_name) > 7 and style_name[7].isdigit():
        mapped = f"Heading {style_name[7:]}"
    elif style_name == "ListBullet":
        mapped = "List Bullet"
    elif style_name == "ListNumber":
        mapped = "List Number"

    try:
        return doc.Styles(mapped)
    except Exception:
        try:
            return doc.Styles("Normal")
        except Exception:
            return None


def _try_com_write(file_path: str, operations: List[Dict[str, Any]]) -> dict:
    """Primary write path: use Word COM API to write directly into the open document."""
    import win32com.client
    import pythoncom
    import subprocess

    pythoncom.CoInitialize()

    word = None
    target_doc = None
    path = Path(file_path).resolve()

    word_running = False
    try:
        out = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq winword.exe"], capture_output=True, text=True
        )
        word_running = "winword.exe" in out.stdout.lower()
    except Exception:
        pass

    if word_running:
        try:
            target_doc = win32com.client.GetObject(str(path))
            word = target_doc.Application
        except Exception:
            try:
                word = win32com.client.GetActiveObject("Word.Application")
                for doc in word.Documents:
                    try:
                        if Path(doc.FullName).resolve() == path:
                            target_doc = doc
                            break
                    except Exception:
                        pass
            except Exception:
                pass

    if not target_doc:
        return {"error": f"Document {path.name} is not currently open in Word."}

    # Auto-remove // prompt paragraphs via COM
    prompt_indices = _find_prompt_paragraphs(target_doc, com_mode=True)
    for idx in sorted(prompt_indices, reverse=True):
        try:
            p = target_doc.Paragraphs(idx + 1)
            p.Range.Delete()
            log.info(f"COM: Auto-removed prompt paragraph at index {idx}")
        except Exception as e:
            log.warning(f"COM: Could not remove prompt paragraph {idx}: {e}")

    total_paras = target_doc.Paragraphs.Count
    current_indices = list(range(total_paras))

    applied = []
    errors = []

    for op in operations:
        op_type = op.get("action", op.get("type", ""))

        if "content" in op and "runs" not in op:
            content = op["content"]
            if isinstance(content, list):
                op["runs"] = [{"text": str(c)} for c in content]
            else:
                op["runs"] = [{"text": str(content)}]

        try:
            if op_type in ("insert_paragraph", "append", "insert_after_heading"):
                _com_insert_paragraph(target_doc, op, current_indices)
                applied.append(op)
            elif op_type == "replace_paragraph":
                _com_replace_paragraph(target_doc, op, current_indices)
                applied.append(op)
            elif op_type == "append_to_run":
                _com_append_to_run(target_doc, op, current_indices)
                applied.append(op)
            elif op_type == "insert_table":
                _com_insert_table_com(target_doc, op, current_indices)
                applied.append(op)
            elif op_type in ("delete_paragraph", "remove_prompt"):
                _com_delete_paragraph(target_doc, op, current_indices)
                applied.append(op)
            else:
                errors.append(f"Unknown operation type: {op_type!r}")
        except Exception as e:
            errors.append(f"Operation {op_type!r} failed: {e}")
            log.error(f"COM Operation error: {e}\n{traceback.format_exc()}")

    # Save the document via COM
    try:
        target_doc.Save()
    except Exception as e:
        log.warning(f"COM Save failed (may be auto-saved): {e}")

    return {
        "applied_count": len(applied),
        "errors": errors,
        "path": str(path),
    }


def _com_set_paragraph_content(doc, p, text, style, runs_data):
    """Sets the text and style of a paragraph without overwriting its paragraph mark."""
    start = p.Range.Start
    end = max(start, p.Range.End - 1)
    rng = doc.Range(Start=start, End=end)
    rng.Text = text
    if runs_data:
        _com_apply_runs(rng, runs_data)
    if style:
        com_style = _get_com_style(doc, style)
        if com_style:
            p.Style = com_style


def _com_insert_paragraph(doc, op, current_indices):
    op_type = op.get("action", op.get("type", ""))
    runs_data = op.get("runs", [])
    text = "".join(r.get("text", "") for r in runs_data)
    style = op.get("style", "Normal")

    curr_idx = -1
    if op_type == "insert_after_heading":
        heading_text = op.get("heading_text", "")
        for i in range(1, doc.Paragraphs.Count + 1):
            p_text = doc.Paragraphs(i).Range.Text.strip()
            if heading_text.lower() in p_text.lower():
                curr_idx = i - 1
                break
    else:
        idx = op.get("after_paragraph_index", op.get("index", -1))
        if 0 <= idx < len(current_indices):
            curr_idx = current_indices[idx]

    count_before = doc.Paragraphs.Count

    if curr_idx == -1:
        # Append to end of document
        p = doc.Paragraphs(doc.Paragraphs.Count)
        if p.Range.Text.strip() == "":
            new_p = p
        else:
            rng = p.Range
            rng.Collapse(0)  # Collapse to end
            rng.InsertParagraphAfter()
            new_p = doc.Paragraphs(doc.Paragraphs.Count)
    else:
        p = doc.Paragraphs(curr_idx + 1)
        next_p = p.Next()
        if next_p:
            doc.Paragraphs.Add(Range=next_p.Range)
            new_p = doc.Paragraphs(curr_idx + 2)
        else:
            rng = p.Range
            rng.Collapse(0)
            rng.InsertParagraphAfter()
            new_p = doc.Paragraphs(doc.Paragraphs.Count)

    _com_set_paragraph_content(doc, new_p, text, style, runs_data)

    count_after = doc.Paragraphs.Count
    diff = count_after - count_before
    if diff != 0 and curr_idx != -1:
        for i in range(len(current_indices)):
            if current_indices[i] > curr_idx:
                current_indices[i] += diff


def _com_replace_paragraph(doc, op, current_indices):
    idx = op.get("paragraph_index", op.get("index", -1))
    runs_data = op.get("runs", [])
    text = "".join(r.get("text", "") for r in runs_data)
    style = op.get("style", "")

    if 0 <= idx < len(current_indices):
        curr_idx = current_indices[idx]
        if curr_idx != -1:
            p = doc.Paragraphs(curr_idx + 1)
            _com_set_paragraph_content(doc, p, text, style, runs_data)


def _com_delete_paragraph(doc, op, current_indices):
    idx = op.get("paragraph_index", op.get("index", -1))
    if 0 <= idx < len(current_indices):
        curr_idx = current_indices[idx]
        if curr_idx != -1:
            count_before = doc.Paragraphs.Count
            p = doc.Paragraphs(curr_idx + 1)
            p.Range.Delete()
            count_after = doc.Paragraphs.Count
            diff = count_after - count_before

            # Mark deleted paragraph
            current_indices[idx] = -1

            # Shift subsequent paragraphs
            if diff != 0:
                for i in range(len(current_indices)):
                    if current_indices[i] > curr_idx:
                        current_indices[i] += diff


def _com_append_to_run(doc, op, current_indices):
    idx = op.get("paragraph_index", op.get("index", -1))
    runs_data = op.get("runs", [])
    text = "".join(r.get("text", "") for r in runs_data)
    if 0 <= idx < len(current_indices):
        curr_idx = current_indices[idx]
        if curr_idx != -1:
            p = doc.Paragraphs(curr_idx + 1)
            # Find the end of the range, but keep paragraph mark
            start = p.Range.End - 1
            rng = doc.Range(Start=start, End=start)
            rng.Text = text
            if runs_data:
                _com_apply_runs(rng, runs_data)


def _com_insert_table_com(doc, op, current_indices):
    after_idx = op.get("after_paragraph_index", op.get("index", -1))
    headers = op.get("headers", [])
    rows_data = op.get("rows", [])

    cols = len(headers) if headers else (max(len(r) for r in rows_data) if rows_data else 0)
    rows = len(rows_data) + (1 if headers else 0)
    if cols == 0 or rows == 0:
        return

    curr_idx = -1
    if 0 <= after_idx < len(current_indices):
        curr_idx = current_indices[after_idx]

    count_before = doc.Paragraphs.Count

    if curr_idx == -1:
        rng = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
    else:
        p = doc.Paragraphs(curr_idx + 1)
        rng = doc.Range(p.Range.End, p.Range.End)

    table = doc.Tables.Add(rng, rows, cols)
    try:
        table.Style = "Table Grid"
    except Exception:
        pass

    r_idx = 1
    if headers:
        for c_idx, h in enumerate(headers):
            table.Cell(r_idx, c_idx + 1).Range.Text = str(h)
        r_idx += 1

    for r in rows_data:
        for c_idx, cell_text in enumerate(r):
            if c_idx < cols:
                table.Cell(r_idx, c_idx + 1).Range.Text = str(cell_text)
        r_idx += 1

    count_after = doc.Paragraphs.Count
    diff = count_after - count_before
    if diff != 0 and curr_idx != -1:
        for i in range(len(current_indices)):
            if current_indices[i] > curr_idx:
                current_indices[i] += diff


def _com_apply_runs(rng, runs_data):
    start = rng.Start
    for r in runs_data:
        text = r.get("text", "")
        length = len(text)
        if length == 0:
            continue
        run_rng = rng.Document.Range(start, start + length)
        if r.get("bold"):
            run_rng.Bold = True
        if r.get("italic"):
            run_rng.Italic = True
        start += length
