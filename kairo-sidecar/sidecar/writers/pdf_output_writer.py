import logging
from pathlib import Path

log = logging.getLogger("kairo-sidecar.pdf_output_writer")


def write_pdf_output(content: dict, output_type: str, source_path: str) -> str:
    """
    Converts parsed PDF content (paragraphs, tables, metadata) into:
    - Word (.docx) using python-docx
    - PDF (.pdf) using weasyprint or reportlab fallback
    Saves to {source_path_dir}/{source_filename}_kairo_output.{ext}
    Returns the absolute path of the written file.
    """
    source_p = Path(source_path).resolve()
    dest_ext = f".{output_type.lower()}"
    dest_filename = f"{source_p.stem}_kairo_output{dest_ext}"
    dest_path = source_p.parent / dest_filename

    log.info(f"Writing PDF extraction output to {dest_path} in format {output_type}")

    if output_type.lower() == "docx":
        _write_to_docx(content, dest_path)
    elif output_type.lower() == "pdf":
        _write_to_pdf(content, dest_path)
    else:
        raise ValueError(f"Unsupported output type: {output_type}")

    return str(dest_path)


def _write_to_docx(content: dict, dest_path: Path):
    from docx import Document

    doc = Document()

    # Title
    title = content.get("metadata", {}).get("file_path", "PDF Extraction Output")
    doc.add_heading(f"Extracted Content: {Path(title).name}", level=1)

    # Paragraphs & Tables interspersed by index or logically ordered
    paragraphs = content.get("paragraphs", [])
    tables = content.get("tables", [])

    # Sort paragraphs by index
    sorted_paras = sorted(paragraphs, key=lambda x: x.get("index", 0))

    # Keep track of which tables belong after which paragraph
    table_map = {}
    for table in tables:
        after_idx = table.get("after_paragraph_index", -1)
        if after_idx not in table_map:
            table_map[after_idx] = []
        table_map[after_idx].append(table)

    # Write paragraph index -1 tables (before any paragraph)
    if -1 in table_map:
        for table in table_map[-1]:
            _add_docx_table(doc, table)

    for p in sorted_paras:
        text = p.get("text", "")
        style = p.get("style", "Normal")
        level = p.get("level", 0)

        if style.startswith("Heading"):
            doc.add_heading(text, level=level if level > 0 else 1)
        else:
            p_obj = doc.add_paragraph()
            # Restore runs
            runs = p.get("runs", [])
            if runs:
                for run_data in runs:
                    r = p_obj.add_run(run_data.get("text", ""))
                    r.bold = run_data.get("bold", False)
                    r.italic = run_data.get("italic", False)
            else:
                p_obj.add_run(text)

        # Insert any tables that should come after this paragraph
        idx = p.get("index", 0)
        if idx in table_map:
            for table in table_map[idx]:
                _add_docx_table(doc, table)

    doc.save(str(dest_path))
    log.info(f"Saved DOCX output file successfully to {dest_path}")


def _add_docx_table(doc, table_data: dict):
    rows_data = table_data.get("rows", [])
    if not rows_data:
        return
    cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    for r in rows_data:
        row = table.add_row()
        for col_idx, cell_data in enumerate(r):
            # cell_data could be a string or a dict/object depending on MinerU/Docling
            if isinstance(cell_data, dict):
                cell_text = cell_data.get("text", "")
            else:
                cell_text = str(cell_data)
            if col_idx < cols:
                row.cells[col_idx].text = cell_text


def _write_to_pdf(content: dict, dest_path: Path):
    # Create HTML
    html_content = _build_html_string(content)

    # Try WeasyPrint
    try:
        import weasyprint

        weasyprint.HTML(string=html_content).write_pdf(str(dest_path))
        log.info(f"Saved PDF output file successfully via WeasyPrint to {dest_path}")
        return
    except ImportError:
        log.warning("WeasyPrint is not installed. Trying ReportLab fallback...")

    # Try ReportLab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # noqa: F401
        from reportlab.lib import colors

        doc = SimpleDocTemplate(str(dest_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        title = content.get("metadata", {}).get("file_path", "PDF Extraction Output")
        story.append(Paragraph(f"Extracted Content: {Path(title).name}", styles["Heading1"]))
        story.append(Spacer(1, 12))

        paragraphs = content.get("paragraphs", [])
        tables = content.get("tables", [])
        sorted_paras = sorted(paragraphs, key=lambda x: x.get("index", 0))

        table_map = {}
        for table in tables:
            after_idx = table.get("after_paragraph_index", -1)
            if after_idx not in table_map:
                table_map[after_idx] = []
            table_map[after_idx].append(table)

        for p in sorted_paras:
            text = p.get("text", "")
            style = p.get("style", "Normal")

            p_style = styles["Normal"]
            if style.startswith("Heading"):
                p_style = styles[style] if style in styles else styles["Heading2"]

            story.append(Paragraph(text, p_style))
            story.append(Spacer(1, 6))

            idx = p.get("index", 0)
            if idx in table_map:
                for table in table_map[idx]:
                    rows_data = table.get("rows", [])
                    if rows_data:
                        # Convert all cell values to strings/paragraphs
                        cleaned_rows = []
                        for row in rows_data:
                            cleaned_row = []
                            for cell in row:
                                if isinstance(cell, dict):
                                    cell_val = cell.get("text", "")
                                else:
                                    cell_val = str(cell)
                                cleaned_row.append(Paragraph(cell_val, styles["Normal"]))
                            cleaned_rows.append(cleaned_row)

                        col_widths = [doc.width / len(cleaned_rows[0])] * len(cleaned_rows[0])
                        t = Table(cleaned_rows, colWidths=col_widths)
                        t.setStyle(
                            TableStyle(
                                [
                                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                                    ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                                    ("PADDING", (0, 0), (-1, -1), 6),
                                ]
                            )
                        )
                        story.append(t)
                        story.append(Spacer(1, 12))

        doc.build(story)
        log.info(f"Saved PDF output file successfully via ReportLab to {dest_path}")
        return
    except ImportError:
        log.warning("ReportLab is not installed.")

    # As a last resort, write a beautiful HTML file with system print stylesheets
    html_dest = dest_path.with_suffix(".html")
    with open(html_dest, "w", encoding="utf-8") as f:
        f.write(html_content)
    log.error(
        f"Neither WeasyPrint nor ReportLab was available. Saved beautiful HTML report to {html_dest} instead."
    )
    raise RuntimeError(
        "PDF generation failed: Neither WeasyPrint nor ReportLab was available. Saved HTML version."
    )


def _build_html_string(content: dict) -> str:
    title = content.get("metadata", {}).get("file_path", "PDF Extraction Output")
    filename = Path(title).name

    paragraphs = content.get("paragraphs", [])
    tables = content.get("tables", [])
    sorted_paras = sorted(paragraphs, key=lambda x: x.get("index", 0))

    table_map = {}
    for table in tables:
        after_idx = table.get("after_paragraph_index", -1)
        if after_idx not in table_map:
            table_map[after_idx] = []
        table_map[after_idx].append(table)

    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<meta charset='utf-8'>",
        f"<title>Extracted: {filename}</title>",
        "<style>",
        "body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 40px auto; padding: 0 20px; }",
        "h1 { color: #111; border-bottom: 2px solid #eaecef; padding-bottom: 10px; margin-top: 40px; }",
        "h2, h3, h4 { color: #222; margin-top: 30px; }",
        "p { margin: 1em 0; }",
        "table { width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 0.9em; }",
        "th, td { border: 1px solid #dfe2e5; padding: 8px 12px; text-align: left; }",
        "th { background-color: #f6f8fa; font-weight: bold; }",
        "tr:nth-child(even) { background-color: #f8f9fa; }",
        "footer { margin-top: 60px; font-size: 0.8em; color: #6a737d; text-align: center; border-top: 1px solid #eaecef; padding-top: 20px; }",
        "</style>",
        "</head>",
        "body>",
        f"<h1>Extracted Content: {filename}</h1>",
        f"<p style='color: #6a737d; font-size: 0.9em;'>Parsed on {content.get('metadata', {}).get('parse_timestamp', 'unknown')} using {content.get('metadata', {}).get('tier', 'unknown')} tier</p>",
    ]

    # Interspersed elements
    if -1 in table_map:
        for table in table_map[-1]:
            html_parts.append(_build_html_table(table))

    for p in sorted_paras:
        text = p.get("text", "")
        style = p.get("style", "Normal")
        level = p.get("level", 0)

        if style.startswith("Heading"):
            h_tag = f"h{level}" if 1 <= level <= 6 else "h2"
            html_parts.append(f"<{h_tag}>{text}</{h_tag}>")
        else:
            html_parts.append(f"<p>{text}</p>")

        idx = p.get("index", 0)
        if idx in table_map:
            for table in table_map[idx]:
                html_parts.append(_build_html_table(table))

    html_parts.append("<footer>Generated by Kairo Phantom Document Engine</footer>")
    html_parts.append("</body>")
    html_parts.append("</html>")

    return "\n".join(html_parts)


def _build_html_table(table_data: dict) -> str:
    rows = table_data.get("rows", [])
    if not rows:
        return ""

    parts = ["<table>"]
    for i, row in enumerate(rows):
        parts.append("<tr>")
        for cell in row:
            if isinstance(cell, dict):
                cell_text = cell.get("text", "")
            else:
                cell_text = str(cell)
            tag = "th" if i == 0 else "td"
            parts.append(f"<{tag}>{cell_text}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "\n".join(parts)
