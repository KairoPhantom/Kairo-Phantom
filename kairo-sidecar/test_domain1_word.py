"""
Tests for Domain 1: Word / DOCX Native Track Changes.
All tests are offline (no Adeu or safe-docx installation required).
Gate tests clearly marked: Gate 1, Gate 2, Gate 3, Gate 4.
"""

import json
import os
import shutil
import sys
import tempfile
import textwrap
from pathlib import Path

import pytest

# Add sidecar package to path
sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from sidecar.parsers.adeu_bridge import (
    _adeu_installed,
    _adeu_sdk_available,
    _parse_markdown_to_paragraphs,
    _error as adeu_error,
    _unavailable as adeu_unavailable,
)
from sidecar.parsers.safedocx_bridge import (
    _safedocx_installed,
    _error as safedocx_error,
    _unavailable as safedocx_unavailable,
)
from sidecar.parsers.legal_redline import (
    detect_cuad_clauses,
    generate_redlines_for_clause,
    generate_contract_summary,
    analyze_contract,
    CUAD_HIGH_RISK_CLAUSES,
    _apply_rule_based_redline,
    _estimate_risk_reduction,
)


# ─── Fixtures ────────────────────────────────────────────────────────────────

SAMPLE_CONTRACT = textwrap.dedent("""
SERVICES AGREEMENT

This Services Agreement ("Agreement") is entered into as of January 1, 2025,
between Acme Corp ("Company") and Supplier Inc ("Supplier").

1. TERM AND TERMINATION
The Company may terminate this Agreement at any time for convenience with
immediate effect. The Supplier shall have no claim for damages.

2. INTELLECTUAL PROPERTY
All work product, inventions, and deliverables created under this Agreement
shall be considered work for hire. All right title and interest in such
materials shall be solely owned by the Company.

3. LIMITATION OF LIABILITY
The aggregate liability of Supplier shall not exceed the fees paid in the
preceding three (3) months. This cap applies to all claims.

4. AUTO-RENEWAL
This Agreement shall automatically renew for successive one-year periods
unless either party provides 15 days prior written notice of non-renewal.

5. NON-COMPETE
During the term and for two years thereafter, Supplier agrees not to engage
in any competitive activities in any jurisdiction.

6. CONFIDENTIALITY
Each party agrees to maintain the confidentiality of the other party's
proprietary information and non-disclosure obligations.

7. INDEMNIFICATION
Supplier shall indemnify, defend, and hold harmless the Company from any
and all claims arising from Supplier's performance.

8. GOVERNING LAW
This Agreement shall be governed by the laws of Delaware, and jurisdiction
shall be in the courts of Delaware.
""").strip()

SAMPLE_MARKDOWN = textwrap.dedent("""
# Introduction

This is the first section.

## Background

Some background context here.

### Details

More detailed information.

A regular paragraph with no heading.

| Column A | Column B |
|----------|----------|
| Value 1  | Value 2  |
""").strip()


# ─── Gate 3: Legal Contract Review ──────────────────────────────────────────


class TestCuadClauseDetection:
    """Gate 3: CUAD detection on sample contract."""

    def test_detects_termination_for_convenience(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert (
            "termination_for_convenience" in ids
        ), f"Expected termination_for_convenience, got: {ids}"

    def test_detects_ip_ownership(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "ip_ownership" in ids, f"Expected ip_ownership, got: {ids}"

    def test_detects_liability_cap(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "liability_cap" in ids, f"Expected liability_cap, got: {ids}"

    def test_detects_auto_renewal(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "auto_renewal" in ids, f"Expected auto_renewal, got: {ids}"

    def test_detects_non_compete(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "non_compete" in ids, f"Expected non_compete, got: {ids}"

    def test_detects_confidentiality(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "confidentiality" in ids, f"Expected confidentiality, got: {ids}"

    def test_detects_governing_law(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "governing_law" in ids, f"Expected governing_law, got: {ids}"

    def test_detects_indemnification(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        ids = [c["id"] for c in result["data"]["detected_clauses"]]
        assert "indemnification" in ids, f"Expected indemnification, got: {ids}"

    def test_risk_summary_has_correct_structure(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        summary = result["data"]["risk_summary"]
        assert "HIGH" in summary
        assert "MEDIUM" in summary
        assert "total_flagged" in summary
        assert (
            summary["total_flagged"] >= 4
        ), f"Expected ≥4 flagged clauses, got {summary['total_flagged']}"

    def test_high_risk_count_correct(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        summary = result["data"]["risk_summary"]
        # termination, ip_ownership, liability_cap, non_compete, indemnification → ≥3 HIGH
        assert summary["HIGH"] >= 3, f"Expected ≥3 HIGH risk, got {summary['HIGH']}"

    def test_matched_text_is_excerpt(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        clauses = result["data"]["detected_clauses"]
        for c in clauses:
            assert len(c["matched_text"]) > 0, f"Clause {c['id']} has empty matched_text"

    def test_confidence_in_valid_range(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        for c in result["data"]["detected_clauses"]:
            assert (
                0.0 <= c["confidence"] <= 1.0
            ), f"Clause {c['id']} confidence {c['confidence']} out of [0,1]"

    def test_empty_text_returns_error(self):
        result = detect_cuad_clauses("")
        assert result["ok"] is False
        assert "empty" in result["error"].lower()

    def test_missing_standard_clauses_is_list(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert isinstance(result["data"]["missing_standard_clauses"], list)

    def test_clauses_sorted_high_first(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        clauses = result["data"]["detected_clauses"]
        risk_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
        levels = [risk_order[c["risk_level"]] for c in clauses]
        assert levels == sorted(levels), "Clauses should be sorted HIGH → MEDIUM → LOW"

    def test_no_false_positives_on_lorem_ipsum(self):
        lorem = "Lorem ipsum dolor sit amet consectetur adipiscing elit sed do eiusmod tempor."
        result = detect_cuad_clauses(lorem)
        assert result["ok"] is True
        # No clause should fire on lorem ipsum
        assert result["data"]["risk_summary"]["total_flagged"] == 0


class TestRedlineGeneration:
    """Gate 3: Redline generation correctness."""

    def test_liability_cap_redline_contains_carveout(self):
        clause_text = "The aggregate liability shall not exceed fees paid."
        result = generate_redlines_for_clause(clause_text, "liability_cap", "balanced", "client")
        assert result["ok"] is True
        suggested = result["data"]["suggested_text"]
        # Should add carveout language
        assert "gross negligence" in suggested.lower() or "excluding" in suggested.lower()

    def test_auto_renewal_redline_increases_notice(self):
        clause_text = "Agreement shall renew unless 15 days prior notice is given."
        result = generate_redlines_for_clause(clause_text, "auto_renewal", "balanced", "client")
        assert result["ok"] is True
        suggested = result["data"]["suggested_text"]
        assert "60" in suggested or "sixty" in suggested.lower()

    def test_redline_returns_original_text(self):
        clause_text = "Terminate at any time."
        result = generate_redlines_for_clause(
            clause_text, "termination_for_convenience", "balanced", "client"
        )
        assert result["ok"] is True
        assert result["data"]["original_text"] == clause_text

    def test_redline_has_rationale(self):
        result = generate_redlines_for_clause("standard clause text", "liability_cap", "balanced")
        assert result["ok"] is True
        assert len(result["data"]["rationale"]) > 10

    def test_risk_reduction_present(self):
        result = generate_redlines_for_clause("Some clause text", "liability_cap", "aggressive")
        assert result["ok"] is True
        assert "→" in result["data"]["risk_reduction"]

    def test_empty_clause_text_returns_error(self):
        result = generate_redlines_for_clause("", "liability_cap")
        assert result["ok"] is False

    def test_aggressive_stance_gives_more_reduction(self):
        result_agg = generate_redlines_for_clause("cap on liability", "liability_cap", "aggressive")
        result_con = generate_redlines_for_clause(
            "cap on liability", "liability_cap", "conservative"
        )
        assert result_agg["ok"] and result_con["ok"]
        # Aggressive should reduce more (HIGH → LOW vs HIGH → MEDIUM)
        assert "LOW" in result_agg["data"]["risk_reduction"]


class TestContractSummary:
    """Gate 3: Executive summary generation."""

    def test_summary_text_is_non_empty(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert len(result["data"]["summary_text"]) > 100

    def test_summary_contains_kairo_header(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert "KAIRO" in result["data"]["summary_text"].upper()

    def test_summary_has_high_risk_section(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert "HIGH RISK" in result["data"]["summary_text"].upper()

    def test_action_items_for_high_risk(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert len(result["data"]["action_items"]) >= 1

    def test_action_items_have_priority_field(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        for item in result["data"]["action_items"]:
            assert "priority" in item
            assert item["priority"] in ("URGENT", "HIGH", "MEDIUM", "LOW")


class TestFullContractAnalysis:
    """Gate 3: Full pipeline integration (analyze_contract)."""

    def test_full_analysis_returns_ok(self):
        result = analyze_contract(SAMPLE_CONTRACT)
        assert result["ok"] is True

    def test_full_analysis_detected_clauses(self):
        result = analyze_contract(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert result["data"]["total_clauses_detected"] >= 5

    def test_full_analysis_has_redlines(self):
        result = analyze_contract(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert len(result["data"]["suggested_redlines"]) >= 1

    def test_full_analysis_has_summary_text(self):
        result = analyze_contract(SAMPLE_CONTRACT)
        assert result["ok"] is True
        assert len(result["data"]["summary_text"]) > 50

    def test_full_analysis_empty_text_errors(self):
        result = analyze_contract("")
        assert result["ok"] is False


# ─── Gate 1 & 2: Adeu Bridge Unit Tests ──────────────────────────────────────


class TestAdeuBridgeHelpers:
    """Gate 1 & 2: Adeu bridge internal helpers."""

    def test_parse_markdown_paragraphs_basic(self):
        paragraphs = _parse_markdown_to_paragraphs(SAMPLE_MARKDOWN)
        assert len(paragraphs) >= 4

    def test_parse_markdown_heading_detection(self):
        paragraphs = _parse_markdown_to_paragraphs(SAMPLE_MARKDOWN)
        h1 = [p for p in paragraphs if p["level"] == 1]
        h2 = [p for p in paragraphs if p["level"] == 2]
        assert len(h1) >= 1, "Expected at least 1 H1"
        assert len(h2) >= 1, "Expected at least 1 H2"

    def test_parse_markdown_table_detection(self):
        paragraphs = _parse_markdown_to_paragraphs(SAMPLE_MARKDOWN)
        tables = [p for p in paragraphs if p["is_table"]]
        assert len(tables) >= 1, "Expected at least 1 table block"

    def test_parse_markdown_non_heading_paragraph(self):
        paragraphs = _parse_markdown_to_paragraphs(SAMPLE_MARKDOWN)
        non_heading = [p for p in paragraphs if p["level"] == 0 and not p["is_table"]]
        assert len(non_heading) >= 1

    def test_parse_markdown_index_sequential(self):
        paragraphs = _parse_markdown_to_paragraphs(SAMPLE_MARKDOWN)
        # All paragraphs should be non-empty
        for p in paragraphs:
            assert len(p["text"]) > 0

    def test_error_helper_structure(self):
        e = adeu_error("test error message")
        assert e["ok"] is False
        assert "error" in e
        assert e["error"] == "test error message"

    def test_unavailable_helper_structure(self):
        u = adeu_unavailable("adeu not installed")
        assert u["ok"] is False
        assert u.get("unavailable") is True

    def test_adeu_installed_returns_bool(self):
        result = _adeu_installed()
        assert isinstance(result, bool)

    def test_adeu_sdk_available_returns_bool(self):
        result = _adeu_sdk_available()
        assert isinstance(result, bool)

    def test_adeu_apply_edits_fallback_docx(self):
        from sidecar.parsers.adeu_bridge import adeu_apply_edits
        from unittest.mock import patch
        from docx import Document

        # Create a temp docx file
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, "test.docx")
        doc = Document()
        doc.add_paragraph("This is a sample agreement.")
        doc.save(docx_path)

        edits = [{"target_text": "sample", "new_text": "drafted"}]
        out_path = os.path.join(temp_dir, "test_redlined.docx")

        with (
            patch("sidecar.parsers.adeu_bridge._adeu_installed", return_value=False),
            patch("sidecar.parsers.adeu_bridge._adeu_sdk_available", return_value=False),
            patch("sidecar.parsers.adeu_bridge._word_is_open_with_file", return_value=False),
        ):
            res = adeu_apply_edits(docx_path, edits, output_path=out_path)

        assert res["ok"] is True
        assert res["data"]["backend"] == "python_docx_fallback"
        assert res["data"]["applied_count"] == 1

        # Verify document was modified and exists
        assert os.path.exists(out_path)
        doc2 = Document(out_path)
        # Check that revision tracking was turned on via XML
        from docx.oxml.ns import qn

        assert doc2.settings.element.find(qn("w:trackRevisions")) is not None

        # Clean up
        shutil.rmtree(temp_dir)


# ─── Gate 2: safe-docx Bridge Unit Tests ─────────────────────────────────────


class TestSafeDocxBridgeHelpers:
    """Gate 2: safe-docx bridge internal helpers."""

    def test_error_helper_structure(self):
        e = safedocx_error("test error message")
        assert e["ok"] is False
        assert e["error"] == "test error message"

    def test_unavailable_helper_structure(self):
        u = safedocx_unavailable("npx not found")
        assert u["ok"] is False
        assert u.get("unavailable") is True

    def test_safedocx_installed_returns_bool(self):
        result = _safedocx_installed()
        assert isinstance(result, bool)


# ─── Gate 4: W3 Regression — Zero Role/System Prompt Leakage ─────────────────


class TestW3Regression:
    """Gate 4: Ensure no system prompt / role information leaks into outputs."""

    def test_contract_analysis_no_system_prompt_leakage(self):
        result = analyze_contract(SAMPLE_CONTRACT)
        assert result["ok"] is True
        # Flatten all output text
        output_text = json.dumps(result["data"]).lower()
        forbidden_phrases = [
            "you are kairo",
            "your role is",
            "system prompt",
            "as an ai",
            "i am an ai",
            "as a language model",
        ]
        for phrase in forbidden_phrases:
            assert (
                phrase not in output_text
            ), f"Leaked forbidden phrase '{phrase}' in contract analysis output"

    def test_redline_output_no_role_leakage(self):
        result = generate_redlines_for_clause(
            "The aggregate liability shall not exceed fees paid.",
            "liability_cap",
        )
        assert result["ok"] is True
        output_text = json.dumps(result["data"]).lower()
        assert "you are" not in output_text
        assert "your role" not in output_text
        assert "system prompt" not in output_text

    def test_clause_detection_output_no_role_leakage(self):
        result = detect_cuad_clauses(SAMPLE_CONTRACT)
        assert result["ok"] is True
        output_text = json.dumps(result["data"]).lower()
        assert "you are" not in output_text
        assert "as an ai" not in output_text

    def test_summary_no_meta_information(self):
        result = generate_contract_summary(SAMPLE_CONTRACT)
        assert result["ok"] is True
        output_text = json.dumps(result["data"]).lower()
        forbidden = ["system_prompt", "ghost_session", "alt+m", "kairo phantom internal"]
        for phrase in forbidden:
            assert phrase not in output_text, f"Leaked internal info '{phrase}' in summary"


# ─── CUAD Catalogue Integrity ─────────────────────────────────────────────────


class TestCuadCatalogueIntegrity:
    """Verify the CUAD clause catalogue is structurally correct."""

    def test_all_clauses_have_required_fields(self):
        required = {"id", "label", "keywords", "risk_level", "description"}
        for clause in CUAD_HIGH_RISK_CLAUSES:
            missing = required - set(clause.keys())
            assert not missing, f"Clause {clause.get('id')} missing fields: {missing}"

    def test_all_ids_unique(self):
        ids = [c["id"] for c in CUAD_HIGH_RISK_CLAUSES]
        assert len(ids) == len(set(ids)), "Duplicate clause IDs found"

    def test_risk_levels_valid(self):
        valid = {"HIGH", "MEDIUM", "LOW"}
        for clause in CUAD_HIGH_RISK_CLAUSES:
            assert (
                clause["risk_level"] in valid
            ), f"Clause {clause['id']} has invalid risk_level: {clause['risk_level']}"

    def test_keywords_non_empty(self):
        for clause in CUAD_HIGH_RISK_CLAUSES:
            assert len(clause["keywords"]) >= 1, f"Clause {clause['id']} has no keywords"

    def test_catalogue_has_minimum_clauses(self):
        assert (
            len(CUAD_HIGH_RISK_CLAUSES) >= 15
        ), f"Expected ≥15 CUAD clauses, got {len(CUAD_HIGH_RISK_CLAUSES)}"


# ─── Rule-Based Redline Helpers ───────────────────────────────────────────────


class TestRuleBasedRedlines:
    """Unit-test the deterministic redline rule engine."""

    def test_auto_renewal_increases_days(self):
        text = "Agreement shall renew unless 15 days prior notice is given."
        result = _apply_rule_based_redline(text, "auto_renewal")
        assert "60" in result

    def test_termination_adds_notice_period(self):
        text = "Company may terminate at any time."
        result = _apply_rule_based_redline(text, "termination_for_convenience")
        assert "30" in result or "thirty" in result.lower()

    def test_liability_cap_adds_exclusion(self):
        text = "Supplier liability shall not exceed the fees paid."
        result = _apply_rule_based_redline(text, "liability_cap")
        assert "gross negligence" in result.lower() or "excluding" in result.lower()

    def test_unknown_clause_id_returns_original(self):
        text = "Some clause text."
        result = _apply_rule_based_redline(text, "unknown_clause_xyz")
        assert result == text

    def test_risk_reduction_aggressive_high_to_low(self):
        r = _estimate_risk_reduction("liability_cap", "aggressive")
        assert "LOW" in r

    def test_risk_reduction_balanced_high_to_medium(self):
        r = _estimate_risk_reduction("liability_cap", "balanced")
        assert "MEDIUM" in r

    def test_word_writer_routes_to_adeu_on_track_revisions(self):
        from sidecar.masters.word_master import WordWriter
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import tempfile
        import shutil
        import os

        # Create a temp docx file with track revisions active
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, "track_test.docx")
        doc = Document()
        doc.add_paragraph("This is the original paragraph.")

        # Turn on track revisions
        settings_element = doc.settings.element
        track_revisions = OxmlElement("w:trackRevisions")
        settings_element.append(track_revisions)
        doc.save(docx_path)

        # Call apply_operations
        writer = WordWriter()
        from unittest.mock import MagicMock

        context = MagicMock()
        ops = [
            {
                "type": "replace_paragraph",
                "paragraph_index": 0,
                "runs": [{"text": "This is the updated paragraph."}],
            }
        ]

        # Mock _adeu_installed and _adeu_sdk_available to force python_docx_fallback path
        from unittest.mock import patch

        with (
            patch("sidecar.parsers.adeu_bridge._adeu_installed", return_value=False),
            patch("sidecar.parsers.adeu_bridge._adeu_sdk_available", return_value=False),
        ):
            res = writer.apply_operations(docx_path, ops, context)

        assert "errors" in res
        assert res["applied_count"] == 1

        # Verify the saved document
        doc2 = Document(docx_path)
        # Verify that w:ins exists
        assert doc2.settings.element.find(qn("w:trackRevisions")) is not None
        # Verify w:del and w:ins exist in the XML
        body_xml = doc2._body._element.xml
        assert "w:del" in body_xml
        assert "w:ins" in body_xml

        # Clean up
        shutil.rmtree(temp_dir)


if __name__ == "__main__":
    # Allow running directly for quick validation
    import pytest

    sys.exit(pytest.main([__file__, "-v", "--tb=short"]))
